from __future__ import annotations

import os
import re
import tempfile
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\A1 COOLYEAH\MAGANG\NagariSDLC")
SRC = ROOT / "laporan" / "V.1-Laporan-KP-NagariSDLC.docx"
OUT = ROOT / "laporan" / "V.2-Laporan-KP-NagariSDLC-Final.docx"
ARCH_PNG = ROOT / "laporan" / "qa_v1" / "arsitektur-sistem-terkoreksi.png"

FONT_REG = r"C:\Windows\Fonts\arial.ttf"
FONT_BOLD = r"C:\Windows\Fonts\arialbd.ttf"


def font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)


def rounded_box(draw, xy, fill, outline="#BCD0E8", radius=24, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def centered(draw, xy, text, fnt, fill="#073F78", spacing=5):
    x1, y1, x2, y2 = xy
    box = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=spacing, align="center")
    tw, th = box[2] - box[0], box[3] - box[1]
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2), text, font=fnt, fill=fill, spacing=spacing, align="center")


def arrow(draw, start, end, color="#07589E", width=7):
    draw.line([start, end], fill=color, width=width)
    x, y = end
    draw.polygon([(x, y), (x - 20, y - 14), (x - 20, y + 14)], fill=color)


def make_architecture_diagram(path: Path):
    w, h = 2160, 900
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    blue = "#0A5EA8"
    navy = "#06437D"
    pale = "#EAF2FB"
    gold_pale = "#FFF7E5"
    gray = "#5A6D82"

    # Main layers.
    rounded_box(d, (55, 120, 545, 430), pale)
    rounded_box(d, (880, 120, 1380, 430), blue, outline=blue)
    rounded_box(d, (1690, 145, 2070, 405), navy, outline=navy)

    centered(d, (55, 130, 545, 240), "Peramban (Klien)", font(42, True), navy)
    centered(d, (75, 240, 525, 405), "React 19 · SPA\nVite 8 · Tailwind CSS 4", font(31), "#355C83", 12)
    centered(d, (880, 135, 1380, 250), "Server Aplikasi", font(43, True), "white")
    centered(d, (900, 245, 1360, 405), "Laravel 13 · PHP 8.3\nSanctum · REST API", font(32), "#DDEBFA", 12)
    centered(d, (1690, 155, 2070, 265), "Basis Data", font(42, True), "white")
    centered(d, (1710, 260, 2050, 385), "MySQL\nEloquent ORM", font(32), "#DDEBFA", 12)

    arrow(d, (565, 270), (835, 270))
    arrow(d, (835, 325), (565, 325))
    centered(d, (555, 120, 845, 245), "REST JSON", font(30, True), navy)
    centered(d, (555, 340, 845, 505), "Cookie HttpOnly (utama)\ncredentials: include\nBearer (kompatibilitas)", font(25), "#355C83", 5)

    arrow(d, (1415, 270), (1650, 270))
    arrow(d, (1650, 325), (1415, 325))

    # Security and transport notes.
    rounded_box(d, (150, 520, 790, 655), gold_pale, outline="#E6BD5F", radius=20)
    rounded_box(d, (820, 520, 1460, 655), gold_pale, outline="#E6BD5F", radius=20)
    rounded_box(d, (1490, 520, 2070, 655), gold_pale, outline="#E6BD5F", radius=20)
    centered(d, (160, 525, 780, 650), "Proteksi CSRF\nX-Requested-With", font(28, True), "#80550B", 8)
    centered(d, (830, 525, 1450, 650), "Reverb: rancangan produksi\nEnvironment saat ini: log", font(28, True), "#80550B", 8)
    centered(d, (1500, 525, 2060, 650), "Document Vault\nBerkas + masking nama", font(28, True), "#80550B", 8)

    rounded_box(d, (155, 715, 2005, 820), "#F4F7FB", outline="#D8E4F2", radius=16)
    centered(d, (170, 720, 1990, 815), "Format respons API seragam: { status, message, data, meta? }   ·   CORS: CORS_ALLOWED_ORIGINS   ·   Rate limiting: throttle", font(27), gray)
    im.save(path, format="PNG", optimize=True)


def set_run_font(run, name="Times New Roman", size=12, bold=None, italic=None, color=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_text(paragraph, text, *, bold=False, italic=False, align=None):
    paragraph.clear()
    r = paragraph.add_run(text)
    set_run_font(r, bold=bold, italic=italic)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def find_para(doc, startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise KeyError(startswith)


def set_para_by_prefix(doc, prefix, replacement):
    set_text(find_para(doc, prefix), replacement)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        set_text(p, text)
    return p


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_field(paragraph, instruction, display="1"):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run = OxmlElement("w:r")
        run.append(node)
        paragraph._p.append(run)


def make_pageref_entry(paragraph, label, bookmark):
    paragraph.clear()
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    r = paragraph.add_run(label)
    set_run_font(r)
    r = paragraph.add_run("\t")
    set_run_font(r)
    add_field(paragraph, f"PAGEREF {bookmark} \\h", "1")


def label_matches(text, label):
    return re.match(rf"^{re.escape(label)}(?:\s|$)", text.strip()) is not None


def restart_numbered_group(doc, first_prefix, count):
    paragraphs = doc.paragraphs
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith(first_prefix))
    group = paragraphs[start : start + count]
    old_num_pr = group[0]._p.pPr.numPr
    old_num_id = old_num_pr.numId.val
    numbering = doc.part.numbering_part.element
    old_num = next(n for n in numbering.findall(qn("w:num")) if int(n.get(qn("w:numId"))) == old_num_id)
    abstract_id = old_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    new_num.append(override)
    numbering.append(new_num)
    for p in group:
        p._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = new_num_id


ITALIC_TERMS = [
    "Software Development Life Cycle",
    "System Integration Test",
    "User Acceptance Test",
    "Single Page Application",
    "Quality Assurance",
    "Role-Based Access Control",
    "Object Relational Mapping",
    "segregation of duties",
    "penetration testing",
    "state machine",
    "Backend",
    "Frontend",
    "fullstack",
    "back end",
    "front end",
    "realtime",
    "build",
    "go-live",
    "sign-off",
    "rollback",
    "legacy",
    "approver",
    "Change Request",
    "resolved",
    "display",
]


def italicize_terms(paragraph):
    text = paragraph.text
    if not text or paragraph.style.name.startswith("Heading"):
        return
    pattern = re.compile("(" + "|".join(re.escape(x) for x in sorted(ITALIC_TERMS, key=len, reverse=True)) + ")", re.I)
    parts = pattern.split(text)
    if len(parts) == 1:
        return
    paragraph.clear()
    for part in parts:
        if not part:
            continue
        r = paragraph.add_run(part)
        set_run_font(r, italic=bool(pattern.fullmatch(part)))


def style_document(doc):
    section = doc.sections[0]
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name in ("List Paragraph", "Caption", "TOC 1", "TOC 2", "TOC 3"):
        if name in styles:
            st = styles[name]
            st.font.name = "Times New Roman"
            st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            st.font.size = Pt(12)
            st.paragraph_format.line_spacing = 1.15

    for name, size, before, after, align in [
        ("Heading 1", 14, 6, 12, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 12, 10, 5, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, 8, 4, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15
        st.paragraph_format.alignment = align
        st.paragraph_format.keep_with_next = True
        if name == "Heading 1":
            st.paragraph_format.page_break_before = True

    # Ensure all table geometry is fixed and typography is consistent.
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        set_run_font(r, size=12)


def replace_architecture_image(docx_path: Path, png_path: Path):
    fd, tmp_name = tempfile.mkstemp(suffix=".docx", dir=str(docx_path.parent))
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, "r") as src_zip:
            with zipfile.ZipFile(tmp_name, "w", zipfile.ZIP_DEFLATED) as dst_zip:
                for item in src_zip.infolist():
                    data = png_path.read_bytes() if item.filename == "word/media/image2.png" else src_zip.read(item.filename)
                    dst_zip.writestr(item, data)
        os.replace(tmp_name, docx_path)
    finally:
        if os.path.exists(tmp_name):
            os.unlink(tmp_name)


def main():
    make_architecture_diagram(ARCH_PNG)
    doc = Document(SRC)
    doc.core_properties.title = "Laporan Kerja Praktek NagariSDLC"
    doc.core_properties.subject = "Pengembangan Sistem Informasi Tata Kelola SDLC Berbasis Web pada Bank Nagari"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    # Cover and front matter.
    set_para_by_prefix(doc, "[ LOGO UNIVERSITAS", "[ LOGO UNIVERSITAS ANDALAS - lebar 3 cm x tinggi 3,5 cm ]")
    set_para_by_prefix(
        doc,
        "Proses tata kelola siklus hidup",
        "Proses tata kelola siklus hidup pengembangan perangkat lunak (Software Development Life Cycle, SDLC) di lingkungan perbankan menuntut kontrol yang ketat, keterlacakan, serta pemisahan tugas yang jelas antarunit. Kerja Praktek ini menghasilkan sistem informasi berbasis web untuk mengelola siklus hidup proyek perangkat lunak pada PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari), mulai dari pengajuan kebutuhan, analisis, pengembangan, System Integration Test (SIT), User Acceptance Test (UAT), Quality Assurance (QA), pengujian keamanan siber, hingga gerbang rilis ke produksi. Sistem dibangun dengan arsitektur tiga lapis: Single Page Application menggunakan React 19, server aplikasi Laravel 13 pada PHP 8.3, serta basis data MySQL. Autentikasi menggunakan token Sanctum melalui cookie HttpOnly sebagai jalur utama SPA dan header Bearer sebagai jalur kompatibilitas. Inti sistem berupa state machine proyek dengan 27 status yang seluruh transisinya terpusat sehingga setiap perubahan tercatat, disiarkan, dan menghasilkan notifikasi. Kontrol akses menerapkan 12 peran pengguna. Mahasiswa berperan sebagai pengembang fullstack yang menangani back end dan front end. Snapshot kualitas 26 Agustus 2026 mencatat 236 pengujian otomatis dengan 1.467 asersi yang seluruhnya lulus, ESLint tanpa error, dan build produksi Vite yang berhasil.",
    )

    # Chapter I.
    set_para_by_prefix(
        doc,
        "Industri perbankan merupakan",
        "Industri perbankan menempatkan keandalan, keamanan, dan keterlacakan sebagai kebutuhan penting dalam pengembangan perangkat lunak internal. Kebutuhan tersebut menuntut penerapan siklus hidup pengembangan perangkat lunak (Software Development Life Cycle, SDLC) yang terdisiplin serta pemisahan tugas (segregation of duties) yang jelas antarunit kerja.",
    )
    set_para_by_prefix(
        doc,
        "PT Bank Pembangunan Daerah Sumatera Barat, yang dikenal",
        "PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari) mengembangkan proyek perangkat lunak yang melibatkan perencanaan dan analisis, pengembangan, Quality Assurance (QA), keamanan siber, serta manajemen teknologi informasi. Sistem yang membakukan alur kerja dibutuhkan agar status, tanggung jawab, dan riwayat perubahan setiap proyek dapat ditelusuri secara konsisten.",
    )
    set_para_by_prefix(
        doc,
        "Berangkat dari kondisi tersebut",
        "Berdasarkan kebutuhan tersebut, dikembangkan NagariSDLC, yaitu aplikasi web tata kelola SDLC internal Bank Nagari. Sistem membakukan alur pengajuan, analisis, pengembangan, SIT, UAT, QA, keamanan siber, dan rilis melalui state machine terpusat, kontrol akses berbasis peran, serta pencatatan riwayat. Laporan ini memaparkan hasil pengembangan selama Kerja Praktek dari sudut pandang mahasiswa sebagai pengembang fullstack.",
    )

    # Chapter II: remove unverified institutional claims.
    set_para_by_prefix(
        doc,
        "PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari) merupakan",
        "PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari) merupakan instansi tempat pelaksanaan Kerja Praktek. Kegiatan dilaksanakan pada Divisi Teknologi Informasi di Padang. [ Lengkapi alamat resmi, sejarah singkat, dasar hukum, dan profil terkini berdasarkan dokumen resmi instansi. ]",
    )
    set_para_by_prefix(
        doc,
        "Kerja Praktek dilaksanakan pada Divisi Teknologi Informasi",
        "Pada proyek NagariSDLC, mahasiswa ditempatkan sebagai pengembang fullstack yang menangani back end dan front end. Uraian struktur resmi Divisi Teknologi Informasi tetap perlu disesuaikan dengan dokumen organisasi Bank Nagari.",
    )
    set_para_by_prefix(
        doc,
        "[ Bagian ini diisi dengan visi",
        "[ Isi visi dan misi resmi PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari) berdasarkan dokumen resmi instansi. ]",
    )
    set_para_by_prefix(
        doc,
        "Struktur organisasi unit tempat",
        "Gambar 2.1 menyajikan struktur representatif unit Teknologi Informasi berdasarkan peran yang digunakan dalam NagariSDLC dan menandai posisi mahasiswa sebagai developer pada fungsi Pengembangan. Bagan ini bukan pengganti struktur organisasi resmi; nama unit dan jabatan perlu disesuaikan dengan dokumen perusahaan.",
    )

    # Exact eight-week schedule from the user-provided specification.
    schedule = [
        "Orientasi, pengenalan proses tata kelola SDLC bank, penyiapan lingkungan (PHP 8.3, Laravel 13, React 19, MySQL).",
        "Analisis kebutuhan, penelusuran basis kode, pemahaman model data, peran, dan alur proses.",
        "Backend: model, migrasi, RBAC, autentikasi Sanctum.",
        "Backend: mesin status (ProjectWorkflowService), riwayat status, notifikasi.",
        "Backend: modul SIT/UAT, jalur QA & keamanan siber, gerbang mutu rilis.",
        "Frontend: dasbor, papan Kanban, modul registrasi & manajemen pengguna.",
        "Frontend: wizard SIT/UAT, halaman tugas QA/Siber, matriks persetujuan UAT.",
        "Pengujian menyeluruh (PHPUnit, ESLint, build), perbaikan akhir, dokumentasi, finalisasi.",
    ]
    for i, text in enumerate(schedule, start=1):
        doc.tables[1].cell(i, 0).text = str(i)
        doc.tables[1].cell(i, 1).text = text

    # Chapter III: strengthen literature review without adding project facts.
    set_para_by_prefix(
        doc,
        "SDLC adalah kerangka kerja",
        "SDLC merupakan kerangka kerja yang membagi pengembangan perangkat lunak ke dalam tahap terstruktur, antara lain analisis kebutuhan, perancangan, implementasi, pengujian, dan pemeliharaan (Sommerville, 2016; Pressman dan Maxim, 2020). Pembagian tahap membantu organisasi mengendalikan kualitas, risiko, dan keterlacakan. Dalam NagariSDLC, prinsip tersebut diterapkan sebagai alur status yang membakukan perpindahan tanggung jawab antarperan.",
    )
    set_para_by_prefix(
        doc,
        "System Integration Test (SIT) adalah",
        "System Integration Test (SIT) memverifikasi bahwa komponen sistem bekerja bersama sebagai satu kesatuan, sedangkan User Acceptance Test (UAT) menilai penerimaan sistem dari sudut pandang kebutuhan pengguna. Pada NagariSDLC, SIT dan UAT masing-masing terdiri atas tiga tahap. UAT mencakup skenario, peserta sekaligus matriks approver, undangan, eksekusi per skenario, dan persetujuan final individual.",
    )
    set_para_by_prefix(
        doc,
        "Penjaminan mutu (Quality Assurance",
        "Quality Assurance (QA) mendukung pemenuhan kualitas melalui pengujian dan verifikasi hasil. Pengujian keamanan siber, termasuk penetration testing, digunakan untuk menemukan kelemahan keamanan sebelum rilis (OWASP Foundation, tanpa tahun). Pada NagariSDLC, QA dan keamanan siber menjadi dua jalur paralel dan independen yang masing-masing melalui Pengajuan, Disposisi, Laporan, dan Sign-off.",
    )
    set_para_by_prefix(
        doc,
        "Aplikasi web modern umum",
        "Aplikasi web modern dapat memisahkan antarmuka pengguna, layanan aplikasi, dan basis data. Single Page Application (SPA) berkomunikasi dengan server melalui REST berbasis JSON, sesuai prinsip antarmuka jaringan yang dipaparkan Fielding (2000). NagariSDLC menggunakan React 19 pada klien, Laravel 13 pada server, dan MySQL melalui Eloquent ORM. Token Sanctum dikirim melalui cookie HttpOnly sebagai jalur utama SPA dengan credentials: 'include'; header Authorization: Bearer dipertahankan untuk kompatibilitas klien lama, pengujian, dan tautan approver eksternal. SPA tidak menyimpan token di localStorage, sedangkan header X-Requested-With diwajibkan sebagai proteksi CSRF.",
    )
    set_para_by_prefix(
        doc,
        "Kontrol akses berbasis peran",
        "Role-Based Access Control (RBAC) mengatur hak akses menurut peran, sedangkan state machine membatasi transisi status.",
    )

    # Chapter IV: correct authentication, realtime status, and data model.
    set_para_by_prefix(
        doc,
        "Sistem menerapkan arsitektur tiga lapis",
        "Sistem menerapkan arsitektur tiga lapis. Lapis klien berupa SPA React 19 yang dibangun dengan Vite 8 dan Tailwind CSS 4. Lapis server berupa Laravel 13 pada PHP 8.3 dengan Laravel Sanctum dan REST API. Lapis data berupa MySQL yang diakses melalui Eloquent ORM. Token Sanctum dikirim melalui cookie HttpOnly nagari_sdlc_token sebagai jalur utama dengan credentials: 'include', sedangkan header Authorization: Bearer menjadi jalur kompatibilitas. localStorage hanya memuat profil pengguna dan waktu terbit/kedaluwarsa, bukan token. Proteksi CSRF menggunakan header X-Requested-With. Respons API mengikuti format { status, message, data, meta? }, CORS dikendalikan CORS_ALLOWED_ORIGINS, dan sejumlah route memakai rate limiting. Reverb merupakan rancangan transport realtime produksi; pada environment snapshot, BROADCAST_CONNECTION=log sehingga Reverb belum aktif. Document Vault menangani unggah berkas dan masking nama berkas. Arsitektur disajikan pada Gambar 4.1.",
    )

    role_rows = {
        "super_admin": ("Administrator Sistem", "Manajemen TI", "Administrasi penuh sistem: pengguna, peran, dan master data."),
        "head_of_it": ("Pimpinan Teknologi Informasi", "Manajemen TI", "Persetujuan strategis dan gerbang go-live ke produksi."),
        "lead_group": ("Pimpinan Grup", "Perencanaan & QA", "Pimpinan grup perencanaan dan analisis."),
        "analyst": ("Analis", "Perencanaan & QA", "Analisis kebutuhan dan kelayakan."),
        "development_lead": ("Pimpinan Pengembangan", "Pengembangan", "Pimpinan tim pengembang."),
        "project_manager": ("Manajer Proyek", "Pengembangan", "Perencanaan dan pemantauan proyek."),
        "developer": ("Pengembang", "Pengembangan", "Implementasi perangkat lunak; posisi mahasiswa Kerja Praktek."),
        "qa_lead": ("Pimpinan QA", "Perencanaan & QA", "Disposisi dan sign-off QA."),
        "qa_tester": ("Penguji QA", "Perencanaan & QA", "Eksekusi pengujian QA."),
        "cyber_lead": ("Pimpinan Keamanan Siber", "Keamanan Siber", "Disposisi dan sign-off keamanan siber."),
        "pentester": ("Penguji Penetrasi", "Keamanan Siber", "Pengujian penetrasi."),
        "business_user": ("Pengguna Bisnis", "Pemohon", "Pemohon kebutuhan dan peserta UAT."),
    }
    for row in doc.tables[2].rows[1:]:
        key = row.cells[0].text.strip()
        name, group, desc = role_rows[key]
        row.cells[1].text, row.cells[2].text, row.cells[3].text = name, group, desc

    doc.tables[3].cell(0, 2).text = "Keterangan"
    for row in doc.tables[3].rows[1:]:
        row.cells[2].text = "Pengelompokan tampilan; bukan dasar otorisasi."

    set_para_by_prefix(
        doc,
        "Setelah pengembangan selesai, proyek melewati",
        "Setelah UAT pada Fase 2 selesai dan proyek mencapai DEV_COMPLETED, proyek memasuki dua jalur pengujian yang berjalan paralel dan independen, yaitu QA dan keamanan siber. Masing-masing jalur terdiri atas empat langkah: Pengajuan, Disposisi, Laporan, dan Sign-off. Kegagalan pada satu jalur membuka project return round, mengubah proyek ke RETURN_TO_DEV, dan menahan pengajuan ulang pada jalur tersebut sampai seluruh task perbaikan selesai. Alur kedua jalur disajikan pada Gambar 4.3.",
    )

    set_para_by_prefix(
        doc,
        "Pengujian SIT terdiri atas tiga tahap",
        "SIT terdiri atas tiga tahap, demikian pula UAT. Pada Tahap 1 UAT, sistem mengelola skenario, peserta sekaligus matriks approver, dan undangan. Tahap 2 menjalankan setiap skenario; temuan minor diperbaiki di tempat tanpa rollback dan tanpa perubahan status dari UAT_IN_PROGRESS, tetapi menahan penutupan UAT sampai seluruh Change Request minor berstatus resolved. Temuan mayor menjadi Change Request, memindahkan proyek ke UAT_REVISION_DEV, dilanjutkan perbaikan developer, SIT ulang menyeluruh, dan UAT ulang dari Tahap 1 setelah SIT lulus. Daftar peserta UAT tidak pernah dikosongkan. Tahap 3 memuat persetujuan final individual. Alur ini disajikan pada Gambar 4.4.",
    )

    gate = find_para(doc, "4.1.6  Gerbang Rilis")
    gate.text = "4.1.7  Gerbang Rilis"
    h = gate.insert_paragraph_before("4.1.6  Matriks Persetujuan UAT")
    h.style = "Heading 3"
    p = gate.insert_paragraph_before(
        "Matriks persetujuan UAT memuat tujuh peran approver: requester, requester_group_lead, requester_division_lead, developer, analyst_pm, development_group_lead, dan technology_division_lead. Persetujuan pemohon berada di dalam matriks melalui peran requester, bukan sebagai langkah terpisah."
    )
    p.style = "Normal"
    p = gate.insert_paragraph_before(
        "Enam peran merupakan required single roles, yaitu requester, requester_group_lead, requester_division_lead, analyst_pm, development_group_lead, dan technology_division_lead. Sistem mendukung mode INTERNAL_ACCOUNT dan EXTERNAL_LINK; requester_division_lead menggunakan external link, sedangkan peran lain umumnya memakai akun internal. Approval aktif disimpan pada uat_approval_rounds dan uat_approvers. Putaran yang tidak berlaku lagi ditandai superseded, tidak dihapus, sehingga audit trail tetap tersimpan."
    )
    p.style = "Normal"

    find_para(doc, "4.1.7  Model Data").text = "4.1.8  Model Data"
    find_para(doc, "4.1.8  Antarmuka Aplikasi").text = "4.1.9  Antarmuka Aplikasi"
    set_para_by_prefix(
        doc,
        "Model data berpusat pada entitas projects",
        "Model data memuat 25 tabel: 17 tabel domain dan 8 tabel bawaan framework. Tabel projects menjadi hub yang menautkan entitas anak. Kolom pentingnya meliputi created_by, pm_id, analyst_id, status, qa_status, cyber_status, dan sit_uat_data berbentuk JSON. Relasi anak ke projects menerapkan CASCADE, foreign key audit ke users menerapkan RESTRICT, dan foreign key penugasan menerapkan SET NULL. Diagram relasi entitas ringkas disajikan pada Gambar 4.5.",
    )
    model_para = find_para(doc, "Model data memuat 25 tabel")
    p = insert_after(
        model_para,
        "Tujuh belas tabel domain terdiri atas users, roles, groups, divisions, projects, project_tasks, project_team_members, project_return_rounds, project_status_histories, release_requests, test_reports, uat_approval_rounds, uat_approvers, document_vaults, chat_messages, activity_logs, dan notifications. Delapan tabel bawaan terdiri atas sessions, cache, cache_locks, jobs, job_batches, failed_jobs, password_reset_tokens, dan personal_access_tokens.",
        "Normal",
    )

    set_para_by_prefix(
        doc,
        "Selama Kerja Praktek, penulis berperan",
        "Selama Kerja Praktek, penulis berperan sebagai pengembang fullstack yang menangani back end dan front end. Kegiatan mencakup pemahaman model data dan RBAC, implementasi autentikasi Sanctum, mesin status dan riwayat, notifikasi, modul SIT/UAT, jalur QA dan keamanan siber, gerbang rilis, dasbor, Kanban, wizard, matriks persetujuan UAT, registrasi, serta manajemen pengguna.",
    )
    set_para_by_prefix(
        doc,
        "Sisi back end dibangun dengan Laravel 13",
        "Back end dibangun dengan Laravel 13 pada PHP 8.3. Validasi penulisan menggunakan Form Request, sedangkan logika bisnis lintas endpoint ditempatkan pada service. Seluruh transisi status wajib melalui ProjectWorkflowService::transition() agar otorisasi, prasyarat, riwayat, siaran pembaruan, dan notifikasi ditangani secara konsisten. Autentikasi menggunakan Sanctum: cookie HttpOnly menjadi jalur utama SPA dan header Bearer dipertahankan untuk kompatibilitas. Respons API dipusatkan dalam format { status, message, data, meta? }.",
    )
    set_para_by_prefix(
        doc,
        "Sisi front end dibangun sebagai SPA",
        "Front end dibangun sebagai SPA React 19 dengan Vite 8 dan Tailwind CSS 4. Seluruh panggilan API dipusatkan pada frontend/src/services/api.js. Antarmuka menyediakan dasbor, papan Kanban, wizard tiga tahap SIT/UAT, halaman tugas QA dan keamanan siber, satu inbox persetujuan UAT internal lintas proyek, serta administrasi pengguna. Penyimpanan lokal hanya memuat profil pengguna dan waktu terbit/kedaluwarsa; token tidak disimpan di localStorage.",
    )
    set_para_by_prefix(
        doc,
        "Seluruh 236 pengujian dengan",
        "Pada snapshot 26 Agustus 2026, seluruh 236 pengujian backend dengan 1.467 asersi lulus, ESLint selesai tanpa error, dan build produksi Vite berhasil. Snapshot tersebut menunjukkan kondisi kualitas pada tanggal pemeriksaan dan bukan pengganti pengujian ulang setelah perubahan berikutnya.",
    )
    set_para_by_prefix(
        doc,
        "Kendala utama meliputi kompleksitas",
        "Kendala utama berasal dari banyaknya percabangan status, kebutuhan menjaga independensi jalur QA dan keamanan siber, pengulangan SIT/UAT akibat temuan mayor, serta kewajiban mempertahankan audit trail persetujuan. Solusinya adalah memusatkan transisi proyek pada ProjectWorkflowService, memisahkan status kedua jalur uji, menyimpan putaran persetujuan UAT aktif pada tabel khusus, menandai putaran lama sebagai superseded, dan mempertahankan daftar peserta UAT pada setiap siklus revisi. Keterbatasan environment juga dicatat: broadcasting saat ini menggunakan log, sehingga Reverb belum aktif.",
    )

    # Conclusions and recommendations must stay within the supplied facts.
    set_para_by_prefix(
        doc,
        "Menambahkan pengujian antarmuka",
        "Mengaktifkan Laravel Reverb pada environment yang dituju dengan konfigurasi BROADCAST_CONNECTION=reverb, server Reverb, dan queue worker apabila pembaruan realtime akan digunakan.",
    )
    set_para_by_prefix(
        doc,
        "Meninjau kembali penyimpanan token",
        "Menetapkan konfigurasi produksi untuk basis data, queue, storage, CORS, retensi data, pencadangan, pemantauan, dan kebijakan keamanan sebelum penerapan operasional.",
    )
    set_para_by_prefix(
        doc,
        "Menyusun dokumentasi pengguna",
        "Melengkapi data administratif, visi dan misi resmi, struktur organisasi resmi, tangkapan layar aplikasi, log kegiatan, dan lembar penilaian sebelum laporan dicetak dan dijilid.",
    )

    # References and appendices are not chapters.
    dp = find_para(doc, "BAB \nDAFTAR PUSTAKA")
    set_text(dp, "DAFTAR PUSTAKA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    dp.style = "Heading 1"
    lamp = find_para(doc, "BAB \nLAMPIRAN")
    set_text(lamp, "LAMPIRAN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    lamp.style = "Heading 1"

    refs = [
        "Fielding, R. T. (2000). Architectural Styles and the Design of Network-based Software Architectures. Disertasi doktoral, University of California, Irvine.",
        "Laravel. (2026). Laravel 13.x Documentation. https://laravel.com/docs/13.x [ Diakses: TANGGAL AKSES ].",
        "Meta Open Source. (tanpa tahun). React Documentation. https://react.dev [ Diakses: TANGGAL AKSES ].",
        "OWASP Foundation. (tanpa tahun). OWASP Web Security Testing Guide. https://owasp.org/www-project-web-security-testing-guide/ [ Diakses: TANGGAL AKSES ].",
        "Pressman, R. S., dan Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (edisi ke-9). McGraw-Hill.",
        "Sommerville, I. (2016). Software Engineering (edisi ke-10). Pearson.",
    ]
    dp_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is dp._p)
    for p, text in zip(doc.paragraphs[dp_index + 1 : dp_index + 7], refs):
        set_text(p, text)

    appendix_titles = {
        "Lampiran A": "Lampiran A - Tangkapan Layar Aplikasi",
        "Lampiran B": "Lampiran B - Ringkasan Endpoint API Representatif",
        "Lampiran C": "Lampiran C - Log Kegiatan Harian",
        "Lampiran D": "Lampiran D - Lembar Penilaian Pembimbing Lapangan",
    }
    for prefix in ("Lampiran A", "Lampiran B", "Lampiran C", "Lampiran D"):
        candidates = [p for p in doc.paragraphs if p.text.strip().startswith(prefix) and "\t" not in p.text]
        p = candidates[-1]
        set_text(p, appendix_titles[prefix], bold=True)
        p.style = "Heading 2"
        p.paragraph_format.page_break_before = prefix != "Lampiran A"
        p.paragraph_format.keep_with_next = True

    set_para_by_prefix(
        doc,
        "Tabel berikut memuat sebagian endpoint",
        "Tabel berikut memuat endpoint API representatif sesuai kontrak sistem.",
    )

    # Endpoint appendix.
    api = doc.tables[12]
    api.cell(1, 2).text = "Autentikasi; menerbitkan token Sanctum melalui cookie HttpOnly sebagai jalur utama dan mendukung header Bearer untuk kompatibilitas."
    existing = {row.cells[1].text: row for row in api.rows[1:]}
    additions = [
        ("POST", "/auth/forgot-password", "Permintaan tautan reset kata sandi."),
        ("POST", "/auth/reset-password", "Reset kata sandi."),
    ]
    for method, path, desc in additions:
        row = api.add_row()
        row.cells[0].text, row.cells[1].text, row.cells[2].text = method, path, desc

    # Normalize replacement glyphs inherited from the older draft.
    for p in doc.paragraphs:
        if "�" in p.text:
            set_text(p, p.text.replace("�", "-"))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "�" in p.text:
                        set_text(p, p.text.replace("�", "-"))

    # Captions and lists with live PAGEREF fields.
    table_caps = [
        ("Tabel 2.1", "tbl_2_1", "Tabel 2.1  Jadwal dan uraian kegiatan Kerja Praktek"),
        ("Tabel 4.1", "tbl_4_1", "Tabel 4.1  Daftar peran pengguna sistem"),
        ("Tabel 4.2", "tbl_4_2", "Tabel 4.2  Pengelompokan grup kerja"),
        ("Tabel 4.3", "tbl_4_3", "Tabel 4.3  Daftar status proyek pada mesin status"),
        ("Tabel 4.4", "tbl_4_4", "Tabel 4.4  Ringkasan hasil pengujian sistem"),
    ]
    fig_caps = [
        ("Gambar 2.1", "fig_2_1", "Gambar 2.1  Struktur organisasi representatif Divisi Teknologi Informasi"),
        ("Gambar 4.1", "fig_4_1", "Gambar 4.1  Arsitektur sistem tiga lapis"),
        ("Gambar 4.2", "fig_4_2", "Gambar 4.2  Mesin status proyek per fase"),
        ("Gambar 4.3", "fig_4_3", "Gambar 4.3  Alur jalur QA dan keamanan siber"),
        ("Gambar 4.4", "fig_4_4", "Gambar 4.4  Alur pengujian SIT dan UAT"),
        ("Gambar 4.5", "fig_4_5", "Gambar 4.5  Diagram relasi entitas (ERD) ringkas"),
        ("Gambar 4.6", "fig_4_6", "Gambar 4.6  Antarmuka dasbor (placeholder)"),
        ("Gambar 4.7", "fig_4_7", "Gambar 4.7  Antarmuka papan Kanban proyek (placeholder)"),
        ("Gambar 4.8", "fig_4_8", "Gambar 4.8  Antarmuka wizard SIT/UAT (placeholder)"),
        ("Gambar 4.9", "fig_4_9", "Gambar 4.9  Antarmuka tugas QA/Siber (placeholder)"),
        ("Gambar 4.10", "fig_4_10", "Gambar 4.10  Antarmuka matriks persetujuan UAT (placeholder)"),
        ("Gambar 4.11", "fig_4_11", "Gambar 4.11  Antarmuka administrasi dan manajemen pengguna (placeholder)"),
    ]
    bookmark_id = 100
    for prefix, name, canonical in table_caps + fig_caps:
        candidates = [p for p in doc.paragraphs if label_matches(p.text, prefix)]
        caption = next((p for p in reversed(candidates) if "\t" not in p.text), candidates[-1])
        set_text(caption, canonical, align=WD_ALIGN_PARAGRAPH.CENTER)
        caption.paragraph_format.keep_together = True
        caption.paragraph_format.keep_with_next = prefix.startswith("Tabel ")
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(6)
        add_bookmark(caption, name, bookmark_id)
        bookmark_id += 1

    list_map = {label: (bookmark, canonical) for label, bookmark, canonical in table_caps + fig_caps}
    for p in doc.paragraphs:
        stripped = p.text.strip()
        for label, (bookmark, canonical) in list_map.items():
            if label_matches(stripped, label) and "\t" in p.text:
                make_pageref_entry(p, canonical, bookmark)

    appendix_map = [
        ("Lampiran A", "app_a", "Lampiran A  Tangkapan layar aplikasi"),
        ("Lampiran B", "app_b", "Lampiran B  Ringkasan endpoint API representatif"),
        ("Lampiran C", "app_c", "Lampiran C  Log kegiatan harian"),
        ("Lampiran D", "app_d", "Lampiran D  Lembar penilaian pembimbing lapangan"),
    ]
    for prefix, bookmark, canonical in appendix_map:
        heading = [p for p in doc.paragraphs if p.style.name == "Heading 2" and p.text.strip().startswith(prefix)][-1]
        add_bookmark(heading, bookmark, bookmark_id)
        bookmark_id += 1
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix) and "\t" in p.text:
                make_pageref_entry(p, canonical, bookmark)

    # Restart independent numbered lists at 1.
    restart_numbered_group(doc, "Merancang model 12 peran", 3)
    restart_numbered_group(doc, "Sistem tata kelola SDLC berbasis web berhasil", 3)
    restart_numbered_group(doc, "Mengaktifkan Laravel Reverb", 3)

    # Center titles and captions; keep body justified.
    front_titles = {
        "HALAMAN PERNYATAAN KEASLIAN",
        "HALAMAN PENGESAHAN",
        "PEMBIMBING LAPANGAN",
        "DOSEN PEMBIMBING, PENGUJI, DAN KETUA DEPARTEMEN",
        "ABSTRAK",
        "KATA PENGANTAR",
        "DAFTAR ISI",
        "DAFTAR TABEL",
        "DAFTAR GAMBAR",
        "DAFTAR LAMPIRAN",
    }
    for p in doc.paragraphs:
        if p.text.strip() in front_titles:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, bold=True)
        elif p.style.name in ("Normal", "List Paragraph") and p.text.strip():
            if not p.text.strip().startswith(("Gambar ", "Tabel ")):
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15

    style_document(doc)

    # Apply italics to foreign terms after all content replacements.
    for p in doc.paragraphs:
        italicize_terms(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    italicize_terms(p)

    keywords = find_para(doc, "Kata kunci:")
    keyword_text = keywords.text.removeprefix("Kata kunci:")
    keywords.clear()
    label_run = keywords.add_run("Kata kunci:")
    set_run_font(label_run, bold=True)
    pattern = re.compile("(" + "|".join(re.escape(x) for x in sorted(ITALIC_TERMS, key=len, reverse=True)) + ")", re.I)
    for part in pattern.split(keyword_text):
        if part:
            r = keywords.add_run(part)
            set_run_font(r, italic=bool(pattern.fullmatch(part)))

    # Bibliography entries use a conventional hanging indent.
    dp_index = next(i for i, p in enumerate(doc.paragraphs) if p._p is dp._p)
    for p in doc.paragraphs[dp_index + 1 : dp_index + 7]:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        p.paragraph_format.space_after = Pt(6)

    # Restore table headers after typography normalization.
    for table in doc.tables:
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    set_run_font(r, size=12, bold=True)

    # Ensure Word refreshes TOC, PAGEREF, and page fields on open.
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUT)
    replace_architecture_image(OUT, ARCH_PNG)
    print(OUT)


if __name__ == "__main__":
    main()
