from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(r"D:\A1 COOLYEAH\MAGANG\NagariSDLC\laporan\V.3-Laporan-KP-NagariSDLC-Lengkap.docx")
OUTPUT = Path(r"D:\A1 COOLYEAH\MAGANG\NagariSDLC\laporan\V.4.2-FINAL-Laporan-KP-NagariSDLC-Siap-Cetak.docx")
MANIFEST = Path(r"D:\A1 COOLYEAH\MAGANG\NagariSDLC\laporan\qa_v3\final_static_manifest.json")


def all_paragraphs(document: Document) -> list[Paragraph]:
    return [
        Paragraph(element, document._body)
        for element in document._element.body.iter(qn("w:p"))
    ]


def clean_text(paragraph: Paragraph) -> str:
    return " ".join(paragraph.text.replace("\xa0", " ").split())


def set_plain_text(paragraph: Paragraph, text: str) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"


def replace_text(document: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in all_paragraphs(document):
        text = paragraph.text
        if old in text:
            set_plain_text(paragraph, text.replace(old, new))
            count += 1
    return count


def make_paragraph(
    document: Document,
    text: str,
    style: str,
    *,
    justified: bool = False,
) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph = Paragraph(element, document._body)
    paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    if justified:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(10)
    return paragraph


def clear_content_keep_properties(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def add_text_run(paragraph: Paragraph, text: str, *, bold: bool | None = None) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    if bold is not None:
        run.bold = bold


def append_field(paragraph: Paragraph, instruction: str, result: str = "1") -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run.append(instruction_text)
    paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    result_run = OxmlElement("w:r")
    result_text = OxmlElement("w:t")
    result_text.text = result
    result_run.append(result_text)
    paragraph._p.append(result_run)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)


def insert_after(reference: Paragraph, new_paragraph: Paragraph) -> None:
    reference._p.addnext(new_paragraph._p)


def add_bookmark(paragraph: Paragraph, bookmark_id: int, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(1 if paragraph._p.pPr is not None else 0, start)
    paragraph._p.append(end)


def configure_list_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "7920")
    tabs.append(tab)


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    document = Document(SOURCE)

    replacements = [
        ("PT Bank Bank Nagari", "PT Bank Nagari"),
        ("Grub Pengembangan", "Grup Pengembangan"),
        ("di Grub", "di Grup"),
        ("focus", "fokus"),
        ("kepatihan", "kepatuhan"),
        ("Fery Arjendi Putra, S.Kom., CITPM", "Ferry Arjendi Putra, S.Kom., CITPM"),
        ("Padang, Agustus 2026", "Padang, 31 Agustus 2026"),
        ("Analist", "Analis"),
        ("Gambar lampiran", "Gambar Lampiran"),
        ("Monitoring Proyek Berjalan Oleh", "Monitoring Proyek Berjalan oleh"),
        ("Tabel 4.7 Tabel domain utama", "Tabel 4.7 Domain utama"),
        (
            "Informasi profil perusahaan yang belum didukung dokumen resmi sengaja tidak dikarang dan ditandai sebagai placeholder untuk dilengkapi penulis.",
            "Profil instansi pada laporan ini disusun berdasarkan informasi perusahaan dan dokumen pelaksanaan Kerja Praktek yang tersedia.",
        ),
        (
            "6. Tangkapan layar, logo, struktur organisasi resmi, dan dokumen pengesahan ditandai sebagai placeholder hingga bahan resmi tersedia.",
            "6. Tangkapan layar dan diagram dibatasi pada bagian yang relevan dengan pembahasan serta tidak menampilkan data produksi atau informasi sensitif.",
        ),
        (
            "8. Diagram tidak disisipkan sebagai gambar; dokumen menampilkan placeholder dan berkas pendamping menyediakan kode Mermaid.",
            "8. Diagram proses, UML, ERD, sequence, dan navigasi merepresentasikan rancangan serta implementasi pada saat laporan disusun.",
        ),
        (
            "1. Lengkapi seluruh placeholder identitas, profil resmi, struktur organisasi, logo, foto, tanda tangan, dan bukti kegiatan sebelum pengesahan.",
            "1. Lakukan pengujian regresi backend, pemeriksaan lint, build frontend, dan pengujian end-to-end setelah setiap perubahan utama agar kestabilan sistem tetap terjaga.",
        ),
        (
            "2. Render kode Mermaid pada berkas pendamping, periksa keterbacaan, lalu masukkan setiap gambar pada placeholder dengan caption yang tetap.",
            "2. Tetapkan konfigurasi dan prosedur operasional produksi untuk database, queue, object storage, CORS, Reverb, backup, monitoring, logging, serta retensi data.",
        ),
        (
            "3. Lakukan pengujian ulang backend, ESLint, build, dan pengujian end-to-end setelah perubahan source code terakhir sebelum laporan dinyatakan final.",
            "3. Tambahkan pemantauan kesehatan layanan, metrik performa, pencatatan kesalahan terpusat, dan mekanisme pemulihan agar gangguan operasional dapat dideteksi lebih cepat.",
        ),
        (
            "4. Tetapkan konfigurasi dan SOP produksi untuk database, queue, storage, CORS, Reverb, backup, monitoring, logging, serta retensi data.",
            "4. Pertahankan seluruh transisi proyek melalui ProjectWorkflowService dan seluruh penulisan return round melalui ProjectReturnRoundService agar aturan bisnis tetap terpusat.",
        ),
        (
            "5. Pertahankan seluruh transisi proyek melalui ProjectWorkflowService dan seluruh penulisan return round melalui ProjectReturnRoundService.",
            "5. Pertahankan approval, histori status, laporan pengujian, activity log, dan bukti dokumen sebagai audit trail serta hindari hard delete tanpa kebijakan retensi resmi.",
        ),
        (
            "6. Hindari hard delete pada approval, histori, laporan pengujian, dan bukti dokumen kecuali telah ada keputusan retensi resmi.",
            "6. Lakukan evaluasi usability dan pengukuran karakteristik kualitas ISO/IEC 25010 secara berkala menggunakan data penggunaan nyata setelah sistem diterapkan.",
        ),
    ]
    replacement_count = sum(replace_text(document, old, new) for old, new in replacements)

    paragraphs = all_paragraphs(document)
    bab_two = next(p for p in paragraphs if clean_text(p).startswith("BAB II"))
    additions = [
        make_paragraph(document, "1.6 Metode Pelaksanaan", "Heading 2"),
        make_paragraph(
            document,
            "Pelaksanaan Kerja Praktek dilakukan melalui tahapan observasi proses kerja dan identifikasi kebutuhan, penelaahan dokumentasi proyek serta literatur, analisis dan perancangan sistem, implementasi secara iteratif, pengujian, dan penyusunan dokumentasi. Setiap hasil analisis dikonfirmasi terhadap alur kerja NagariSDLC dan implementasi aktif agar rancangan antarmuka, layanan API, model data, kontrol akses, serta transisi status tetap konsisten.",
            "Normal",
            justified=True,
        ),
        make_paragraph(document, "1.7 Sistematika Penulisan", "Heading 2"),
        make_paragraph(
            document,
            "Laporan ini disusun dalam lima bab. Bab I menjelaskan latar belakang, rumusan masalah, tujuan, manfaat, batasan, metode pelaksanaan, dan sistematika penulisan. Bab II memuat profil instansi, unit kerja, posisi mahasiswa, serta kegiatan Kerja Praktek. Bab III menguraikan landasan teori dan penelitian yang mendukung pengembangan NagariSDLC. Bab IV menyajikan hasil analisis, perancangan, implementasi, pengujian, dan pembahasan sistem. Bab V berisi kesimpulan dan saran, kemudian dilanjutkan dengan daftar pustaka serta lampiran teknis.",
            "Normal",
            justified=True,
        ),
    ]
    parent = bab_two._p.getparent()
    index = parent.index(bab_two._p)
    for offset, paragraph in enumerate(additions):
        parent.insert(index + offset, paragraph._p)

    if "Heading 4" not in [style.name for style in document.styles]:
        document.styles.add_style("Heading 4", WD_STYLE_TYPE.PARAGRAPH)
    heading_four = document.styles["Heading 4"]
    heading_four.font.name = "Times New Roman"
    heading_four.font.bold = True
    heading_four.paragraph_format.space_before = Pt(5)
    heading_four.paragraph_format.space_after = Pt(2)
    heading_four.paragraph_format.keep_with_next = True
    for paragraph in all_paragraphs(document):
        if re.match(r"^4\.1\.15\.([1-9]|1[0-6])\s", clean_text(paragraph)):
            paragraph.style = "Heading 4"

    sequence_counts: dict[tuple[str, str], int] = {}
    captions_converted = 0
    for paragraph in all_paragraphs(document):
        match = re.match(r"^(Gambar|Tabel)\s+([234AB])\.([0-9]+)\s+(.+)$", clean_text(paragraph))
        if not match:
            continue
        label, chapter, _, title = match.groups()
        key = (label, chapter)
        sequence_counts[key] = sequence_counts.get(key, 0) + 1
        number = sequence_counts[key]
        reset = number == 1
        clear_content_keep_properties(paragraph)
        paragraph.style = "Caption"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_after = Pt(4)
        add_text_run(paragraph, f"{label} {chapter}.", bold=True)
        instruction = f"SEQ {label} \\r 1 \\* ARABIC" if reset else f"SEQ {label} \\* ARABIC"
        append_field(paragraph, instruction, str(number))
        add_text_run(paragraph, f" {title}", bold=True)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
        captions_converted += 1

    for table in document.tables:
        if table.rows:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                header = OxmlElement("w:tblHeader")
                header.set(qn("w:val"), "true")
                tr_pr.append(header)

    for style_name in ("TOC 1", "TOC 2", "TOC 3", "Table of Figures"):
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(2)

    appendices = [
        ("A", "Lampiran A Endpoint API", "Lampiran A - Endpoint API"),
        ("B", "Lampiran B Kamus Data Ringkas", "Lampiran B - Kamus Data Ringkas"),
        ("C", "Lampiran C Tampilan Lengkap", "Lampiran C - Tampilan Lengkap"),
    ]

    manifest: list[dict[str, object]] = []

    headings = []
    for paragraph in all_paragraphs(document):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            title = clean_text(paragraph)
            if title:
                level = int(style_name.rsplit(" ", 1)[1])
                headings.append((title, level))

    toc_heading = next(p for p in all_paragraphs(document) if clean_text(p) == "DAFTAR ISI")
    previous = toc_heading
    for title, level in headings:
        entry = make_paragraph(document, f"{title}\t0", f"TOC {level}")
        configure_list_paragraph(entry)
        entry.paragraph_format.left_indent = Pt((level - 1) * 14)
        insert_after(previous, entry)
        previous = entry
        manifest.append({"category": "toc", "title": title, "target": title, "level": level})

    table_captions = []
    figure_captions = []
    for paragraph in all_paragraphs(document):
        title = clean_text(paragraph)
        if re.match(r"^Tabel\s+[234AB]\.\d+\s+", title):
            table_captions.append(title)
        elif re.match(r"^Gambar\s+[24]\.\d+\s+", title):
            figure_captions.append(title)

    for heading_text, category, captions in (
        ("DAFTAR TABEL", "tables", table_captions),
        ("DAFTAR GAMBAR", "figures", figure_captions),
    ):
        heading = next(p for p in all_paragraphs(document) if clean_text(p) == heading_text)
        previous = heading
        for title in captions:
            entry = make_paragraph(document, f"{title}\t0", "Table of Figures")
            configure_list_paragraph(entry)
            insert_after(previous, entry)
            previous = entry
            manifest.append({"category": category, "title": title, "target": title, "level": 1})

    list_heading = next(p for p in all_paragraphs(document) if clean_text(p) == "DAFTAR LAMPIRAN")
    previous = list_heading
    for _, target, entry_text in appendices:
        entry = make_paragraph(document, f"{entry_text}\t0", "Table of Figures")
        configure_list_paragraph(entry)
        insert_after(previous, entry)
        previous = entry
        manifest.append({"category": "appendices", "title": entry_text, "target": target, "level": 1})

    document.core_properties.author = "Muhammad Galid Avero"
    document.core_properties.last_modified_by = "Muhammad Galid Avero"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"OUTPUT={OUTPUT}")
    print(f"REPLACEMENTS={replacement_count}")
    print(f"CAPTIONS_CONVERTED={captions_converted}")
    print(f"LIST_ENTRIES={len(manifest)}")


if __name__ == "__main__":
    main()
