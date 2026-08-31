from docx import Document
import re
from zipfile import ZipFile
from lxml import etree

path = r"D:\A1 COOLYEAH\MAGANG\NagariSDLC\laporan\V.3-Laporan-KP-NagariSDLC-Lengkap.docx"
document = Document(path)

for index, paragraph in enumerate(document.paragraphs):
    text = " ".join(paragraph.text.split())
    if re.match(r"^(Gambar|Tabel|Lampiran)\b", text, re.I) or "placeholder" in text.lower():
        print(f"P{index}|{paragraph.style.name}|{text}")

print("---TABLE CELLS---")
seen_cells = set()
for table_index, table in enumerate(document.tables):
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell_key = id(cell._tc)
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                text = " ".join(paragraph.text.split())
                if re.match(r"^(Gambar|Tabel|Lampiran)\b", text, re.I) or "placeholder" in text.lower():
                    print(
                        f"T{table_index}R{row_index}C{cell_index}P{paragraph_index}|"
                        f"{paragraph.style.name}|{text}"
                    )

print("---OOXML ALL PARAGRAPHS---")
with ZipFile(path) as package:
    root = etree.fromstring(package.read("word/document.xml"))
ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
for index, paragraph in enumerate(root.xpath(".//w:p", namespaces=ns)):
    text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns)).strip()
    if re.match(r"^(Gambar|Tabel|Lampiran)\b", text, re.I) or "placeholder" in text.lower():
        styles = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=ns)
        print(f"X{index}|{styles[0] if styles else ''}|{text}")

print("---HEADINGS---")
for index, paragraph in enumerate(root.xpath(".//w:p", namespaces=ns)):
    styles = paragraph.xpath("./w:pPr/w:pStyle/@w:val", namespaces=ns)
    if styles and styles[0] in {"2", "3", "4", "Heading1", "Heading2", "Heading3", "Heading4"}:
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns)).strip()
        print(f"H{index}|{styles[0]}|{text}")

print("---SELECTED BODY PARAGRAPHS---")
for index in list(range(44, 86)) + list(range(388, 398)):
    if index < len(document.paragraphs):
        paragraph = document.paragraphs[index]
        print(f"S{index}|{paragraph.style.name}|{' '.join(paragraph.text.split())}")
