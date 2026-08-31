from __future__ import annotations

from pathlib import Path
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\A1 COOLYEAH\MAGANG\NagariSDLC")
OUT = ROOT / "laporan" / "V.3-Laporan-KP-NagariSDLC-Lengkap.docx"
PLACEHOLDER_DIR = ROOT / "laporan" / "qa_v3" / "placeholder_assets"

FONT = "Times New Roman"
BLACK = "000000"
BLUE = "00529C"
LIGHT_BLUE = "DCEAF7"
LIGHT_GRAY = "F2F2F2"
PLACEHOLDER_FILL = "F7F7F7"
CONTENT_DXA = 7938  # 14 cm: A4 21 cm - margin kiri 4 cm - kanan 3 cm.


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, *, size=12, bold=None, italic=None, color=BLACK, font=FONT) -> None:
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text_run(paragraph, text, *, bold=False, italic=False, size=12, color=BLACK):
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic, color=color)
    return run


def style_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, after=0, before=0, line=1.15):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.first_line_indent = Cm(1.27) if indent else Cm(0)
    return paragraph


def add_para(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, after=0, before=0, line=1.15, bold=False, italic=False, keep=False):
    p = doc.add_paragraph()
    style_paragraph(p, align=align, indent=indent, after=after, before=before, line=line)
    if text:
        add_text_run(p, text, bold=bold, italic=italic)
    p.paragraph_format.keep_with_next = keep
    return p


def add_mixed_para(doc, parts, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, after=0, line=1.15):
    p = doc.add_paragraph()
    style_paragraph(p, align=align, indent=indent, after=after, line=line)
    for text, bold, italic in parts:
        add_text_run(p, text, bold=bold, italic=italic)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_field(paragraph, instruction: str, display="1") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    run_text = OxmlElement("w:r")
    run_text.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.append(begin)
    run.append(instr)
    run.append(separate)
    paragraph._p.append(run)
    paragraph._p.append(run_text)
    paragraph._p.append(end)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def set_page_number_format(section, fmt: str, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))


def set_footer_page_number(section, *, different_first=False) -> None:
    section.different_first_page_header_footer = different_first
    if different_first:
        section.first_page_footer.paragraphs[0].clear()
    p = section.footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    add_field(p, "PAGE", "1")
    for run in p.runs:
        set_run(run, size=12)


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.right_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_chapter_heading(doc, roman: str, title: str):
    if roman != "I":
        page_break(doc)
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = False
    p.paragraph_format.keep_with_next = True
    add_text_run(p, f"BAB {roman}\n{title}", size=14, bold=True)
    return p


def add_heading(doc, text: str, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level == 2 else 5)
    p.paragraph_format.space_after = Pt(2)
    add_text_run(p, text, size=12, bold=True)
    return p


def add_numbering_definition(doc, kind="decimal") -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"{(abstract_id * 2654435761) & 0xFFFFFFFF:08X}")
    abstract.append(tmpl)
    name = OxmlElement("w:name")
    name.set(qn("w:val"), f"NagariListGroup{abstract_id}")
    abstract.append(name)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    unique_pos = 720 + abstract_id
    tab.set(qn("w:pos"), str(unique_pos))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(unique_pos))
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def add_list(doc, items: list[str], *, kind="decimal") -> None:
    # Use visible prefixes instead of Word's list state. Microsoft Word merges
    # separately generated numbering definitions while refreshing fields,
    # which made later lists continue from earlier ones (for example, 6–11).
    # A hanging-indent prefix keeps the rendered report deterministic and lets
    # every semantic list restart at 1 as required by academic formatting.
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        style_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False, after=0, line=1.15)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-0.64)
        prefix = f"{index}." if kind == "decimal" else "•"
        add_text_run(p, f"{prefix}\t{item}")


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], caption: str | None = None, bookmark: str | None = None, bookmark_id: int | None = None):
    if caption:
        cap = add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=4, line=1.0, keep=True)
        if bookmark and bookmark_id is not None:
            add_bookmark(cap, bookmark, bookmark_id)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        add_text_run(p, header, size=11, bold=True)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = ""
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_text_run(p, str(value), size=11)
        set_table_geometry(table, widths)
    add_para(doc, "", indent=False, line=1.0)
    return table


def create_placeholder_asset(bookmark: str, instruction: str, code_id: str | None, lines: int) -> Path:
    """Create a single unsplittable placeholder graphic for Word."""
    PLACEHOLDER_DIR.mkdir(parents=True, exist_ok=True)
    path = PLACEHOLDER_DIR / f"{bookmark}.png"
    width = 1400
    height = 320 + lines * 20
    image = Image.new("RGB", (width, height), "#F7F7F7")
    draw = ImageDraw.Draw(image)
    draw.rectangle((3, 3, width - 4, height - 4), outline="#808080", width=3)

    regular = r"C:\Windows\Fonts\times.ttf"
    bold = r"C:\Windows\Fonts\timesbd.ttf"
    italic = r"C:\Windows\Fonts\timesi.ttf"
    title_font = ImageFont.truetype(bold, 40)
    instruction_font = ImageFont.truetype(italic, 30)
    code_font = ImageFont.truetype(bold, 29)

    def centered(text: str, y: int, font, fill: str, wrap: int) -> int:
        for line in textwrap.wrap(text, width=wrap) or [""]:
            box = draw.textbbox((0, 0), line, font=font)
            line_width = box[2] - box[0]
            draw.text(((width - line_width) / 2, y), line, font=font, fill=fill)
            y += (box[3] - box[1]) + 10
        return y

    y = 48
    y = centered("PLACEHOLDER GAMBAR / DIAGRAM", y, title_font, "#00529C", 60)
    y += 8
    y = centered(instruction, y, instruction_font, "#202020", 78)
    if code_id:
        y += 8
        centered(f"Kode Mermaid: {code_id} (lihat berkas pendamping)", y, code_font, "#202020", 72)
    image.save(path, dpi=(220, 220))
    return path


def add_figure_placeholder(doc, caption: str, bookmark: str, bookmark_id: int, instruction: str, code_id: str | None = None, lines=7):
    # Keep the atomic picture and its caption in one non-splittable borderless
    # table row. This avoids Word placing an inline picture baseline above the
    # top margin when a placeholder happens to fall at a page boundary.
    asset = create_placeholder_asset(bookmark, instruction, code_id, lines)
    if code_id == "D-10":
        # This diagram follows another large figure. Word otherwise begins the
        # row in the final sliver of that page and clips it at the next page.
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    keep_table_row_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    picture = p.add_run().add_picture(str(asset), width=Cm(14))
    picture._inline.docPr.set("descr", f"Placeholder {caption}. {instruction}")
    picture._inline.docPr.set("title", caption)
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.keep_together = True
    add_text_run(cap, caption)
    cap.paragraph_format.keep_with_next = False
    add_bookmark(cap, bookmark, bookmark_id)
    return cap


def add_list_entry(doc, label: str, bookmark: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    add_text_run(p, label)
    add_text_run(p, "\t")
    add_field(p, f"PAGEREF {bookmark} \\h", "0")
    return p


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for name, size, before, after in (("Heading 1", 14, 0, 12), ("Heading 2", 12, 8, 2), ("Heading 3", 12, 5, 2)):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True


def add_front_title(doc, title: str):
    p = add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=12, line=1.15, bold=True)
    p.paragraph_format.keep_with_next = True
    return p


def build():
    doc = Document()
    doc.core_properties.title = "Laporan Kerja Praktek NagariSDLC"
    doc.core_properties.subject = "Pengembangan Sistem Informasi Tata Kelola SDLC Berbasis Web pada Bank Nagari"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    style_document(doc)
    configure_section(doc.sections[0])

    # Sampul akademik: editorial_cover dengan override academic_unand_cover dari PDF referensi.
    for _ in range(2):
        add_para(doc, "", indent=False, line=1.0)
    add_para(doc, "LAPORAN KERJA PRAKTEK", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=14, line=1.15, bold=True)
    add_para(doc, "PENGEMBANGAN SISTEM INFORMASI TATA KELOLA", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.15, bold=True)
    add_para(doc, "SIKLUS HIDUP PENGEMBANGAN PERANGKAT LUNAK (SDLC)", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.15, bold=True)
    add_para(doc, "BERBASIS WEB PADA PT BANK PEMBANGUNAN DAERAH SUMATERA BARAT", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.15, bold=True)
    add_para(doc, "(BANK NAGARI)", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=10, line=1.15, bold=True)
    add_para(doc, "Periode: [ TANGGAL MULAI ] - [ TANGGAL SELESAI ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=12, line=1.15)
    add_para(doc, "Oleh", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=8, line=1.0)
    add_para(doc, "[ NAMA MAHASISWA ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0, bold=True)
    add_para(doc, "[ NIM ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=12, line=1.0, bold=True)
    add_para(doc, "Dosen Pembimbing", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0)
    add_para(doc, "[ NAMA DOSEN PEMBIMBING, GELAR ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0, bold=True)
    add_para(doc, "NIP/NIKU. [ NOMOR IDENTITAS ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=10, line=1.0, bold=True)
    add_para(doc, "[ LOGO UNIVERSITAS ANDALAS - sisipkan di sini ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, after=12, line=1.0, bold=True)
    add_para(doc, "DEPARTEMEN INFORMATIKA", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0, bold=True)
    add_para(doc, "FAKULTAS TEKNOLOGI INFORMASI", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0, bold=True)
    add_para(doc, "UNIVERSITAS ANDALAS", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0, bold=True)
    add_para(doc, "[ TAHUN ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, line=1.0, bold=True)
    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front_section)
    front_section.header.is_linked_to_previous = False
    front_section.footer.is_linked_to_previous = False
    set_page_number_format(front_section, "lowerRoman", 1)
    set_footer_page_number(front_section)

    add_front_title(doc, "SURAT PERNYATAAN\nLAPORAN KERJA PRAKTEK")
    add_para(doc, "Yang bertanda tangan di bawah ini:", indent=False)
    add_table(doc, ["Keterangan", "Data"], [["Nama", "[ NAMA MAHASISWA ]"], ["NIM", "[ NIM ]"]], [1900, 6038])
    add_para(doc, "Menyatakan dengan sesungguhnya bahwa:", indent=False)
    add_list(doc, [
        "Laporan Kerja Praktek ini disusun berdasarkan kegiatan, dokumentasi, dan sumber yang dapat dipertanggungjawabkan.",
        "Data instansi yang dicantumkan akan disesuaikan dengan dokumen resmi dan persetujuan instansi tempat Kerja Praktek.",
        "Laporan ini bukan hasil plagiarisme dan seluruh sumber yang digunakan dicantumkan pada daftar pustaka.",
        "Bagian yang masih diberi tanda kurung siku merupakan placeholder yang wajib dilengkapi penulis sebelum pengesahan.",
        "Kode, tangkapan layar, dan bukti kegiatan yang bersifat internal hanya akan dimasukkan setelah memperoleh izin yang diperlukan.",
    ])
    add_para(doc, "Padang, [ TANGGAL PENGESAHAN ]", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, before=14)
    add_para(doc, "Yang membuat pernyataan,", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    for _ in range(3):
        add_para(doc, "", indent=False, line=1.0)
    add_para(doc, "[ NAMA MAHASISWA ]\n[ NIM ]", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, bold=True, line=1.0)
    page_break(doc)

    add_front_title(doc, "LEMBAR PENGESAHAN\nLAPORAN KERJA PRAKTEK")
    add_para(doc, "PENGEMBANGAN SISTEM INFORMASI TATA KELOLA SIKLUS HIDUP PENGEMBANGAN PERANGKAT LUNAK (SDLC) BERBASIS WEB PADA PT BANK PEMBANGUNAN DAERAH SUMATERA BARAT (BANK NAGARI)", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, line=1.15)
    add_para(doc, "Periode: [ TANGGAL MULAI ] - [ TANGGAL SELESAI ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    add_para(doc, "Oleh\n[ NAMA MAHASISWA ]\n[ NIM ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, line=1.15)
    add_para(doc, "Mengetahui,", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, before=10)
    add_table(doc, ["Dosen Penguji", "Dosen Pembimbing"], [["\n\n\n[ NAMA DOSEN PENGUJI ]\nNIP/NIKU. [ NOMOR ]", "\n\n\n[ NAMA DOSEN PEMBIMBING ]\nNIP/NIKU. [ NOMOR ]"]], [3969, 3969])
    add_para(doc, "Ketua Departemen Informatika", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    for _ in range(3):
        add_para(doc, "", indent=False, line=1.0)
    add_para(doc, "[ NAMA KETUA DEPARTEMEN ]\nNIP. [ NOMOR ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True, line=1.0)
    page_break(doc)

    add_front_title(doc, "HALAMAN PENGESAHAN PEMBIMBING LAPANGAN")
    add_para(doc, "[ SISIPKAN HASIL PEMINDAIAN LEMBAR PENGESAHAN YANG TELAH DITANDATANGANI PEMBIMBING LAPANGAN ]", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
    for _ in range(16):
        add_para(doc, "", indent=False, line=1.0)
    page_break(doc)

    add_front_title(doc, "ABSTRAK")
    add_para(doc, "Tata kelola pengembangan perangkat lunak di lingkungan perbankan memerlukan alur yang terdokumentasi, pembagian wewenang yang tegas, kontrol keamanan, serta jejak audit yang dapat ditelusuri. Kerja Praktek ini berfokus pada pengembangan NagariSDLC, yaitu sistem informasi internal berbasis web untuk mengelola proyek teknologi sejak pengajuan kebutuhan sampai rilis ke produksi. Sistem dibangun menggunakan Laravel 13 pada PHP 8.3 sebagai REST API, React 19 dengan Vite 8 dan Tailwind CSS 4 sebagai antarmuka Single Page Application, serta MySQL melalui Eloquent ORM sebagai lapisan data. Metode pekerjaan meliputi studi dokumentasi dan literatur, analisis kebutuhan, penelusuran proses bisnis dan source code, perancangan model peran dan mesin status, implementasi modul proyek dan task, pengembangan wizard System Integration Test dan User Acceptance Test, pengujian paralel Quality Assurance dan keamanan siber, pengelolaan dokumen, persetujuan individual, serta gerbang rilis. Landasan analisis menggunakan standar ISO/IEC/IEEE mengenai lifecycle, testing, dan kualitas produk; panduan NIST dan OWASP mengenai secure development serta security testing; artikel jurnal mengenai agile, continuous software engineering, dan DevOps; serta buku rekayasa perangkat lunak dan arsitektur. Hasil pengembangan membentuk alur SDLC terpusat dengan 12 peran pengguna, 27 status proyek, pengujian SIT/UAT bertahap, dua jalur QA dan keamanan siber yang independen, putaran pengembalian yang tidak menghapus histori, dan quality gate sebelum produksi. Analisis menunjukkan bahwa pemusatan state transition, project scope, approval round, transaksi, dan audit trail mendukung functional suitability, security, reliability, serta maintainability, sementara performance, usability formal, CI/CD, monitoring, backup, dan kesiapan operasi produksi masih memerlukan verifikasi serta keputusan lanjutan. Snapshot kualitas tanggal 26 Agustus 2026 mencatat 236 pengujian backend dengan 1.467 asersi yang lulus, pemeriksaan ESLint tanpa error, dan build produksi Vite yang berhasil. Laporan ini menyediakan placeholder untuk diagram dan tangkapan layar; kode Mermaid disiapkan secara terpisah agar seluruh diagram dapat dirender ulang secara konsisten.")
    add_mixed_para(doc, [("Kata kunci: ", True, False), ("audit trail", False, True), (", Laravel, React, RBAC, SDLC, SIT, UAT.", False, False)], indent=False)
    page_break(doc)

    add_front_title(doc, "KATA PENGANTAR")
    add_para(doc, "Puji dan syukur penulis panjatkan ke hadirat Allah Swt. karena atas rahmat dan karunia-Nya laporan Kerja Praktek berjudul \"Pengembangan Sistem Informasi Tata Kelola Siklus Hidup Pengembangan Perangkat Lunak (SDLC) Berbasis Web pada PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari)\" dapat diselesaikan.")
    add_para(doc, "Laporan ini disusun sebagai salah satu bentuk pertanggungjawaban akademik atas pelaksanaan Kerja Praktek pada [ NAMA UNIT/DIVISI RESMI ]. Kegiatan tersebut memberikan kesempatan kepada penulis untuk memahami penerapan pengembangan perangkat lunak, tata kelola proyek, pengujian, keamanan, dokumentasi, dan kolaborasi lintas peran pada lingkungan kerja nyata.")
    add_para(doc, "Penulis menyampaikan terima kasih kepada:", indent=False)
    add_list(doc, [
        "[ NAMA KETUA DEPARTEMEN ], selaku Ketua Departemen Informatika Universitas Andalas.",
        "[ NAMA DOSEN PEMBIMBING ], selaku dosen pembimbing Kerja Praktek.",
        "[ NAMA DOSEN PENGUJI ], selaku dosen penguji laporan Kerja Praktek.",
        "[ NAMA PEMBIMBING LAPANGAN ], selaku pembimbing lapangan pada Bank Nagari.",
        "Seluruh pegawai [ NAMA UNIT/DIVISI RESMI ] yang telah memberikan arahan dan kesempatan belajar.",
        "Orang tua, keluarga, dan rekan-rekan yang memberikan dukungan selama kegiatan dan penyusunan laporan.",
    ])
    add_para(doc, "Penulis menyadari laporan ini masih dapat disempurnakan. Kritik dan saran yang membangun sangat diharapkan agar laporan ini memberi manfaat bagi pengembangan pengetahuan dan praktik tata kelola perangkat lunak.")
    add_para(doc, "Padang, [ BULAN TAHUN ]\nPenulis\n\n\n[ NAMA MAHASISWA ]", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False, line=1.15)
    page_break(doc)

    # Daftar isi dan daftar otomatis lainnya.
    add_front_title(doc, "DAFTAR ISI")
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u', "Klik kanan lalu pilih Update Field")

    table_entries = [
        ("Tabel 1.1 Tahapan pelaksanaan Kerja Praktek", "tbl_1_1"),
        ("Tabel 2.1 Identitas pelaksanaan Kerja Praktek", "tbl_2_1"),
        ("Tabel 2.2 Jadwal kegiatan Kerja Praktek", "tbl_2_2"),
        ("Tabel 3.1 Perbandingan model pengembangan", "tbl_3_1"),
        ("Tabel 3.2 Pemetaan model kualitas ISO/IEC 25010:2023", "tbl_3_2"),
        ("Tabel 3.3 Sintesis landasan teori", "tbl_3_3"),
        ("Tabel 3.4 Matriks literatur dan relevansi NagariSDLC", "tbl_3_4"),
        ("Tabel 4.1 Aktor dan tanggung jawab sistem", "tbl_4_1"),
        ("Tabel 4.2 Kebutuhan fungsional", "tbl_4_2"),
        ("Tabel 4.3 Kebutuhan nonfungsional", "tbl_4_3"),
        ("Tabel 4.4 Daftar 12 peran pengguna", "tbl_4_4"),
        ("Tabel 4.5 Daftar 27 status proyek", "tbl_4_5"),
        ("Tabel 4.6 Matriks approver UAT", "tbl_4_6"),
        ("Tabel 4.7 Tabel domain utama", "tbl_4_7"),
        ("Tabel 4.8 Endpoint API representatif", "tbl_4_8"),
        ("Tabel 4.9 Ringkasan hasil verifikasi", "tbl_4_9"),
        ("Tabel 4.10 Matriks keterlacakan kebutuhan", "tbl_4_10"),
        ("Tabel 4.11 Pemetaan kontrol dan risiko keamanan", "tbl_4_11"),
        ("Tabel 4.12 Evaluasi kualitatif berdasarkan ISO/IEC 25010:2023", "tbl_4_12"),
        ("Tabel 4.13 Ringkasan keputusan arsitektur", "tbl_4_13"),
        ("Tabel 4.14 Kendala dan solusi", "tbl_4_14"),
        ("Tabel 4.15 Roadmap pengembangan", "tbl_4_15"),
        ("Tabel C.1 Daftar endpoint API", "tbl_c_1"),
        ("Tabel D.1 Kamus data ringkas", "tbl_d_1"),
        ("Tabel E.1 Log kegiatan", "tbl_e_1"),
    ]
    table_list_title = add_front_title(doc, "DAFTAR TABEL")
    # The automatic TOC already ends at a natural page boundary. A forced
    # break here creates an extra blank page after Word expands the TOC.
    table_list_title.paragraph_format.page_break_before = False
    for label, bookmark in table_entries:
        add_list_entry(doc, label, bookmark)
    page_break(doc)

    figure_entries = [
        ("Gambar 2.1 Logo Bank Nagari", "fig_2_1"),
        ("Gambar 2.2 Gedung/lokasi Kerja Praktek", "fig_2_2"),
        ("Gambar 2.3 Struktur organisasi instansi", "fig_2_3"),
        ("Gambar 2.4 Struktur unit Teknologi Informasi", "fig_2_4"),
        ("Gambar 4.1 Proses bisnis sebelum sistem terpusat", "fig_4_1"),
        ("Gambar 4.2 Proses bisnis yang diusulkan", "fig_4_2"),
        ("Gambar 4.3 Arsitektur sistem", "fig_4_3"),
        ("Gambar 4.4 Use case aktor utama", "fig_4_4"),
        ("Gambar 4.5 Mesin status proyek", "fig_4_5"),
        ("Gambar 4.6 Alur SIT", "fig_4_6"),
        ("Gambar 4.7 Alur UAT", "fig_4_7"),
        ("Gambar 4.8 Jalur QA dan keamanan siber", "fig_4_8"),
        ("Gambar 4.9 Relasi approver UAT", "fig_4_9"),
        ("Gambar 4.10 Putaran pengembalian pengujian", "fig_4_10"),
        ("Gambar 4.11 ERD NagariSDLC", "fig_4_11"),
        ("Gambar 4.12 Class relationship inti", "fig_4_12"),
        ("Gambar 4.13 Sequence login", "fig_4_13"),
        ("Gambar 4.14 Sequence persetujuan UAT eksternal", "fig_4_14"),
        ("Gambar 4.15 Peta navigasi halaman", "fig_4_15"),
    ]
    for n, label in enumerate([
        "Antarmuka login", "Antarmuka dashboard", "Daftar proyek", "Detail proyek", "Papan Kanban", "Detail task",
        "Wizard SIT", "Wizard UAT", "Tugas QA", "Tugas keamanan siber", "Matriks persetujuan UAT",
        "Quality Gate", "Manajemen pengguna", "Manajemen grup dan role", "Document Vault", "Activity Log",
    ], start=16):
        figure_entries.append((f"Gambar 4.{n} {label}", f"fig_4_{n}"))
    add_front_title(doc, "DAFTAR GAMBAR")
    for label, bookmark in figure_entries:
        add_list_entry(doc, label, bookmark)
    page_break(doc)

    appendix_entries = [
        ("Lampiran A Daftar placeholder tangkapan layar", "app_a"),
        ("Lampiran B Indeks kode Mermaid", "app_b"),
        ("Lampiran C Endpoint API representatif", "app_c"),
        ("Lampiran D Kamus data ringkas", "app_d"),
        ("Lampiran E Log kegiatan Kerja Praktek", "app_e"),
        ("Lampiran F Formulir bimbingan dan penilaian", "app_f"),
    ]
    add_front_title(doc, "DAFTAR LAMPIRAN")
    for label, bookmark in appendix_entries:
        add_list_entry(doc, label, bookmark)

    # Section baru: isi utama menggunakan angka Arab mulai 1.
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    set_page_number_format(body_section, "decimal", 1)
    set_footer_page_number(body_section)

    # Isi bab dan lampiran ditambahkan oleh fungsi terpisah agar builder tetap terbaca.
    build_chapters(doc)
    build_appendices(doc)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def build_chapters(doc: Document):
    bid = 1000

    add_chapter_heading(doc, "I", "PENDAHULUAN")
    add_heading(doc, "1.1 Latar Belakang", 2)
    add_para(doc, "Transformasi digital membuat perangkat lunak menjadi bagian penting dalam operasional organisasi, termasuk perbankan. Aplikasi tidak hanya dituntut berfungsi, tetapi juga harus dikembangkan melalui proses yang aman, terdokumentasi, dapat diaudit, dan memiliki pembagian tanggung jawab yang jelas. Kebutuhan tersebut mendorong penerapan Software Development Life Cycle (SDLC) sebagai kerangka untuk mengelola pekerjaan sejak kebutuhan dirumuskan sampai sistem dioperasikan [1], [2].")
    add_para(doc, "Proses SDLC yang dikelola melalui dokumen dan komunikasi terpisah menimbulkan risiko ketidaksamaan informasi, sulitnya memantau status, keterlambatan persetujuan, serta hilangnya bukti keputusan. Pada proyek dengan banyak pemangku kepentingan, masalah tersebut semakin besar karena setiap tahap memerlukan pelaksana, reviewer, dokumen, dan keputusan yang berbeda. Sistem tata kelola terpusat dibutuhkan agar alur kerja tidak hanya terlihat, tetapi juga ditegakkan oleh aturan aplikasi.")
    add_para(doc, "PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari) menjadi tempat pelaksanaan Kerja Praktek. Proyek yang dikerjakan adalah NagariSDLC, yaitu sistem informasi internal untuk mengelola siklus proyek teknologi dari pengajuan kebutuhan, review, analisis, pengembangan, SIT, UAT, QA, audit keamanan siber, pengajuan rilis, sampai live production. Informasi profil perusahaan yang belum didukung dokumen resmi sengaja tidak dikarang dan ditandai sebagai placeholder untuk dilengkapi penulis.")
    add_para(doc, "NagariSDLC dibangun sebagai aplikasi web tiga lapis. Antarmuka menggunakan React 19, Vite 8, dan Tailwind CSS 4; layanan aplikasi menggunakan Laravel 13 pada PHP 8.3; data dikelola pada MySQL melalui Eloquent ORM. Komunikasi antarlapis menggunakan REST API berbasis JSON. Autentikasi SPA menggunakan token Sanctum di dalam cookie HttpOnly, sedangkan header Authorization Bearer dipertahankan untuk kompatibilitas klien dan pengujian. Mesin status terpusat menjadi inti yang membatasi transisi, mencatat riwayat, dan menghasilkan notifikasi.")
    add_para(doc, "Dari sisi tata kelola, sistem memadukan Role-Based Access Control (RBAC), cakupan akses proyek, persetujuan individual, Document Vault, activity log, dan histori status. Dua jalur pengujian setelah UAT internal—QA dan keamanan siber—berjalan paralel serta memiliki status masing-masing. Setelah keduanya lulus, Project Manager mengajukan migrasi dan rilis, kemudian Head of IT mengambil keputusan pada quality gate. Tidak terdapat UAT final setelah QA dan keamanan siber.")
    add_para(doc, "Berdasarkan ruang lingkup tersebut, laporan ini disusun untuk mendokumentasikan latar belakang, landasan teori, analisis, perancangan, implementasi, pengujian, dan hasil proyek. Kedalaman penyajian mengikuti pola laporan Kerja Praktek yang menjadi referensi pengguna, tetapi seluruh substansi teknis ditulis khusus berdasarkan dokumentasi dan implementasi NagariSDLC.")

    add_heading(doc, "1.2 Rumusan Masalah", 2)
    add_list(doc, [
        "Bagaimana merancang sistem informasi yang memusatkan proses tata kelola proyek teknologi dari pengajuan sampai produksi?",
        "Bagaimana menerapkan mesin status dan kontrol akses agar setiap transisi dilakukan oleh peran yang berwenang dan tercatat sebagai jejak audit?",
        "Bagaimana mengelola SIT dan UAT bertahap, termasuk persetujuan individual serta revisi minor dan mayor?",
        "Bagaimana mengelola QA dan keamanan siber sebagai dua jalur paralel yang independen tanpa menimbulkan status yang ambigu?",
        "Bagaimana menyediakan antarmuka web, layanan API, dokumen, notifikasi, dan informasi proyek yang konsisten antara front end dan back end?",
    ])

    add_heading(doc, "1.3 Tujuan", 2)
    add_list(doc, [
        "Membangun sistem informasi tata kelola SDLC berbasis web untuk mendukung proses proyek teknologi di Bank Nagari.",
        "Menerapkan state machine proyek, RBAC, cakupan akses, histori status, dan activity log sebagai kontrol tata kelola.",
        "Mengimplementasikan pengelolaan task, SIT, UAT, persetujuan UAT per orang, QA, keamanan siber, dan quality gate rilis.",
        "Menyediakan antarmuka responsif berbahasa Indonesia dengan satu lapisan layanan API terpusat.",
        "Mendokumentasikan arsitektur, alur, model data, hubungan peran, endpoint, pengujian, dan keterbatasan sistem secara terstruktur.",
    ])

    add_heading(doc, "1.4 Manfaat Kerja Praktek", 2)
    add_heading(doc, "1.4.1 Manfaat bagi Mahasiswa", 3)
    add_list(doc, [
        "Memperoleh pengalaman menerapkan pengembangan fullstack pada aplikasi dengan aturan bisnis dan audit trail yang kompleks.",
        "Meningkatkan kemampuan menelusuri kebutuhan, merancang model data, mengembangkan REST API, dan membangun antarmuka React.",
        "Memahami praktik pengujian, keamanan aplikasi, dokumentasi teknis, dan kolaborasi lintas peran.",
    ])
    add_heading(doc, "1.4.2 Manfaat bagi Instansi", 3)
    add_list(doc, [
        "Mendukung digitalisasi proses tata kelola proyek perangkat lunak melalui satu sumber informasi terpusat.",
        "Meningkatkan keterlacakan status, dokumen, task, pengujian, persetujuan, dan keputusan rilis.",
        "Mengurangi risiko transisi tidak sah melalui state machine, kontrol peran, dan validasi prasyarat.",
    ])
    add_heading(doc, "1.4.3 Manfaat bagi Perguruan Tinggi", 3)
    add_list(doc, [
        "Menjadi bukti penerapan pengetahuan akademik pada proyek perangkat lunak di lingkungan industri.",
        "Memperkaya contoh studi kasus mengenai SDLC governance, RBAC, workflow, pengujian, dan audit trail.",
        "Mendukung hubungan kerja sama dan pertukaran pengetahuan antara perguruan tinggi dan instansi.",
    ])

    add_heading(doc, "1.5 Batasan Masalah", 2)
    add_list(doc, [
        "Laporan membahas aplikasi NagariSDLC dan tidak membahas sistem inti perbankan maupun transaksi keuangan.",
        "Daftar peran otorisasi dibatasi pada 12 role tetap yang aktif pada implementasi.",
        "Grup kerja dipakai sebagai pengelompokan tampilan dan tidak menjadi sumber otorisasi fase.",
        "Data SIT/UAT disimpan pada projects.sit_uat_data dengan approval aktif UAT pada uat_approval_rounds dan uat_approvers.",
        "Realtime produksi, database produksi, queue, object storage, CI/CD, backup, monitoring, SLA, dan kebijakan retensi final masih memerlukan keputusan operasional.",
        "Tangkapan layar, logo, struktur organisasi resmi, dan dokumen pengesahan ditandai sebagai placeholder hingga bahan resmi tersedia.",
        "Snapshot pengujian yang dicantumkan merupakan kondisi historis tanggal 26 Agustus 2026 dan bukan hasil menjalankan test pada penyusunan laporan ini.",
        "Diagram tidak disisipkan sebagai gambar; dokumen menampilkan placeholder dan berkas pendamping menyediakan kode Mermaid.",
    ])

    add_heading(doc, "1.6 Metodologi Pelaksanaan dan Penyusunan Laporan", 2)
    add_heading(doc, "1.6.1 Pendekatan Pelaksanaan", 3)
    add_para(doc, "Pelaksanaan Kerja Praktek menggunakan pendekatan rekayasa perangkat lunak iteratif. Setiap iterasi dimulai dengan memahami kebutuhan dan aturan bisnis, menelusuri implementasi yang sudah ada, merancang perubahan pada lapisan yang tepat, mengintegrasikan front end dengan back end, kemudian memeriksa hasilnya secara terarah. Pendekatan ini sesuai dengan prinsip bahwa proses pengembangan perlu dapat menyesuaikan perubahan kebutuhan tanpa kehilangan disiplin dokumentasi dan verifikasi [25], [26].")
    add_para(doc, "NagariSDLC sendiri tidak diposisikan sebagai implementasi satu metodologi SDLC tertentu. Sistem berfungsi sebagai lapisan tata kelola yang dapat menaungi pekerjaan berurutan maupun iteratif. Status, role, dokumen, task, pengujian, approval, dan quality gate menjadi bukti proses, sedangkan cara tim mengorganisasi sprint atau urutan implementasi dapat disesuaikan dengan karakter proyek. Pemisahan ini penting agar tata kelola tidak disamakan dengan metode pengembangan.")
    add_heading(doc, "1.6.2 Teknik Pengumpulan Data", 3)
    add_para(doc, "Data laporan diperoleh melalui studi dokumentasi proyek, observasi terhadap perilaku aplikasi dan struktur source code, penelusuran route, service, model, migration, serta konfigurasi aktif, dan diskusi kebutuhan yang terekam selama pengembangan. Informasi teknis divalidasi dengan membandingkan dokumentasi resmi proyek terhadap implementasi aktual. Apabila ditemukan perbedaan, source code, migration, route, konfigurasi, dan keputusan proyek terbaru ditempatkan sebagai sumber kebenaran.")
    add_para(doc, "Landasan ilmiah diperoleh melalui studi pustaka terarah. Sumber diprioritaskan dari standar internasional, publikasi lembaga resmi, artikel jurnal peer-reviewed, buku akademik, dan dokumentasi resmi teknologi. Literatur tidak hanya dirangkum, tetapi disintesis untuk menjelaskan alasan pemilihan arsitektur, kontrol akses, workflow, mekanisme pengujian, keamanan, dan audit trail pada NagariSDLC.")
    add_heading(doc, "1.6.3 Tahapan Pekerjaan", 3)
    methodology_rows = [
        ["1", "Identifikasi konteks", "Memahami tujuan produk, aktor, masalah tata kelola, dan ruang lingkup kerja."],
        ["2", "Studi kebutuhan dan literatur", "Menganalisis PRD, workflow, arsitektur, model data, standar, jurnal, dan buku."],
        ["3", "Analisis sistem", "Memetakan proses bisnis, kebutuhan fungsional/nonfungsional, role, status, serta risiko."],
        ["4", "Perancangan", "Menyusun arsitektur, model data, kontrak API, navigasi, state machine, dan rancangan pengujian."],
        ["5", "Implementasi", "Mengembangkan back end Laravel dan front end React secara terintegrasi."],
        ["6", "Verifikasi", "Memeriksa workflow, otorisasi, kontrak data, pengujian, linting, build, dan bukti audit."],
        ["7", "Dokumentasi dan evaluasi", "Menyusun laporan, diagram, keterbatasan, kesimpulan, serta rekomendasi."],
    ]
    add_table(doc, ["Tahap", "Kegiatan", "Uraian"], methodology_rows, [750, 2050, 5138], "Tabel 1.1 Tahapan pelaksanaan Kerja Praktek", "tbl_1_1", bid); bid += 1

    add_heading(doc, "1.7 Sistematika Penulisan", 2)
    add_para(doc, "BAB I Pendahuluan memuat latar belakang, rumusan masalah, tujuan, manfaat, batasan, metodologi pelaksanaan, dan sistematika. BAB II Profil Instansi dan Pelaksanaan Kerja Praktek memuat profil resmi yang dapat diverifikasi, struktur organisasi dalam bentuk placeholder, penempatan, uraian tugas, dan jadwal kegiatan. BAB III Tinjauan Pustaka menguraikan teori, standar, hasil penelitian, dan sintesis literatur yang mendasari pengembangan. BAB IV Hasil dan Pembahasan menyajikan analisis, perancangan, implementasi, pengujian, evaluasi, serta keterkaitan hasil dengan literatur. BAB V Penutup berisi kesimpulan serta saran. Bagian akhir memuat daftar pustaka dan lampiran pendukung.")

    add_chapter_heading(doc, "II", "PROFIL INSTANSI DAN PELAKSANAAN KERJA PRAKTEK")
    add_heading(doc, "2.1 Profil Instansi", 2)
    add_heading(doc, "2.1.1 Identitas Instansi", 3)
    add_para(doc, "Kerja Praktek dilaksanakan pada Bank Nagari, bank pembangunan daerah yang berakar di Sumatera Barat. Situs resmi perusahaan mencantumkan kantor pusat di Jalan Pemuda Nomor 21, Padang, Sumatera Barat [35]. Identitas unit penempatan, periode, dan pembimbing lapangan tetap dibiarkan sebagai placeholder karena hanya dapat dipastikan melalui surat penempatan atau dokumen internal mahasiswa.")
    add_table(doc, ["Keterangan", "Isi"], [
        ["Nama instansi", "PT Bank Pembangunan Daerah Sumatera Barat (Bank Nagari)"],
        ["Unit/divisi", "[ NAMA RESMI UNIT/DIVISI TEMPAT KERJA PRAKTEK ]"],
        ["Alamat kantor pusat", "Jl. Pemuda No. 21, Padang, Sumatera Barat [35]"],
        ["Periode", "[ TANGGAL MULAI ] - [ TANGGAL SELESAI ]"],
        ["Pembimbing lapangan", "[ NAMA DAN JABATAN ]"],
        ["Posisi mahasiswa", "Pengembang fullstack pada proyek NagariSDLC"],
    ], [2200, 5738], "Tabel 2.1 Identitas pelaksanaan Kerja Praktek", "tbl_2_1", bid); bid += 1

    add_heading(doc, "2.1.2 Sejarah Singkat", 3)
    add_para(doc, "Berdasarkan profil resmi perusahaan, Bank Pembangunan Daerah Sumatera Barat didirikan pada 12 Maret 1962 dan disahkan melalui akta notaris Hasan Qalbi di Padang. Pembentukannya dipelopori oleh pemerintah daerah bersama tokoh masyarakat dan pelaku usaha sebagai lembaga keuangan yang mendukung pembangunan daerah. Operasi awal berlokasi di Jalan Batang Arau Nomor 54 Padang dengan modal awal Rp50.000.000 [36].")
    add_para(doc, "Perusahaan membuka kantor cabang pertama di Payakumbuh pada 1965. Bentuk badan hukum kemudian berubah menjadi Perusahaan Daerah pada 1973, dan sebutan Bank Nagari mulai digunakan pada 1996 untuk memperkuat pengenalan merek. Pada 2006 badan hukum kembali diubah menjadi Perseroan Terbatas. Profil resmi juga mencatat keputusan perubahan nama perseroan menjadi PT Bank Nagari pada 2021 [36]. Uraian ini digunakan sebatas konteks instansi dan perlu dicocokkan kembali dengan laporan tahunan yang berlaku pada tahun pengesahan laporan.")
    add_para(doc, "Perjalanan tersebut menunjukkan bahwa Bank Nagari berkembang bersama perubahan kebutuhan kelembagaan dan layanan. Dalam konteks Kerja Praktek, pengembangan NagariSDLC dapat dipahami sebagai bagian dari upaya memperkuat proses internal berbasis teknologi: bukan sebagai aplikasi transaksi nasabah, melainkan sebagai sarana tata kelola pengembangan sistem agar keputusan, pengujian, persetujuan, dan bukti proyek lebih terstruktur.")
    add_heading(doc, "2.1.3 Visi dan Misi", 3)
    add_para(doc, "Visi resmi Bank Nagari adalah menjadi bank pembangunan daerah yang terkemuka dan tepercaya di Indonesia [35]. Makna terkemuka berkaitan dengan pengenalan dan posisi bank pada tingkat nasional, sedangkan tepercaya dikaitkan dengan penerapan manajemen perusahaan yang baik, layanan yang memuaskan, kejujuran, dan kepatuhan.")
    add_para(doc, "Misi Bank Nagari menekankan kontribusi terhadap pertumbuhan ekonomi dan kesejahteraan masyarakat serta pemenuhan kepentingan para pemangku kepentingan secara konsisten dan seimbang [35]. Misi tersebut dijabarkan melalui upaya menjaga pertumbuhan bank yang sehat, memberikan layanan prima, memberikan keuntungan yang memadai bagi pemegang saham, dan memberikan manfaat maksimal bagi masyarakat.")
    add_para(doc, "Hubungan visi dan misi dengan proyek NagariSDLC bersifat tidak langsung tetapi relevan. Sistem tata kelola pengembangan perangkat lunak yang dapat diaudit mendukung konsistensi proses internal, memperjelas tanggung jawab, dan membantu menjaga mutu perubahan teknologi. Laporan ini tidak menyatakan bahwa NagariSDLC merupakan satu-satunya instrumen pencapaian visi perusahaan; sistem hanya menjadi salah satu sarana pendukung pengelolaan proyek teknologi.")
    add_heading(doc, "2.1.4 Logo dan Lokasi", 3)
    add_figure_placeholder(doc, "Gambar 2.1 Logo Bank Nagari", "fig_2_1", bid, "Masukkan logo resmi beresolusi tinggi setelah memperoleh izin penggunaan.", lines=4); bid += 1
    add_figure_placeholder(doc, "Gambar 2.2 Gedung/lokasi Kerja Praktek", "fig_2_2", bid, "Masukkan foto gedung atau lokasi unit tempat Kerja Praktek dan cantumkan sumber/dokumentasi.", lines=5); bid += 1

    add_heading(doc, "2.2 Struktur Organisasi", 2)
    add_para(doc, "Struktur organisasi diperlukan untuk menjelaskan kedudukan unit kerja dan jalur koordinasi proyek. Karena struktur resmi belum dilampirkan, laporan menyediakan placeholder. Kode Mermaid D-02 hanya bersifat bagan representatif berdasarkan peran aplikasi dan wajib disesuaikan dengan nomenklatur resmi sebelum digunakan.")
    add_figure_placeholder(doc, "Gambar 2.3 Struktur organisasi instansi", "fig_2_3", bid, "Masukkan bagan struktur organisasi Bank Nagari yang berlaku pada periode Kerja Praktek.", lines=6); bid += 1
    add_figure_placeholder(doc, "Gambar 2.4 Struktur unit Teknologi Informasi", "fig_2_4", bid, "Render kode Mermaid D-02 atau ganti dengan bagan resmi unit Teknologi Informasi.", "D-02", lines=7); bid += 1

    add_heading(doc, "2.3 Unit Kerja dan Posisi Mahasiswa", 2)
    add_heading(doc, "2.3.1 Unit Kerja", 3)
    add_para(doc, "Kegiatan dilaksanakan pada [ NAMA RESMI UNIT/DIVISI ]. Dalam konteks NagariSDLC, pekerjaan berhubungan dengan fungsi perencanaan, pengembangan, pengujian mutu, keamanan siber, manajemen teknologi informasi, dan pemohon bisnis. Uraian ini menjelaskan domain aplikasi, bukan menetapkan struktur organisasi formal perusahaan.")
    add_heading(doc, "2.3.2 Posisi sebagai Pengembang Fullstack", 3)
    add_para(doc, "Mahasiswa berperan sebagai pengembang fullstack yang mengerjakan sisi back end dan front end. Tanggung jawab back end mencakup model, migrasi, Form Request, controller, resource, service, state machine, autentikasi, dan pengujian otomatis. Tanggung jawab front end mencakup halaman, komponen, routing, role guard, context, layanan API, validasi antarmuka, dan integrasi dengan endpoint.")
    add_heading(doc, "2.3.3 Uraian Tugas", 3)
    add_list(doc, [
        "Mempelajari kebutuhan produk, dokumentasi proyek, basis kode, dan model data yang aktif.",
        "Mengembangkan dan menyempurnakan alur proyek, task, SIT, UAT, QA, keamanan siber, dan rilis.",
        "Menjaga konsistensi kontrak field dan endpoint antara front end dan back end.",
        "Menerapkan validasi, otorisasi, histori, notifikasi, dan perlindungan audit trail.",
        "Melakukan pemeriksaan statis, pengujian terarah, dokumentasi, dan perbaikan berdasarkan temuan.",
    ])

    add_heading(doc, "2.4 Jadwal dan Kegiatan Kerja Praktek", 2)
    add_heading(doc, "2.4.1 Kegiatan Mingguan", 3)
    schedule = [
        ["1", "Orientasi, pengenalan proses tata kelola SDLC bank, dan penyiapan lingkungan PHP 8.3, Laravel 13, React 19, serta MySQL.", "Lingkungan kerja dan pemahaman domain awal."],
        ["2", "Analisis kebutuhan, penelusuran basis kode, pemahaman model data, peran, dan alur proses.", "Catatan analisis dan pemetaan modul."],
        ["3", "Back end: model, migrasi, RBAC, dan autentikasi Sanctum.", "Fondasi data, akses, dan sesi pengguna."],
        ["4", "Back end: mesin status melalui ProjectWorkflowService, riwayat status, dan notifikasi.", "Workflow terpusat dan audit trail status."],
        ["5", "Back end: modul SIT/UAT, jalur QA dan keamanan siber, serta gerbang mutu rilis.", "Alur pengujian dan quality gate."],
        ["6", "Front end: dashboard, papan Kanban, modul registrasi, dan manajemen pengguna.", "Antarmuka kerja utama dan administrasi."],
        ["7", "Front end: wizard SIT/UAT, halaman tugas QA/Siber, dan matriks persetujuan UAT.", "Antarmuka pengujian dan approval."],
        ["8", "Pengujian menyeluruh, perbaikan akhir, dokumentasi, dan finalisasi.", "Snapshot kualitas dan dokumentasi proyek."],
    ]
    add_table(doc, ["Minggu", "Kegiatan", "Luaran"], schedule, [850, 4800, 2288], "Tabel 2.2 Jadwal kegiatan Kerja Praktek", "tbl_2_2", bid); bid += 1
    add_heading(doc, "2.4.2 Hasil Kegiatan", 3)
    add_para(doc, "Hasil kegiatan berupa implementasi aplikasi NagariSDLC, dokumentasi teknis, pemetaan workflow, perbaikan integrasi front end dan back end, serta bukti verifikasi yang tercatat pada dokumentasi proyek. Log harian rinci tetap perlu diisi menggunakan catatan kegiatan asli dan ditempatkan pada Lampiran E.")

    add_chapter_heading(doc, "III", "TINJAUAN PUSTAKA")
    theory_sections = [
        ("3.1 Sistem Informasi dan Tata Kelola Teknologi Informasi", [
            "Sistem informasi merupakan kesatuan manusia, prosedur, data, perangkat lunak, perangkat keras, dan jaringan yang bekerja untuk menghasilkan informasi bagi operasi dan pengambilan keputusan. Karena itu, keberhasilan sistem tidak hanya ditentukan oleh kemampuan menyimpan data. Sistem harus membantu pengguna memahami pekerjaan, mengurangi ketidakpastian, menjaga kualitas informasi, dan menyediakan bukti yang cukup untuk mengevaluasi keputusan [1], [2].",
            "Dalam konteks tata kelola teknologi informasi, proses pengembangan perangkat lunak perlu memiliki pemilik, pelaksana, reviewer, kriteria masuk dan keluar, serta bukti hasil. Tata kelola berbeda dari manajemen harian: tata kelola menetapkan arah, batas kewenangan, akuntabilitas, dan mekanisme pengawasan; manajemen melaksanakan pekerjaan di dalam batas tersebut. NagariSDLC menghubungkan keduanya melalui state machine, penugasan, persetujuan, histori, dan quality gate.",
            "Proyek menjadi agregat utama pada NagariSDLC. Di sekelilingnya terdapat kebutuhan, task, anggota tim, dokumen, SIT, UAT, QA, pengujian keamanan siber, approval, release request, notifikasi, dan activity log. Struktur ini mengubah data yang semula berpotensi tersebar menjadi rangkaian informasi yang saling terkait. Nilai sistem terletak pada kemampuan merekonstruksi alasan sebuah proyek berada pada status tertentu dan tindakan apa yang masih diperlukan.",
        ]),
        ("3.2 Software Development Life Cycle", [
            "Software Development Life Cycle (SDLC) adalah kerangka untuk mengelola perangkat lunak sejak gagasan dan kebutuhan, perancangan, implementasi, verifikasi, validasi, transisi, operasi, pemeliharaan, hingga penghentian. Sommerville serta Pressman dan Maxim menempatkan aktivitas spesifikasi, pengembangan, validasi, dan evolusi sebagai unsur penting rekayasa perangkat lunak [1], [2]. SDLC membuat pekerjaan kompleks dapat diuraikan menjadi keluaran, tanggung jawab, dan keputusan yang lebih terukur.",
            "ISO/IEC/IEEE 12207:2026 menyediakan kerangka proses siklus hidup perangkat lunak yang dapat digunakan untuk pekerjaan internal maupun hubungan pemasok–pengakuisisi. Standar tersebut bersifat method-agnostic dan mengakui bahwa proses dapat dilakukan secara berulang, rekursif, serta paralel [15]. Artinya, disiplin lifecycle tidak identik dengan model air terjun; organisasi dapat menggunakan pendekatan iteratif selama tanggung jawab, hasil, dan keterlacakan tetap jelas.",
            "NagariSDLC menerjemahkan lifecycle menjadi governance workflow. Tahap pengajuan, review, analisis, pengembangan, SIT, UAT, QA, keamanan siber, rilis, dan produksi memiliki aktor dan bukti masing-masing. Aplikasi tidak menentukan bagaimana developer mengorganisasi pekerjaan internal sehari-hari, tetapi menegakkan bahwa suatu quality gate hanya dapat dilewati ketika prasyarat yang disepakati telah terpenuhi.",
        ]),
        ("3.3 Model Pengembangan Perangkat Lunak", [
            "Model sekuensial menekankan penyelesaian tahap secara relatif berurutan dan cocok ketika kebutuhan stabil serta bukti formal dibutuhkan. Model incremental membagi produk menjadi bagian yang dapat dikembangkan dan dievaluasi bertahap. Model spiral Boehm menempatkan identifikasi serta mitigasi risiko sebagai penggerak setiap putaran [29]. Model agile menekankan umpan balik cepat, kolaborasi, perangkat lunak yang bekerja, dan kemampuan merespons perubahan [25].",
            "Tidak ada model yang selalu paling tepat. Penelitian sistematis Dybå dan Dingsøyr menunjukkan adanya manfaat sekaligus keterbatasan pendekatan agile, serta menekankan pentingnya konteks ketika menilai bukti empiris [26]. Proyek dengan tuntutan audit tinggi dapat membutuhkan dokumen dan approval formal, sementara ketidakpastian kebutuhan tetap menuntut iterasi. Oleh sebab itu, pendekatan hibrida sering lebih realistis daripada memaksakan satu label metode.",
            "NagariSDLC dirancang untuk mendukung pola hibrida tersebut. Task dapat dikerjakan iteratif, revisi UAT dapat membuka pekerjaan perbaikan, dan jalur QA–Siber dapat berjalan paralel. Pada saat yang sama, sistem mempertahankan gerbang formal untuk sign-off, perubahan status, dan rilis. Dengan demikian fleksibilitas implementasi tidak menghilangkan akuntabilitas proses.",
        ]),
        ("3.4 Agile, Continuous Software Engineering, dan DevOps", [
            "Agile Manifesto menghargai individu dan interaksi, perangkat lunak yang bekerja, kolaborasi dengan pengguna, serta respons terhadap perubahan tanpa meniadakan nilai proses dan dokumentasi [25]. Pada proyek yang diatur ketat, prinsip agile tidak boleh diterjemahkan sebagai penghapusan bukti. Dokumentasi perlu dibuat secukupnya, dekat dengan aktivitas yang dibuktikan, dan mudah diperbarui ketika keadaan berubah.",
            "Fitzgerald dan Stol menjelaskan continuous software engineering sebagai upaya mengurangi keterputusan antara perencanaan, pengembangan, integrasi, pengujian, rilis, dan operasi [27]. Lwakatare dan rekan-rekan menunjukkan bahwa DevOps berkaitan dengan kemampuan memperbarui sistem operasional secara sering dan andal melalui praktik lintas fungsi [28]. Humble dan Farley menekankan pipeline rilis berulang dan berisiko rendah, sedangkan Forsgren, Humble, dan Kim menghubungkan praktik delivery dengan pengukuran kinerja teknologi serta organisasi [32], [33].",
            "NagariSDLC belum merupakan platform CI/CD dan tidak mengotomatisasi deployment. Kontribusinya berada pada continuity of governance: hasil analisis mengalir ke pengembangan, task menjadi dasar SIT/UAT, hasil pengujian menjadi prasyarat rilis, dan keputusan produksi memiliki histori. Pengembangan lanjutan dapat mengintegrasikan pipeline otomatis, tetapi integrasi tersebut harus tetap tunduk pada quality gate dan otorisasi yang berlaku.",
        ]),
        ("3.5 Workflow, State Machine, dan Quality Gate", [
            "Workflow menjelaskan aktivitas, urutan, pelaksana, input, keluaran, serta kondisi perpindahan. State machine memodelkan sistem sebagai himpunan keadaan dan transisi yang sah. Pemodelan ini bermanfaat ketika sebuah nilai status bukan sekadar label tampilan, tetapi menentukan aksi yang boleh dilakukan, aktor yang berwenang, dan bukti yang harus tersedia.",
            "Quality gate adalah titik keputusan yang menahan perpindahan sampai kriteria tertentu dipenuhi. Gate yang baik harus memiliki kriteria objektif, pemilik keputusan, bukti, dan hasil yang dapat diaudit. Gate yang hanya ditampilkan pada antarmuka mudah dilewati melalui endpoint lain; karena itu aturan perlu ditegakkan pada lapisan domain atau service yang menjadi sumber kebenaran.",
            "Pada NagariSDLC, allowedTransitions mendefinisikan bentuk perpindahan, rolePermissions membatasi pelaku, dan validator prasyarat menilai kondisi bisnis. Semua transisi proyek dipusatkan pada ProjectWorkflowService. Keputusan ini menghindari status yang diubah secara bebas oleh controller atau komponen UI dan memudahkan audit terhadap satu jalur perubahan yang konsisten.",
        ]),
        ("3.6 Arsitektur Perangkat Lunak dan Pemisahan Tanggung Jawab", [
            "Arsitektur perangkat lunak menggambarkan struktur utama sistem, tanggung jawab elemen, serta hubungan di antaranya. Bass, Clements, dan Kazman menekankan bahwa keputusan arsitektur berkaitan erat dengan atribut kualitas seperti modifiability, availability, security, dan performance [30]. Arsitektur yang baik bukan sekadar diagram teknologi; keputusan harus menjelaskan trade-off yang ingin dikelola.",
            "Pemisahan presentation, application/domain logic, dan persistence mengurangi ketergantungan langsung. Perubahan tampilan tidak seharusnya mengubah aturan workflow, sedangkan perubahan database tidak seharusnya memaksa komponen UI mengetahui detail penyimpanan. Batas yang jelas juga mempermudah pengujian karena perilaku bisnis dapat diuji melalui service atau API tanpa bergantung pada seluruh antarmuka.",
            "NagariSDLC menggunakan React SPA sebagai lapis klien, Laravel REST API sebagai lapis aplikasi, dan MySQL/Eloquent sebagai lapis data. Di back end, Form Request menangani validasi input, controller mengorkestrasi request, service menampung aturan lintas endpoint, model merepresentasikan data, dan Resource menormalisasi respons. Struktur tersebut meningkatkan maintainability walaupun tetap membutuhkan disiplin agar service tidak berkembang menjadi objek yang terlalu besar.",
        ]),
        ("3.7 REST, HTTP, dan Desain API", [
            "REST merupakan gaya arsitektur untuk sistem jaringan yang menekankan client–server, stateless interaction, cacheability, uniform interface, layered system, dan code-on-demand opsional [3]. HTTP sendiri merupakan protokol application-level yang stateless dengan pola request–response, resource yang diidentifikasi URI, method yang menyatakan maksud, serta status code yang menyampaikan hasil [22]. REST bukan sinonim dari JSON, tetapi JSON umum digunakan sebagai representasi data.",
            "Desain API perlu konsisten pada penamaan resource, penggunaan method, status code, format validasi, pagination, dan penanganan error. Request yang valid secara sintaksis masih dapat ditolak oleh aturan bisnis; dalam kasus tersebut respons harus membedakan kegagalan autentikasi, otorisasi, data tidak ditemukan, dan prasyarat workflow. Konsistensi kontrak mengurangi logika khusus pada klien dan mempercepat diagnosis integrasi.",
            "NagariSDLC memakai REST API berbasis JSON dengan envelope { status, message, data, meta? }. Kode 401 digunakan untuk sesi tidak sah, 403 untuk larangan akses, 404 untuk resource yang tidak ditemukan, dan 422 untuk validasi atau aturan workflow. Front end memusatkan request pada satu service API sehingga cookie, header, serialisasi, parsing error, dan perilaku logout dapat dikelola secara seragam.",
        ]),
        ("3.8 Laravel sebagai Kerangka Back End", [
            "Laravel menyediakan routing, middleware, Form Request, Eloquent ORM, API Resource, event, notification, queue, broadcasting, cache, dan fasilitas pengujian [4]. Kerangka ini mempercepat pengembangan, tetapi kualitas hasil tetap ditentukan oleh penempatan tanggung jawab. Controller yang terlalu banyak memuat aturan bisnis akan sulit digunakan ulang dan diuji, sedangkan model yang memuat seluruh orkestrasi dapat menimbulkan ketergantungan tersembunyi.",
            "Validasi write pada Form Request memisahkan pemeriksaan bentuk input dari orkestrasi endpoint. Service digunakan ketika aturan melibatkan beberapa model, transaksi, notifikasi, dan perubahan status. Resource membentuk kontrak output sehingga perubahan internal model tidak langsung bocor ke klien. Pola ini sejalan dengan kebutuhan NagariSDLC yang memiliki alur approval dan pengujian lintas tabel.",
            "ProjectWorkflowService menjadi pusat transisi status proyek. UatExecutionService menangani draft/final eksekusi dan revisi; UatApprovalService mengelola putaran persetujuan; TestingTrackService mengelola jalur QA dan keamanan siber; ProjectReturnRoundService menjaga putaran pengembalian. Pembagian ini membuat aturan audit kritis memiliki lokasi yang dapat ditelusuri.",
        ]),
        ("3.9 React, Vite, Tailwind CSS, dan Usability", [
            "React membangun antarmuka melalui komponen deklaratif dan state yang mendorong pembentukan UI berdasarkan data [5]. Vite menyediakan development server serta proses build [6], sedangkan Tailwind CSS menyediakan utility class untuk menyusun tampilan [8]. Kombinasi tersebut mendukung SPA yang responsif, tetapi perlu aturan struktur agar komponen tidak terlalu terikat pada pemanggilan API dan logika bisnis.",
            "Nielsen menjelaskan usability melalui kemudahan dipelajari, efisiensi, kemudahan diingat, tingkat kesalahan, dan kepuasan [34]. Pada sistem workflow, usability juga dipengaruhi kejelasan status, aksi berikutnya, alasan sebuah aksi ditolak, serta konsistensi istilah. Wizard membantu memecah proses panjang, sedangkan pencarian dan filter membantu pengguna menemukan pekerjaan sesuai perannya.",
            "Front end NagariSDLC menggunakan router, protected route, context untuk state lintas komponen, dan service API terpusat. Antarmuka berbahasa Indonesia, menampilkan action sesuai role, serta menerima indikator can_resubmit dan resubmit_blocker dari server. UI mencegah aksi yang pasti tidak sah, tetapi otorisasi akhir tetap berada di back end agar penyembunyian tombol tidak disalahartikan sebagai kontrol keamanan.",
        ]),
        ("3.10 Role-Based Access Control, Least Privilege, dan Separation of Duties", [
            "RBAC menghubungkan permission dengan role dan menugaskan pengguna pada role tersebut. Model Sandhu dan rekan-rekan membedakan role dari group serta menunjukkan bagaimana hierarki dan constraint dapat membentuk kebijakan yang lebih kaya [21]. RBAC mengurangi kebutuhan mengelola izin satu per satu, tetapi role yang terlalu luas dapat menciptakan hak akses berlebihan.",
            "Prinsip least privilege memberikan akses minimum yang diperlukan untuk menjalankan tugas. Separation of duties memisahkan aktivitas yang berpotensi menimbulkan konflik, misalnya pelaksana pengujian dan pemberi sign-off. Pada aplikasi proyek, role saja belum cukup karena dua pengguna dengan role sama belum tentu terlibat pada proyek yang sama. Otorisasi perlu menggabungkan role, assignment, membership, kepemilikan, dan keadaan workflow.",
            "NagariSDLC memiliki 12 role tetap dan lima group bawaan. Group hanya mengelompokkan role untuk tampilan; group tidak memberi wewenang fase. ProjectWorkflowService mengatur hak transisi, ProjectAccessService membatasi cakupan proyek, dan service pengujian membatasi assignee. Model ini memperkecil risiko bahwa perubahan menu atau group tanpa sengaja mengubah otorisasi inti.",
        ]),
        ("3.11 Autentikasi, Sesi, Cookie, CSRF, dan CORS", [
            "Autentikasi membuktikan identitas, sedangkan manajemen sesi mempertahankan konteks pengguna pada request berikutnya. RFC 6265 menjelaskan mekanisme server mengirim Set-Cookie dan user agent mengembalikan Cookie pada request selanjutnya [23]. Atribut HttpOnly mencegah JavaScript membaca cookie, Secure membatasi pengiriman melalui koneksi aman, dan SameSite membantu mengendalikan pengiriman lintas situs. Tidak satu pun atribut tersebut menggantikan validasi server.",
            "Pada aplikasi SPA, risiko utama meliputi pencurian token melalui XSS, cross-site request forgery, origin yang terlalu longgar, session fixation, serta sesi yang tidak dicabut. OWASP ASVS menyediakan persyaratan verifikasi untuk autentikasi, session management, access control, validasi, dan kontrol keamanan web [20]. NIST SSDF menempatkan keamanan sebagai praktik yang perlu diintegrasikan ke dalam seluruh lifecycle, bukan pemeriksaan akhir [18].",
            "NagariSDLC menggunakan token Sanctum di cookie HttpOnly sebagai jalur utama SPA, credentials include pada request, dan X-Requested-With untuk jalur cookie [11]. Token tidak disimpan di localStorage. Header Authorization Bearer dipertahankan untuk kompatibilitas klien serta pengujian. CORS dibatasi melalui konfigurasi environment. Rancangan ini harus dilengkapi konfigurasi HTTPS, domain cookie, SameSite, rotasi token, rate limiting, dan monitoring pada produksi.",
        ]),
        ("3.12 Basis Data Relasional, ORM, dan Integritas Transaksi", [
            "Basis data relasional merepresentasikan entitas dalam tabel dan hubungan melalui primary key serta foreign key. Normalisasi membantu mengurangi redundansi dan anomali pembaruan, sedangkan indeks mempercepat pola pencarian tertentu. Constraint menjaga kondisi yang dapat ditegakkan pada tingkat database. MySQL menyediakan transaksi dan integritas referensial yang dibutuhkan untuk operasi multi-tabel [12].",
            "ORM memetakan tabel ke objek dan relasi sehingga kode aplikasi lebih ekspresif. Namun ORM tidak menghilangkan kebutuhan memahami query, eager loading, indeks, transaksi, dan perilaku delete. Pilihan CASCADE, SET NULL, dan RESTRICT memiliki makna domain: CASCADE sesuai untuk anak yang tidak bermakna tanpa induk; SET NULL sesuai untuk referensi penugasan yang boleh hilang; RESTRICT sesuai untuk bukti audit yang tidak boleh terhapus diam-diam.",
            "NagariSDLC menggunakan transaksi ketika satu aksi mengubah proyek, status jalur, histori, approval, return round, atau notifikasi. Soft delete diterapkan pada entitas yang perlu dipulihkan atau ditelusuri. Data wizard SIT/UAT berada pada projects.sit_uat_data untuk kompatibilitas evolusi, sedangkan approval aktif UAT dinormalisasi ke uat_approval_rounds dan uat_approvers agar snapshot serta keputusan per orang dapat diaudit.",
        ]),
        ("3.13 Konsep dan Tingkatan Pengujian Perangkat Lunak", [
            "Pengujian adalah kegiatan mengevaluasi artefak perangkat lunak untuk menemukan perbedaan antara hasil aktual dan hasil yang diharapkan, serta memberikan informasi tentang kualitas dan risiko. Pengujian tidak dapat membuktikan ketiadaan seluruh cacat, tetapi dapat meningkatkan keyakinan melalui cakupan yang dirancang sesuai risiko. Verifikasi menilai apakah produk dibangun dengan benar terhadap spesifikasi; validasi menilai apakah produk yang dibangun sesuai kebutuhan pengguna [1], [2].",
            "ISO/IEC/IEEE 29119-2:2021 mendefinisikan proses pengujian generik yang dapat digunakan untuk mengatur, mengelola, dan melaksanakan testing pada berbagai model lifecycle [17]. Proses tersebut menghubungkan perencanaan, monitoring, desain, implementasi, eksekusi, pelaporan insiden, dan penyelesaian. Test case perlu memiliki tujuan, prasyarat, data, langkah, hasil yang diharapkan, hasil aktual, bukti, dan status.",
            "Unit test memeriksa bagian kecil; integration test memeriksa interaksi komponen; system test memeriksa sistem lengkap; acceptance test menilai penerimaan pemangku kepentingan. Black-box berfokus pada input dan output, white-box menggunakan struktur internal, sedangkan risk-based testing memprioritaskan area berdasarkan kemungkinan dan dampak kegagalan. NagariSDLC membutuhkan kombinasi tersebut karena risiko utamanya terdapat pada workflow, otorisasi, konsistensi data, dan audit trail.",
        ]),
        ("3.14 System Integration Test dan User Acceptance Test", [
            "SIT memverifikasi bahwa komponen yang telah dikembangkan dapat bekerja sebagai kesatuan, termasuk interaksi API, database, layanan, dan antarmuka. UAT dilakukan dari sudut pandang kebutuhan bisnis untuk menentukan apakah sistem dapat diterima. Keduanya berbeda dari unit test: SIT dan UAT membutuhkan skenario end-to-end, data uji, lingkungan, bukti, serta pihak yang bertanggung jawab terhadap hasil [17].",
            "UAT bukan hanya satu tombol persetujuan. Proses penerimaan perlu menjelaskan siapa peserta dan penanda tangan, skenario apa yang diuji, bukti apa yang dihasilkan, bagaimana revisi dicatat, serta apakah keputusan berlaku pada versi hasil yang sama. Ketika hasil berubah setelah revisi, approval lama tidak boleh dianggap mewakili artefak terbaru; diperlukan putaran baru atau mekanisme supersession.",
            "NagariSDLC membagi SIT dan UAT menjadi tiga tahap. SIT memerlukan bukti per task dan sign-off developer, PM, serta Development Lead. UAT mengelola peserta, undangan, skenario, Change Request, dan tujuh peran approval. Revisi minor menahan approval sampai task perbaikan selesai; revisi mayor mengulang SIT penuh dan UAT dari tahap pertama. Putaran lama disimpan sebagai histori.",
        ]),
        ("3.15 Quality Assurance dan Model Kualitas Produk", [
            "Quality Assurance (QA) berorientasi pada keyakinan bahwa proses dan produk memenuhi kebutuhan mutu, sedangkan quality control lebih berfokus pada pemeriksaan hasil. QA tidak sama dengan testing; testing adalah salah satu teknik yang memberikan bukti. Praktik QA juga mencakup review kebutuhan, standar pengembangan, definisi selesai, pengelolaan cacat, keterlacakan, dan evaluasi proses.",
            "ISO/IEC 25010:2023 mendefinisikan model kualitas produk yang terdiri atas sembilan karakteristik: functional suitability, performance efficiency, compatibility, interaction capability, reliability, security, maintainability, flexibility, dan safety [16]. Karakteristik tersebut membantu mengubah istilah 'berkualitas' menjadi sasaran yang dapat dijelaskan dan diuji. Tidak semua karakteristik memiliki bobot sama; prioritas mengikuti konteks dan risiko produk.",
            "Pada NagariSDLC, functional suitability berkaitan dengan kelengkapan workflow; security dengan autentikasi, otorisasi, serta audit; reliability dengan transaksi dan konsistensi; maintainability dengan pemisahan service; interaction capability dengan wizard dan pesan status; compatibility dengan dukungan data legacy serta Bearer token. Pengukuran kuantitatif penuh belum dilakukan, sehingga laporan membahas alignment desain dan bukti yang tersedia, bukan mengklaim sertifikasi ISO.",
        ]),
        ("3.16 Secure SDLC dan Pengujian Keamanan Siber", [
            "NIST Secure Software Development Framework versi 1.1 mengelompokkan praktik ke dalam Prepare the Organization, Protect the Software, Produce Well-Secured Software, dan Respond to Vulnerabilities [18]. Kerangka tersebut dirancang agar dapat ditambahkan pada berbagai model SDLC. Tujuannya mengurangi jumlah kerentanan yang dirilis, mengurangi dampak eksploitasi, serta menangani akar penyebab agar masalah serupa tidak berulang.",
            "NIST SP 800-115 membahas perencanaan assessment, pelaksanaan pengujian teknis, analisis temuan, dan pengembangan strategi mitigasi [19]. OWASP WSTG menyediakan kategori pengujian web seperti identity, authentication, authorization, session, input validation, error handling, business logic, client-side, dan API [7]. ASVS melengkapi WSTG sebagai basis persyaratan verifikasi kontrol [20].",
            "NagariSDLC memisahkan QA dan keamanan siber menjadi dua jalur independen. Pemisahan ini menghindari asumsi bahwa lulus pengujian fungsional berarti aman. Setiap jalur memiliki pengajuan, disposisi, laporan pelaksana, sign-off Lead, kegagalan, return round, task perbaikan, dan resubmission gate. Tested scenarios disimpan sebagai teks karena ruang lingkup perlu disesuaikan dengan risiko proyek.",
        ]),
        ("3.17 Audit Trail, Traceability, dan Tata Kelola Dokumen", [
            "Traceability menghubungkan kebutuhan dengan desain, implementasi, pengujian, temuan, perbaikan, dan keputusan rilis. Keterlacakan dua arah membantu menjawab dua pertanyaan: apakah setiap kebutuhan telah diwujudkan dan diuji, serta apakah setiap artefak implementasi memiliki alasan kebutuhan. Pada proyek yang melibatkan banyak approval, hubungan tersebut mengurangi risiko keputusan berdasarkan versi artefak yang keliru.",
            "Audit trail merekam siapa melakukan apa, kapan, terhadap objek mana, dari kondisi apa menjadi kondisi apa, dan dengan hasil apa. Audit trail harus dilindungi dari perubahan atau penghapusan yang tidak sah. Soft delete, immutable event, status superseded, foreign key RESTRICT, dan snapshot approval adalah beberapa mekanisme untuk menjaga konteks historis.",
            "NagariSDLC menyediakan project_status_histories, activity_logs, test_reports, uat_approval_rounds, uat_approvers, project_return_rounds, release_requests, dan Document Vault. Putaran approval yang tidak berlaku ditandai superseded, bukan dihapus. Baris approved tetap disimpan. Dokumen yang masih menjadi evidence tidak dapat dihapus sembarangan. Rancangan ini memungkinkan kronologi keputusan direkonstruksi sekaligus mempertahankan kompatibilitas data lama.",
        ]),
        ("3.18 Pemodelan UML, ERD, dan Mermaid", [
            "UML merupakan bahasa pemodelan standar untuk merepresentasikan struktur dan perilaku sistem. Spesifikasi UML 2.5.1 mencakup elemen untuk use case, activity, state machine, sequence, class, dan bentuk model lain [24]. Fowler menekankan bahwa diagram seharusnya dipilih sesuai tujuan komunikasi, bukan dibuat sebanyak mungkin [31]. Model yang baik menyederhanakan aspek penting tanpa mengklaim menampilkan seluruh detail implementasi.",
            "ERD menggambarkan entitas, atribut kunci, dan kardinalitas hubungan data. State diagram menyoroti status serta transisi; sequence diagram menjelaskan urutan pesan; use case menjelaskan interaksi aktor dengan tujuan sistem; class diagram menjelaskan struktur tipe serta relasi. Satu diagram tidak dapat menggantikan diagram lain karena masing-masing menjawab pertanyaan berbeda.",
            "Mermaid menggunakan teks sebagai sumber diagram sehingga perubahan dapat ditelusuri dan dirender ulang [13]. Dalam laporan ini, kode Mermaid dipisahkan pada berkas pendamping. Placeholder di dokumen menentukan lokasi gambar, sementara kode dapat dirender di Mermaid Live, diekspor sebagai SVG/PNG resolusi tinggi, dan disesuaikan orientasinya agar terbaca pada media cetak.",
        ]),
        ("3.19 Notifikasi, Realtime, Queue, dan Operasional", [
            "Notifikasi membantu pengguna mengetahui perubahan yang memerlukan perhatian, tetapi volume yang berlebihan dapat menimbulkan alert fatigue. Notifikasi perlu memiliki penerima yang tepat, konteks proyek, jenis peristiwa, status baca, dan tautan menuju tindakan. Event bisnis sebaiknya dipisahkan dari cara penyampaiannya agar satu kejadian dapat dikirim melalui database notification, email, atau kanal realtime tanpa menduplikasi aturan domain.",
            "Queue memindahkan pekerjaan yang tidak harus selesai dalam response request, misalnya email, pemrosesan file, dan integrasi eksternal. Broadcasting mengirim pembaruan ke klien secara realtime. Pola ini meningkatkan responsivitas, tetapi menambah kebutuhan worker, retry, idempotency, observability, dan pengelolaan kegagalan [4], [32].",
            "NagariSDLC telah memiliki notification database dan rancangan Reverb, tetapi konfigurasi environment aktif masih menggunakan broadcast log. Karena itu laporan tidak mengklaim realtime produksi telah berjalan. Chat menggunakan polling. Aktivasi Reverb dan queue worker perlu didahului keputusan arsitektur produksi, kapasitas, keamanan koneksi, monitoring, serta prosedur pemulihan.",
        ]),
        ("3.20 Sintesis Literatur dan Posisi NagariSDLC", [
            "Literatur menunjukkan tiga kebutuhan yang sering tarik-menarik: fleksibilitas menghadapi perubahan, disiplin lifecycle untuk menghasilkan bukti, dan integrasi keamanan sejak awal. Agile serta continuous software engineering mendorong umpan balik dan aliran kerja yang tidak terputus [25]–[28]. ISO 12207 dan 29119 menekankan proses lifecycle serta testing yang dapat ditata [15], [17]. SSDF, ASVS, dan WSTG menempatkan keamanan sebagai tanggung jawab sepanjang pengembangan [7], [18]–[20].",
            "NagariSDLC mengambil posisi sebagai governance-enabling system. Sistem tidak menggantikan keahlian analis, developer, tester, pentester, PM, atau pengambil keputusan. Sistem membuat batas wewenang, urutan, prasyarat, evidence, dan histori dapat dilihat serta ditegakkan. Nilainya tidak hanya berupa otomatisasi administrasi, tetapi pengurangan ruang bagi transisi tanpa bukti dan keputusan yang tidak memiliki konteks.",
            "Kesenjangan yang masih ada juga terlihat dari literatur. Otomasi CI/CD, pengukuran karakteristik kualitas, baseline ASVS, threat modeling, monitoring produksi, disaster recovery, dan evaluasi usability formal belum menjadi keluaran lengkap proyek. Kesenjangan tersebut tidak menghapus hasil saat ini; justru menjadi dasar roadmap agar tata kelola berkembang dari workflow internal menuju ekosistem software delivery yang lebih terukur.",
        ]),
    ]
    for heading, paragraphs in theory_sections:
        add_heading(doc, heading, 2)
        for paragraph in paragraphs:
            add_para(doc, paragraph)

    add_heading(doc, "3.21 Perbandingan Model Pengembangan", 2)
    model_rows = [
        ["Sekuensial", "Kebutuhan stabil dan gate formal", "Dokumen dan tahap jelas", "Perubahan terlambat relatif mahal"],
        ["Incremental", "Produk dapat dibagi menjadi bagian", "Nilai dapat diberikan bertahap", "Integrasi dan prioritas perlu dijaga"],
        ["Spiral", "Risiko tinggi dan ketidakpastian", "Risiko menjadi penggerak iterasi", "Memerlukan kemampuan analisis risiko"],
        ["Agile", "Perubahan dan umpan balik sering", "Adaptif dan kolaboratif", "Butuh disiplin teknis dan keterlibatan aktif"],
        ["Hibrida governance", "Kebutuhan iteratif dengan audit formal", "Fleksibel tetapi tetap terlacak", "Gate dapat menjadi bottleneck bila kriterianya kabur"],
    ]
    add_table(doc, ["Model", "Konteks", "Kekuatan", "Risiko/Keterbatasan"], model_rows, [1400, 2050, 2100, 2388], "Tabel 3.1 Perbandingan model pengembangan", "tbl_3_1", bid); bid += 1

    add_heading(doc, "3.22 Pemetaan Karakteristik Kualitas", 2)
    quality_rows = [
        ["Functional suitability", "Cakupan fungsi dan ketepatan workflow", "Kebutuhan fungsional, status, prasyarat, dan test case"],
        ["Performance efficiency", "Waktu respons dan penggunaan sumber daya", "Belum diukur formal; perlu load/performance test"],
        ["Compatibility", "Koeksistensi dan interoperabilitas", "REST/JSON, Bearer compatibility, pembacaan data legacy"],
        ["Interaction capability", "Kemudahan interaksi pengguna", "Wizard, filter, pesan blocker, Bahasa Indonesia"],
        ["Reliability", "Konsistensi dan pemulihan", "Transaksi, constraint, retry operasional yang direncanakan"],
        ["Security", "Kerahasiaan, integritas, akuntabilitas", "HttpOnly, RBAC, project scope, audit trail, jalur Siber"],
        ["Maintainability", "Kemudahan analisis dan perubahan", "Form Request, service, Resource, API terpusat"],
        ["Flexibility", "Kemampuan beradaptasi", "Status paralel, return round, JSON wizard kompatibel"],
        ["Safety", "Pencegahan dampak yang tidak dapat diterima", "Tidak dinilai sebagai sistem safety-critical; release gate mengurangi risiko perubahan"],
    ]
    add_table(doc, ["Karakteristik ISO 25010", "Makna pada Sistem", "Bukti/Implikasi NagariSDLC"], quality_rows, [2150, 2450, 3338], "Tabel 3.2 Pemetaan model kualitas ISO/IEC 25010:2023", "tbl_3_2", bid); bid += 1

    add_heading(doc, "3.23 Ringkasan Landasan Teori", 2)
    theory_rows = [
        ["SDLC dan standar lifecycle", "Proses dari kebutuhan hingga operasi dapat iteratif dan paralel", "Workflow end-to-end dan quality gate"],
        ["Agile/continuous engineering", "Umpan balik, adaptasi, dan kontinuitas aktivitas", "Task iteratif serta revisi tanpa kehilangan histori"],
        ["Arsitektur dan REST", "Pemisahan tanggung jawab dan antarmuka seragam", "React–Laravel–MySQL dan envelope API"],
        ["RBAC dan least privilege", "Akses berdasarkan peran, assignment, serta constraint", "12 role, project scope, dan separation of duties"],
        ["Testing dan kualitas", "Bukti berbasis risiko pada berbagai tingkat", "SIT/UAT bertahap dan QA independen"],
        ["Secure SDLC", "Keamanan terintegrasi sepanjang lifecycle", "Jalur Siber, HttpOnly, audit, dan return round"],
        ["Traceability", "Kebutuhan–implementasi–test–keputusan terhubung", "Histori, approval round, dokumen, dan release request"],
        ["UML/Mermaid", "Model dipilih sesuai pertanyaan yang dijawab", "Flow, state, sequence, class, use case, dan ERD"],
    ]
    add_table(doc, ["Konsep", "Pokok Literatur", "Penerapan pada NagariSDLC"], theory_rows, [1900, 2920, 3118], "Tabel 3.3 Sintesis landasan teori", "tbl_3_3", bid); bid += 1

    add_heading(doc, "3.24 Matriks Literatur dan Relevansi Penelitian", 2)
    literature_rows = [
        ["Boehm (1988) [29]", "Spiral berbasis risiko", "Risiko perlu mendorong iterasi", "Return round dan revisi mayor memicu siklus kerja baru"],
        ["Dybå & Dingsøyr (2008) [26]", "Systematic review agile", "Manfaat agile bergantung konteks", "Governance tidak mengunci satu metode pengembangan"],
        ["Fitzgerald & Stol (2017) [27]", "Continuous software engineering", "Kurangi keterputusan antarfase", "Data analisis–task–test–rilis berada pada satu rantai"],
        ["Lwakatare et al. (2019) [28]", "Multiple case study DevOps", "Kolaborasi development–operations mendukung update andal", "Menjadi dasar roadmap integrasi CI/CD dan operasi"],
        ["Sandhu et al. (1996) [21]", "Model RBAC", "Role, permission, hierarchy, dan constraint", "12 role ditambah cakupan proyek serta assignment"],
        ["NIST SSDF (2022) [18]", "Secure development practices", "Security terintegrasi pada SDLC", "Jalur Siber, kontrol sesi, audit, dan roadmap secure pipeline"],
        ["ISO 29119-2 (2021) [17]", "Proses pengujian", "Testing perlu direncanakan, dilaksanakan, dan dilaporkan", "Wizard SIT/UAT serta laporan QA/Siber"],
        ["ISO 25010 (2023) [16]", "Model kualitas produk", "Kualitas dinilai melalui karakteristik", "Evaluasi kualitatif dan identifikasi metrik yang belum ada"],
    ]
    add_table(doc, ["Sumber", "Fokus", "Temuan/Gagasan", "Relevansi"], literature_rows, [1850, 1750, 2200, 2138], "Tabel 3.4 Matriks literatur dan relevansi NagariSDLC", "tbl_3_4", bid); bid += 1
    add_para(doc, "Matriks menunjukkan bahwa NagariSDLC tidak dibangun dari satu teori tunggal. Model risiko menjelaskan perlunya siklus revisi; agile dan continuous engineering menjelaskan kebutuhan adaptasi serta aliran informasi; RBAC menjelaskan pembagian akses; standar testing dan secure development menjelaskan evidence serta keamanan; model kualitas menyediakan kerangka evaluasi. Sintesis lintas sumber tersebut menjadi dasar pembahasan Bab IV.")

    build_chapter_four(doc, bid)


def build_chapter_four(doc: Document, bid: int):
    # Implementasi dipisahkan karena Bab IV memuat mayoritas tabel dan placeholder.
    add_chapter_heading(doc, "IV", "HASIL DAN PEMBAHASAN")
    add_heading(doc, "4.1 Hasil", 2)
    add_para(doc, "Hasil Kerja Praktek berupa sistem informasi NagariSDLC beserta dokumentasi alur, model data, endpoint, dan mekanisme pengujian. Bagian ini mengikuti urutan analisis proses bisnis, perancangan sistem, implementasi, dan verifikasi agar hubungan antara kebutuhan dan hasil dapat ditelusuri.")

    add_heading(doc, "4.1.1 Analisis Proses Bisnis", 3)
    add_para(doc, "Sebelum tersedia satu sistem governance terpusat, informasi kebutuhan, status, task, dokumen, pengujian, dan approval berisiko tersebar pada media yang berbeda. Kondisi tersebut membuat pemangku kepentingan harus melakukan konfirmasi manual dan menyulitkan pelacakan bukti keputusan. Gambar 4.1 menggambarkan kondisi konseptual tersebut dan bukan klaim mengenai prosedur resmi instansi yang belum diberikan.")
    add_figure_placeholder(doc, "Gambar 4.1 Proses bisnis sebelum sistem terpusat", "fig_4_1", bid, "Render kode Mermaid D-01A. Sesuaikan istilah dengan hasil wawancara/proses resmi sebelum finalisasi.", "D-01A", lines=7); bid += 1
    add_para(doc, "Proses yang diusulkan menempatkan NagariSDLC sebagai pusat koordinasi. Business User mengajukan kebutuhan, pihak perencanaan melakukan review dan analisis, pengembangan mengalokasikan tim serta task, pengujian internal dilakukan melalui SIT dan UAT, jalur QA dan keamanan siber berjalan paralel, kemudian pengajuan rilis diputuskan pada quality gate. Semua perpindahan status terjadi melalui service workflow.")
    add_figure_placeholder(doc, "Gambar 4.2 Proses bisnis yang diusulkan", "fig_4_2", bid, "Render kode Mermaid D-01 untuk menampilkan alur end-to-end NagariSDLC.", "D-01", lines=8); bid += 1

    add_heading(doc, "4.1.2 Aktor dan Lingkup Sistem", 3)
    actor_rows = [
        ["Business User", "Mengajukan kebutuhan, memantau proyek, dan menjadi pihak peminta pada UAT."],
        ["Lead Group / Analyst", "Review dan analisis kebutuhan pada fase perencanaan."],
        ["Development Lead / PM", "Disposisi pengembangan, alokasi tim, task, SIT/UAT, dan pengajuan rilis."],
        ["Developer", "Mengerjakan task, revisi, bukti SIT, dan approval yang ditugaskan."],
        ["QA Lead / QA Tester", "Disposisi, pelaksanaan, laporan, dan sign-off QA."],
        ["Cyber Lead / Pentester", "Disposisi, pelaksanaan, laporan, dan sign-off keamanan siber."],
        ["Head of IT", "Mengambil keputusan quality gate sebelum live production."],
        ["Super Admin", "Mengelola pengguna, divisi, grup, role, dan akses administrasi."],
    ]
    add_table(doc, ["Aktor", "Tanggung Jawab"], actor_rows, [2200, 5738], "Tabel 4.1 Aktor dan tanggung jawab sistem", "tbl_4_1", bid); bid += 1

    add_heading(doc, "4.1.3 Kebutuhan Fungsional", 3)
    functional_rows = [
        ["F-01", "Autentikasi dan sesi", "Login, logout, refresh, profil, ganti/lupa/reset kata sandi."],
        ["F-02", "Pengajuan proyek", "Mencatat kebutuhan, pemohon, divisi, prioritas, RBB/Non-RBB, dan target."],
        ["F-03", "Workflow proyek", "Memvalidasi transisi status, peran, penugasan, prasyarat, histori, dan notifikasi."],
        ["F-04", "Manajemen task", "CRUD task, assignee, status, prioritas, revisi, dan task perbaikan."],
        ["F-05", "SIT", "Persiapan, eksekusi per task, bukti, revisi, review, dan sign-off."],
        ["F-06", "UAT", "Skenario, peserta, undangan, draft/final eksekusi, Change Request, dan approval."],
        ["F-07", "QA", "Pengajuan, disposisi, laporan, bukti, sign-off, dan pengembalian."],
        ["F-08", "Keamanan siber", "Pentest/secure code, disposisi, laporan, bukti, sign-off, dan pengembalian."],
        ["F-09", "Rilis", "Pengajuan target rilis, downtime, rollback, quality gate, dan produksi."],
        ["F-10", "Dokumen", "Upload, masking nama, daftar, unduh, dan penghapusan yang melindungi bukti."],
        ["F-11", "Komunikasi", "Chat proyek, notifikasi, dan informasi aktivitas."],
        ["F-12", "Administrasi", "Pengguna, divisi, grup, role, dan pembatasan menu."],
    ]
    add_table(doc, ["ID", "Kebutuhan", "Deskripsi"], functional_rows, [650, 1900, 5388], "Tabel 4.2 Kebutuhan fungsional", "tbl_4_2", bid); bid += 1

    add_heading(doc, "4.1.4 Kebutuhan Nonfungsional", 3)
    nonfunctional_rows = [
        ["Keamanan", "Cookie HttpOnly untuk sesi SPA, RBAC, project scope, validasi input, CORS berbasis environment, serta proteksi audit trail."],
        ["Keterlacakan", "Histori status, activity log, test report, approval round, return round, dan document evidence."],
        ["Konsistensi", "Envelope API seragam dan satu service API pada front end."],
        ["Reliabilitas", "Transaksi database pada operasi workflow dan pengujian yang mengubah beberapa tabel."],
        ["Kompatibilitas", "Dukungan Bearer token dan pembacaan data legacy tanpa menghapus histori."],
        ["Usability", "Antarmuka Bahasa Indonesia, responsif, pencarian, filter, wizard, dan status yang mudah dipahami."],
        ["Maintainability", "Pemisahan route, Form Request, controller, Resource, model, dan service."],
        ["Deployability", "Template environment produksi tersedia; keputusan infrastruktur final masih perlu ditetapkan."],
    ]
    add_table(doc, ["Aspek", "Kebutuhan"], nonfunctional_rows, [1800, 6138], "Tabel 4.3 Kebutuhan nonfungsional", "tbl_4_3", bid); bid += 1

    add_heading(doc, "4.1.5 Arsitektur Sistem", 3)
    add_para(doc, "Lapis klien berupa React 19 SPA yang dibangun dengan Vite 8 dan Tailwind CSS 4. Lapis aplikasi berupa Laravel 13 pada PHP 8.3 dengan REST API, Sanctum, Form Request, Resource, service, event, dan notification. Lapis data menggunakan MySQL melalui Eloquent ORM serta penyimpanan dokumen pada disk yang dikonfigurasi. Front end memanggil API melalui frontend/src/services/api.js. Respons mengikuti format { status, message, data, meta? }.")
    add_para(doc, "Autentikasi aktif memakai cookie HttpOnly nagari_sdlc_token sebagai jalur utama SPA dengan credentials include. Header X-Requested-With diwajibkan pada jalur cookie. Header Authorization Bearer tetap dipertahankan untuk klien lama dan pengujian. localStorage hanya memuat metadata sesi dan profil yang diperlukan antarmuka, bukan token. Broadcasting Reverb tersedia dalam rancangan, tetapi template environment saat ini memakai BROADCAST_CONNECTION=log sehingga pengaktifan realtime produksi memerlukan keputusan dan konfigurasi operasional.")
    add_figure_placeholder(doc, "Gambar 4.3 Arsitektur sistem", "fig_4_3", bid, "Render kode Mermaid D-03 untuk arsitektur klien, API, data, dokumen, dan realtime.", "D-03", lines=8); bid += 1

    add_heading(doc, "4.1.6 Model Peran dan Use Case", 3)
    roles = [
        ["super_admin", "Super Admin", "Manajemen TI", "Administrasi dan akses global."],
        ["head_of_it", "Head of IT", "Manajemen TI", "Quality gate dan pengawasan."],
        ["lead_group", "Lead Group", "Perencanaan-QA", "Review dan pengawasan grup."],
        ["analyst", "Analyst", "Perencanaan-QA", "Analisis kebutuhan Fase 1."],
        ["development_lead", "Development Lead", "Pengembangan", "Disposisi dan pengawasan pengembangan."],
        ["project_manager", "Project Manager/Dev Analyst", "Pengembangan", "Pengelolaan proyek Fase 2."],
        ["developer", "Developer", "Pengembangan", "Implementasi dan perbaikan."],
        ["qa_lead", "QA Lead", "Perencanaan-QA", "Disposisi dan sign-off QA."],
        ["qa_tester", "QA Tester", "Perencanaan-QA", "Pelaksanaan pengujian QA."],
        ["cyber_lead", "Cyber Lead", "Keamanan Siber", "Disposisi dan sign-off Siber."],
        ["pentester", "Pentester", "Keamanan Siber", "Pentest atau secure code review."],
        ["business_user", "Business User", "Pemohon", "Pengajuan dan penerimaan UAT."],
    ]
    add_table(doc, ["Role", "Nama Tampilan", "Grup Bawaan", "Tanggung Jawab"], roles, [1500, 1850, 1700, 2888], "Tabel 4.4 Daftar 12 peran pengguna", "tbl_4_4", bid); bid += 1
    add_para(doc, "Lima grup bawaan—Perencanaan-QA, Pengembangan, Keamanan Siber, Manajemen TI, dan Pemohon—mengelompokkan role untuk tampilan. Grup tidak memberi wewenang. Hak transisi tetap berada di ProjectWorkflowService, cakupan proyek di ProjectAccessService, dan hak pelaksana pengujian di TestingTrack. Role baru yang dibuat melalui administrasi belum otomatis berfungsi karena daftar otorisasi tetap berada di kode.")
    add_figure_placeholder(doc, "Gambar 4.4 Use case aktor utama", "fig_4_4", bid, "Render kode Mermaid D-04 untuk relasi aktor dan fungsi utama.", "D-04", lines=8); bid += 1

    add_heading(doc, "4.1.7 Mesin Status Proyek", 3)
    statuses = [
        "PENDING", "IN_REVIEW", "ANALYSIS_APPROVED", "REJECTED", "READY_FOR_DEVELOPMENT", "DEV_ANALYSIS", "DEV_ANALYSIS_DONE",
        "IN_DEVELOPMENT", "SIT_IN_PROGRESS", "SIT_PASSED", "SIT_REVISION", "UAT_IN_PROGRESS", "UAT_REVISION_SIT", "UAT_REVISION_DEV",
        "DEV_COMPLETED", "READY_FOR_QA", "QA_IN_PROGRESS", "RETURN_TO_DEV", "QA_PASSED", "CYBER_IN_PROGRESS", "CYBER_PASSED",
        "READY_FOR_UAT", "UAT_PASSED", "PENDING_GOLIVE", "LIVE_PRODUCTION", "ON_HOLD", "CANCELLED",
    ]
    phase_map = {
        "PENDING": "Inisiasi", "IN_REVIEW": "Review", "ANALYSIS_APPROVED": "Analisis", "REJECTED": "Nonlinear",
        "READY_FOR_DEVELOPMENT": "Pengembangan", "DEV_ANALYSIS": "Pengembangan", "DEV_ANALYSIS_DONE": "Pengembangan", "IN_DEVELOPMENT": "Pengembangan",
        "SIT_IN_PROGRESS": "SIT", "SIT_PASSED": "SIT", "SIT_REVISION": "SIT", "UAT_IN_PROGRESS": "UAT", "UAT_REVISION_SIT": "Legacy/revisi",
        "UAT_REVISION_DEV": "UAT/revisi", "DEV_COMPLETED": "Pasca-UAT", "READY_FOR_QA": "QA/Siber", "QA_IN_PROGRESS": "QA",
        "RETURN_TO_DEV": "Pengembalian", "QA_PASSED": "QA", "CYBER_IN_PROGRESS": "Siber", "CYBER_PASSED": "Siber",
        "READY_FOR_UAT": "Legacy", "UAT_PASSED": "Legacy/opsional", "PENDING_GOLIVE": "Rilis", "LIVE_PRODUCTION": "Produksi",
        "ON_HOLD": "Nonlinear", "CANCELLED": "Terminal",
    }
    status_rows = [[str(i), status, phase_map[status]] for i, status in enumerate(statuses, 1)]
    add_table(doc, ["No.", "Status", "Fase/Makna"], status_rows, [600, 3300, 4038], "Tabel 4.5 Daftar 27 status proyek", "tbl_4_5", bid); bid += 1
    add_para(doc, "READY_FOR_UAT dipertahankan sebagai status legacy agar histori lama tetap terbaca, tetapi tidak memiliki transisi masuk maupun keluar. LIVE_PRODUCTION dan CANCELLED bersifat terminal pada matriks aktif. Status revisi dapat membawa proyek kembali ke pengembangan; revisi tidak direpresentasikan sebagai pembatalan.")
    add_figure_placeholder(doc, "Gambar 4.5 Mesin status proyek", "fig_4_5", bid, "Render kode Mermaid D-05. Untuk laporan cetak, ekspor dalam orientasi landscape atau bagi menjadi dua bagian bila teks terlalu kecil.", "D-05", lines=10); bid += 1

    add_heading(doc, "4.1.8 Alur System Integration Test", 3)
    add_para(doc, "SIT dimulai setelah seluruh task aktif selesai; task TAKE DOWN tidak dihitung. Tahap 1 menyimpan URL staging dan menyiapkan skenario. Tahap 2 menampilkan setiap task, status OK, komentar, bukti, dan revisi. Draft dapat disimpan tanpa berpindah tahap. Tahap 3 memuat ringkasan, catatan review, dokumen hasil review/berita acara, serta approval seluruh developer tim, PM, dan Development Lead.")
    add_para(doc, "Jika UAT menemukan revisi mayor, SIT diulang secara menyeluruh pada seluruh task aktif kecuali TAKE DOWN. Perbedaannya bukan pada daftar task, melainkan pada ketatnya bukti: task harus selesai, di-OK, memiliki lampiran yang sah, dan approval tim harus lengkap. Bukti per task dibekukan setelah SIT lulus agar berita acara tidak berubah setelah keputusan.")
    add_figure_placeholder(doc, "Gambar 4.6 Alur SIT", "fig_4_6", bid, "Render kode Mermaid D-06 untuk gate, tiga tahap, revisi, dan SIT ulang.", "D-06", lines=9); bid += 1

    add_heading(doc, "4.1.9 Alur User Acceptance Test", 3)
    add_para(doc, "Tahap 1 UAT memuat skenario dari task, unit peminta, tanggal pelaksanaan, penyiap, peserta, dan undangan. Peserta menjadi sumber kandidat approver. Roster uat1_participants tidak pernah dikosongkan oleh jalur kode; PM boleh menambah atau memperbaiki, tetapi pengosongan ditolak karena penanda tangan harus terbawa ke setiap putaran revisi.")
    add_para(doc, "Tahap 2 mencatat hasil setiap skenario sebagai accepted atau revision. Revisi minor membuat Change Request dan membuka kembali task tanpa rollback atau SIT ulang, tetapi menahan approval sampai seluruh perbaikan resolved. Revisi mayor mengarsipkan putaran UAT, memindahkan proyek ke UAT_REVISION_DEV, membuka task perbaikan, menjalankan SIT ulang menyeluruh, lalu mengulang UAT dari Tahap 1. Tahap 3 membuka putaran approval baru atas hasil terbaru.")
    add_para(doc, "Kompatibilitas penanda pengulangan UAT pada data legacy dibaca hanya melalui Project::isUatRestartPending(), bukan dengan mengakses key JSON yang telah dipensiunkan secara langsung.")
    add_figure_placeholder(doc, "Gambar 4.7 Alur UAT", "fig_4_7", bid, "Render kode Mermaid D-07 untuk persiapan, eksekusi, revisi minor/mayor, dan approval.", "D-07", lines=10); bid += 1

    add_heading(doc, "4.1.10 Matriks Persetujuan UAT", 3)
    approvers = [
        ["requester", "Pihak peminta", "External link", "Ya, satu"],
        ["requester_group_lead", "Pihak peminta", "External link", "Ya, satu"],
        ["requester_division_lead", "Pihak peminta", "External link", "Ya, satu"],
        ["developer", "Pihak IT", "Internal account", "Minimal satu"],
        ["analyst_pm", "Pihak IT", "Internal account", "Ya, satu"],
        ["development_group_lead", "Pihak IT", "Internal account", "Ya, satu"],
        ["technology_division_lead", "Pihak IT", "Internal account", "Ya, satu"],
    ]
    add_table(doc, ["Approval Role", "Sisi", "Mode", "Kewajiban"], approvers, [2250, 1500, 1900, 2288], "Tabel 4.6 Matriks approver UAT", "tbl_4_6", bid); bid += 1
    add_para(doc, "Approval aktif berada pada uat_approval_rounds dan uat_approvers. Pihak peminta memakai link pribadi dan verifikasi nomor HP, sedangkan pihak IT memakai akun. Setiap orang memberikan keputusan secara individual dan dapat paralel. Putaran yang tidak berlaku lagi ditandai superseded; token dicabut, baris pending menjadi revoked, dan baris approved tetap disimpan sebagai audit trail.")
    add_figure_placeholder(doc, "Gambar 4.9 Relasi approver UAT", "fig_4_9", bid, "Render kode Mermaid D-09 untuk sisi peminta, sisi IT, mode akses, dan putaran approval.", "D-09", lines=8); bid += 1

    add_heading(doc, "4.1.11 Jalur QA dan Keamanan Siber", 3)
    add_para(doc, "Setelah DEV_COMPLETED, PM mengajukan proyek ke QA dan keamanan siber. Masing-masing jalur memiliki kolom status sendiri: qa_status dan cyber_status. Empat langkah pada setiap jalur adalah pengajuan, disposisi oleh Lead, laporan pelaksana, dan sign-off Lead. Cakupan pengujian ditulis pada test_reports.tested_scenarios sebagai teks bebas; checklist tetap hanya dipertahankan untuk membaca laporan lama.")
    add_para(doc, "Sign-off fail menandai jalur FAILED, memindahkan proyek ke RETURN_TO_DEV, dan membuka project_return_rounds. PM harus membuat task perbaikan yang menunjuk return round. Pengajuan ulang ditolak bila belum ada task, masih ada task tanpa assignee, atau ada task yang belum done/take_down. Penilaian dilakukan per jalur sehingga pengembalian Siber tidak otomatis menahan QA yang belum berjalan, dan sebaliknya.")
    add_figure_placeholder(doc, "Gambar 4.8 Jalur QA dan keamanan siber", "fig_4_8", bid, "Render kode Mermaid D-08 untuk dua jalur paralel dan empat langkah per jalur.", "D-08", lines=9); bid += 1
    add_figure_placeholder(doc, "Gambar 4.10 Putaran pengembalian pengujian", "fig_4_10", bid, "Render kode Mermaid D-10 untuk kegagalan, return round, task perbaikan, dan pengajuan ulang.", "D-10", lines=8); bid += 1

    add_heading(doc, "4.1.12 Gerbang Rilis", 3)
    add_para(doc, "PENDING_GOLIVE hanya dapat dicapai bila qa_status dan cyber_status sama-sama PASSED. PM mengajukan release request yang memuat target tanggal, estimasi downtime, rollback plan, dan catatan. Head of IT melihat antrean quality gate, lalu menyetujui menuju LIVE_PRODUCTION atau menolak dengan alasan yang tercatat. Tidak terdapat UAT final setelah kedua jalur pengujian lulus.")

    add_heading(doc, "4.1.13 Model Data dan Relationship", 3)
    domain_tables = [
        ["users", "Akun, role, divisi, nomor HP, status aktif, soft delete"],
        ["roles", "Nama role, nama tampilan, grup, dan pembatasan menu"],
        ["groups", "Pengelompokan role untuk tampilan"],
        ["divisions", "Master unit/divisi pengguna dan proyek"],
        ["projects", "Entitas pusat proyek, status, penugasan, jalur uji, dan sit_uat_data"],
        ["project_tasks", "Task, assignee, status, revisi, dan return_round_id"],
        ["project_team_members", "Keanggotaan tim proyek"],
        ["project_status_histories", "Riwayat transisi status"],
        ["project_return_rounds", "Putaran pengembalian QA/Siber"],
        ["test_reports", "Laporan pengujian dan sign-off Lead"],
        ["uat_approval_rounds", "Putaran persetujuan final UAT"],
        ["uat_approvers", "Approver individual, mode akses, dan keputusan"],
        ["release_requests", "Pengajuan, rencana, dan keputusan rilis"],
        ["document_vaults", "Metadata dan lokasi dokumen"],
        ["chat_messages", "Pesan diskusi proyek"],
        ["activity_logs", "Aktivitas pengguna dan metadata audit"],
        ["notifications", "Kotak masuk pemberitahuan pengguna"],
    ]
    add_table(doc, ["Tabel", "Fungsi"], domain_tables, [2600, 5338], "Tabel 4.7 Tabel domain utama", "tbl_4_7", bid); bid += 1
    add_para(doc, "Selain 17 tabel domain, terdapat delapan tabel framework: sessions, cache, cache_locks, jobs, job_batches, failed_jobs, password_reset_tokens, dan personal_access_tokens. Dengan demikian model data mencakup 25 tabel. projects.sit_uat_data menyimpan data wizard yang berstruktur JSON, sedangkan approval aktif UAT dipindahkan ke tabel terstruktur agar snapshot per putaran dan keputusan per orang dapat diaudit.")
    add_figure_placeholder(doc, "Gambar 4.11 ERD NagariSDLC", "fig_4_11", bid, "Render kode Mermaid D-11. Bila terlalu padat, ekspor pada orientasi landscape dan gunakan resolusi tinggi.", "D-11", lines=10); bid += 1
    add_figure_placeholder(doc, "Gambar 4.12 Class relationship inti", "fig_4_12", bid, "Render kode Mermaid D-12 untuk relasi service, model, controller, dan resource inti.", "D-12", lines=8); bid += 1

    add_heading(doc, "4.1.14 Desain API dan Sequence", 3)
    api_rows = [
        ["POST", "/auth/login", "Login dan penerbitan sesi Sanctum"],
        ["GET/POST", "/projects", "Daftar dan pengajuan proyek"],
        ["PATCH", "/projects/{id}/status", "Transisi melalui workflow"],
        ["GET/POST", "/projects/{id}/tasks", "Daftar dan pembuatan task"],
        ["POST", "/projects/{id}/sit-approval", "Approval SIT"],
        ["PUT/POST", "/projects/{id}/uat-execution[/draft]", "Draft/final eksekusi UAT"],
        ["POST", "/qa-requests/*", "Empat langkah jalur QA"],
        ["POST", "/cyber-requests/*", "Empat langkah jalur Siber"],
        ["POST", "/release-requests", "Pengajuan migrasi dan rilis"],
        ["POST", "/quality-gate/approve|reject", "Keputusan Head of IT"],
    ]
    add_table(doc, ["Method", "Endpoint", "Fungsi"], api_rows, [1100, 3400, 3438], "Tabel 4.8 Endpoint API representatif", "tbl_4_8", bid); bid += 1
    add_para(doc, "Seluruh error API menggunakan envelope seragam. Kegagalan validasi menambahkan errors per field; 401 menandai sesi tidak sah, 403 menandai larangan otorisasi, 404 untuk data tidak ditemukan, 422 untuk validasi atau aturan workflow, dan 500 untuk kegagalan internal. Detail implementasi endpoint ditempatkan pada Lampiran C.")
    add_figure_placeholder(doc, "Gambar 4.13 Sequence login", "fig_4_13", bid, "Render kode Mermaid D-13 untuk login, cookie HttpOnly, metadata sesi, dan request berikutnya.", "D-13", lines=8); bid += 1
    add_figure_placeholder(doc, "Gambar 4.14 Sequence persetujuan UAT eksternal", "fig_4_14", bid, "Render kode Mermaid D-14 untuk link pribadi, verifikasi nomor HP, token akses, detail, dan keputusan.", "D-14", lines=9); bid += 1

    add_heading(doc, "4.1.15 Perancangan Navigasi dan Antarmuka", 3)
    add_para(doc, "Navigasi ditentukan oleh role guard dan menuConfig. Pembatasan roles.menu_access hanya dapat mengurangi menu yang telah tersedia; menyembunyikan menu tidak menggantikan otorisasi route maupun backend. Halaman dikelompokkan menjadi dashboard, proyek, workspace per peran, task, SIT/UAT, QA/Siber, approval, release, quality gate, notifikasi, activity log, dan administrasi.")
    add_figure_placeholder(doc, "Gambar 4.15 Peta navigasi halaman", "fig_4_15", bid, "Render kode Mermaid D-15 untuk sitemap halaman utama berdasarkan peran.", "D-15", lines=9); bid += 1

    ui_specs = [
        (16, "Antarmuka login", "Masukkan tangkapan layar halaman login; sembunyikan email/akun nyata."),
        (17, "Antarmuka dashboard", "Masukkan dashboard yang menunjukkan kartu ringkasan sesuai role."),
        (18, "Daftar proyek", "Masukkan daftar proyek beserta pencarian/filter tanpa data sensitif."),
        (19, "Detail proyek", "Masukkan ringkasan status, metadata, timeline, dan aksi yang relevan."),
        (20, "Papan Kanban", "Masukkan papan task berdasarkan status todo, in progress, hold, done, dan take down."),
        (21, "Detail task", "Masukkan detail task, assignee, prioritas, due date, serta banner revisi bila ada."),
        (22, "Wizard SIT", "Masukkan rangkaian persiapan, eksekusi per task, dan review/sign-off."),
        (23, "Wizard UAT", "Masukkan persiapan peserta, eksekusi skenario, dan approval final."),
        (24, "Tugas QA", "Masukkan halaman pencarian/filter, laporan, bukti, dan status jalur QA."),
        (25, "Tugas keamanan siber", "Masukkan halaman pentest/secure code beserta ruang lingkup pengujian."),
        (26, "Matriks persetujuan UAT", "Masukkan daftar approver per orang, mode akses, dan keputusan."),
        (27, "Quality Gate", "Masukkan antrean PENDING_GOLIVE dan detail rencana rilis."),
        (28, "Manajemen pengguna", "Masukkan daftar akun, role, divisi, status aktif, dan aksi admin."),
        (29, "Manajemen grup dan role", "Masukkan pengelompokan role serta pembatasan menu."),
        (30, "Document Vault", "Masukkan daftar dokumen, tipe, nama masking, dan aksi yang diizinkan."),
        (31, "Activity Log", "Masukkan daftar aktivitas dengan filter yang relevan."),
    ]
    for no, title, instruction in ui_specs:
        add_heading(doc, f"4.1.15.{no-15} {title}", 3)
        add_para(doc, f"Placeholder berikut disediakan untuk dokumentasi {title.lower()}. Sebelum memasukkan gambar, potong area yang tidak relevan, gunakan resolusi yang cukup, dan sensor data internal, token, nomor telepon, alamat surel, atau dokumen rahasia.")
        add_figure_placeholder(doc, f"Gambar 4.{no} {title}", f"fig_4_{no}", bid, instruction, lines=5); bid += 1

    add_heading(doc, "4.1.16 Implementasi Back End", 3)
    add_para(doc, "Back end memakai pola route–Form Request–controller–service–model–Resource. Validasi write dilakukan pada Form Request. Logika bisnis lintas endpoint ditempatkan pada service. ProjectWorkflowService::transition() menjadi satu-satunya jalur transisi status proyek. TestingTrackService mengelola empat langkah QA/Siber, UatExecutionService mengelola draft/final UAT dan hold revisi, UatApprovalService mengelola approval round, sedangkan ProjectReturnRoundService mengelola putaran pengembalian.")
    add_para(doc, "Operasi yang mengubah beberapa data dibungkus transaksi. Status proyek, kolom jalur, histori, notifikasi, approval, dan return round tidak boleh berada pada keadaan setengah jadi. Resource menormalisasi data untuk front end, termasuk key task_ pada sit2_task_approvals agar JSON tetap berbentuk object. Endpoint penghapusan melindungi dokumen dan data audit yang masih dirujuk.")

    add_heading(doc, "4.1.17 Implementasi Front End", 3)
    add_para(doc, "Front end menggunakan React Router dengan ProtectedRoute, context untuk autentikasi, proyek, chat, notifikasi, dan master data, serta service API terpusat. Komponen wizard SIT/UAT mengatur tiga tahap, mode read-only, bukti per task/skenario, dan approval. Halaman QA/Siber menyediakan pencarian dan filter per assignee bagi role pengawas yang sesuai. UI menggunakan Bahasa Indonesia dan warna identitas #00529C.")
    add_para(doc, "Keputusan bisnis tidak dihitung ulang secara bebas pada front end. Contohnya, can_resubmit dan resubmit_blocker berasal dari server; angka ringkasan UAT dihitung backend; dan gate status tetap ditegakkan service. Front end bertugas menyajikan kondisi dan mencegah aksi yang pasti ditolak, sedangkan back end tetap menjadi sumber kebenaran otorisasi.")

    add_heading(doc, "4.1.18 Keamanan dan Audit Trail", 3)
    add_list(doc, [
        "Token sesi SPA disimpan pada cookie HttpOnly; Bearer tetap didukung sebagai jalur kompatibilitas.",
        "CORS dibatasi oleh origin environment dan mendukung credentials untuk cookie.",
        "Form Request dan middleware role memvalidasi input serta akses endpoint.",
        "ProjectAccessService mengisolasi proyek berdasarkan peran, penugasan, dan keterlibatan personal.",
        "Soft delete, RESTRICT, SET NULL, dan CASCADE dipilih sesuai makna audit setiap relasi.",
        "Approval dan histori tidak di-hard-delete; putaran lama ditandai completed, superseded, revoked, atau resubmitted.",
        "Document Vault mencegah penghapusan bukti yang masih dirujuk laporan atau berita acara.",
    ])

    add_heading(doc, "4.1.19 Pengujian dan Verifikasi", 3)
    verification_rows = [
        ["Backend test", "236 test; 1.467 assertions", "Lulus pada snapshot 26 Agustus 2026"],
        ["ESLint", "npx eslint src", "Tidak ada error pada snapshot"],
        ["Build front end", "npx vite build", "Berhasil; terdapat peringatan lama ukuran chunk"],
        ["Database test", "SQLite :memory:", "Tidak menyentuh database pengembangan"],
        ["Verifikasi laporan ini", "Audit source dan dokumentasi", "Tidak menjalankan ulang test/build aplikasi"],
    ]
    add_table(doc, ["Jenis", "Cakupan", "Hasil/Keterangan"], verification_rows, [1850, 2450, 3638], "Tabel 4.9 Ringkasan hasil verifikasi", "tbl_4_9", bid); bid += 1
    add_para(doc, "Snapshot kualitas bersifat historis dan dicantumkan untuk menggambarkan kondisi yang telah dilaporkan dokumentasi proyek. Karena penyusunan laporan tidak mengubah source code dan pengguna memilih menjalankan pengujian sendiri, test suite, ESLint, dan build tidak dijalankan ulang pada pekerjaan ini.")

    add_heading(doc, "4.1.20 Matriks Keterlacakan Kebutuhan", 3)
    add_para(doc, "Matriks keterlacakan menghubungkan kebutuhan dengan modul, data, kontrol, dan bukti verifikasi. Matriks ini bukan daftar seluruh endpoint, tetapi menunjukkan bahwa fungsi utama tidak berdiri sendiri. Setiap kebutuhan memiliki lokasi implementasi, objek data, dan bukti yang dapat digunakan untuk menilai hasil.")
    trace_rows = [
        ["F-01", "Auth dan profil", "users, personal_access_tokens, sessions", "Cookie/Bearer, middleware", "Test autentikasi dan respons sesi"],
        ["F-02", "Registrasi proyek", "projects, divisions", "Form Request, project scope", "Data pengajuan dan histori awal"],
        ["F-03", "Workflow", "projects, project_status_histories", "ProjectWorkflowService", "Transisi, from/to, actor, timestamp"],
        ["F-04", "Task", "project_tasks, team_members", "Assignee dan status gate", "Task selesai/take down dan histori revisi"],
        ["F-05", "SIT", "sit_uat_data, documents", "Tiga tahap dan sign-off", "Bukti task, review, berita acara"],
        ["F-06", "UAT", "sit_uat_data, approval_rounds, approvers", "Putaran, token, individual approval", "Skenario, CR, keputusan tiap orang"],
        ["F-07/F-08", "QA dan Siber", "test_reports, return_rounds", "Jalur independen dan resubmit gate", "Laporan, sign-off, task perbaikan"],
        ["F-09", "Rilis", "release_requests, projects", "Kedua jalur harus passed", "Rencana rilis dan keputusan Head of IT"],
        ["F-10/F-11", "Dokumen/komunikasi", "document_vaults, notifications, chat_messages", "Access scope dan proteksi evidence", "File, notifikasi, serta pesan proyek"],
        ["F-12", "Administrasi", "users, roles, groups, divisions", "Super admin dan menu restriction", "Perubahan master dan activity log"],
    ]
    add_table(doc, ["ID", "Modul", "Data Utama", "Kontrol", "Bukti"], trace_rows, [650, 1250, 1900, 1900, 2238], "Tabel 4.10 Matriks keterlacakan kebutuhan", "tbl_4_10", bid); bid += 1

    add_heading(doc, "4.1.21 Pemetaan Kontrol Keamanan", 3)
    add_para(doc, "Kontrol keamanan dipetakan terhadap risiko yang relevan, bukan digunakan untuk mengklaim kepatuhan penuh terhadap ASVS atau SSDF. Pemetaan menunjukkan titik kontrol yang sudah ada dan pekerjaan lanjutan yang masih diperlukan sebelum produksi.")
    security_rows = [
        ["Identitas dan sesi", "Pencurian/penyalahgunaan sesi", "Sanctum, cookie HttpOnly, logout, Bearer compatibility", "HTTPS, konfigurasi cookie, expiry, dan monitoring produksi"],
        ["Otorisasi", "Akses proyek atau aksi tanpa hak", "RBAC, middleware, ProjectAccessService, assignee gate", "Uji negatif per role dan review matriks akses"],
        ["Validasi", "Input tidak sah dan mass assignment", "Form Request, validation errors, model fillable", "Baseline ASVS dan fuzz/abuse case"],
        ["Integritas workflow", "Status dilompati atau bukti tidak lengkap", "ProjectWorkflowService dan prasyarat", "Property/transition coverage lebih luas"],
        ["Audit", "Histori diubah atau dihapus", "Soft delete, RESTRICT, superseded, activity log", "Retensi, immutability, dan ekspor audit"],
        ["Dokumen", "Akses atau penghapusan evidence", "Masking nama, project scope, referential protection", "Malware scan, object storage, encryption, retention"],
        ["Transport dan origin", "Penyadapan/permintaan lintas origin", "CORS environment dan credentials", "TLS, CSP, SameSite, CSRF review"],
        ["Operasional", "Kegagalan tak terdeteksi", "Log dan failed job framework", "Centralized logging, alerting, backup, DR drill"],
    ]
    add_table(doc, ["Area", "Risiko", "Kontrol Saat Ini", "Tindak Lanjut"], security_rows, [1450, 1850, 2500, 2138], "Tabel 4.11 Pemetaan kontrol dan risiko keamanan", "tbl_4_11", bid); bid += 1

    add_heading(doc, "4.1.22 Evaluasi Kualitas Produk", 3)
    quality_eval_rows = [
        ["Functional suitability", "Kuat", "12 role, 27 status, task, SIT/UAT, QA/Siber, dan rilis tercakup", "Validasi dengan pengguna resmi tetap diperlukan"],
        ["Performance efficiency", "Belum dinilai", "Belum ada angka response time, throughput, atau beban", "Lakukan performance test berbasis skenario"],
        ["Compatibility", "Cukup", "REST/JSON, Bearer, normalisasi key legacy", "Uji browser dan integrasi environment produksi"],
        ["Interaction capability", "Cukup", "Wizard, filter, pesan blocker, UI Bahasa Indonesia", "Usability test dan accessibility audit formal"],
        ["Reliability", "Cukup", "Transaksi dan constraint menjaga operasi multi-tabel", "Uji kegagalan, retry, backup, dan recovery"],
        ["Security", "Cukup", "HttpOnly, RBAC, project scope, audit, pengujian Siber", "Threat model, ASVS baseline, pentest independen"],
        ["Maintainability", "Kuat", "Service domain, Form Request, Resource, API service terpusat", "Pantau kompleksitas service dan coverage"],
        ["Flexibility", "Cukup", "Parallel track, return round, kompatibilitas legacy", "Konfigurasi role dinamis masih terbatas"],
        ["Safety", "Di luar fokus", "Bukan sistem kontrol keselamatan", "Tetap evaluasi risiko perubahan terhadap layanan bank"],
    ]
    add_table(doc, ["Karakteristik", "Penilaian Kualitatif", "Dasar", "Kesenjangan"], quality_eval_rows, [1750, 1350, 2850, 1988], "Tabel 4.12 Evaluasi kualitatif berdasarkan ISO/IEC 25010:2023", "tbl_4_12", bid); bid += 1
    add_para(doc, "Penilaian di atas bersifat analisis desain dan bukti proyek, bukan hasil sertifikasi. Istilah kuat atau cukup menunjukkan tingkat dukungan relatif berdasarkan artefak yang tersedia. Karakteristik yang belum memiliki metrik dinyatakan belum dinilai agar laporan tidak mengubah asumsi menjadi fakta.")

    add_heading(doc, "4.1.23 Keputusan Arsitektur dan Trade-off", 3)
    adr_rows = [
        ["State transition terpusat", "Konsistensi dan audit", "Service dapat menjadi kompleks", "Pisahkan validator dan tambahkan transition tests"],
        ["Approval UAT ternormalisasi", "Snapshot/keputusan per orang jelas", "Migrasi dari JSON legacy perlu kompatibilitas", "Helper normalisasi dan supersession"],
        ["SIT/UAT sebagian dalam JSON", "Fleksibel terhadap evolusi wizard", "Query dan constraint lebih terbatas", "Normalisasi bertahap bila pola stabil"],
        ["Cookie HttpOnly + Bearer", "Keamanan SPA dan kompatibilitas", "Dua jalur autentikasi perlu diuji", "Kontrak middleware dan dokumentasi eksplisit"],
        ["QA/Siber independen", "Status tidak ambigu dan dapat paralel", "Koordinasi menuju gate lebih kompleks", "Prasyarat kedua jalur passed"],
        ["Soft delete dan supersession", "Audit trail terlindungi", "Volume data bertambah", "Kebijakan retensi dan arsip resmi"],
        ["Reverb belum aktif", "Menghindari klaim operasional palsu", "Polling menambah request", "Aktifkan setelah kapasitas dan SOP diputuskan"],
    ]
    add_table(doc, ["Keputusan", "Manfaat", "Trade-off", "Mitigasi/Roadmap"], adr_rows, [1900, 1950, 2050, 2038], "Tabel 4.13 Ringkasan keputusan arsitektur", "tbl_4_13", bid); bid += 1

    add_heading(doc, "4.2 Pembahasan", 2)
    add_heading(doc, "4.2.1 Kesesuaian terhadap Tujuan", 3)
    add_para(doc, "Tujuan pertama adalah memusatkan tata kelola proyek teknologi dari pengajuan sampai produksi. Hasil implementasi menunjukkan bahwa kebutuhan, review, analisis, pengembangan, task, SIT, UAT, QA, keamanan siber, release request, dan quality gate berada pada satu agregat proyek. Pemusatan ini mengurangi kebutuhan mencocokkan status dari beberapa media dan memungkinkan pengguna melihat konteks sebelum mengambil tindakan.")
    add_para(doc, "Tujuan kedua adalah menegakkan alur dan pembagian wewenang. Keberadaan 27 status saja tidak cukup; pencapaian tujuan ditentukan oleh transisi terpusat, role permission, cakupan proyek, assignment, serta validasi prasyarat. Karena ProjectWorkflowService menjadi satu-satunya jalur transisi, aturan tidak bergantung pada tombol yang terlihat di front end. Hal ini membuat workflow memiliki sifat enforceable, bukan hanya informatif.")
    add_para(doc, "Tujuan ketiga berkaitan dengan pengujian, approval, dan audit. Wizard SIT/UAT, putaran persetujuan, jalur QA/Siber, return round, Document Vault, histori status, dan activity log membentuk bukti berlapis. Hasil tersebut memenuhi sasaran desain. Walaupun demikian, keberhasilan operasional masih memerlukan data pengguna nyata, SOP, infrastruktur produksi, pengujian ulang, dan evaluasi setelah sistem digunakan.")

    add_heading(doc, "4.2.2 Kesesuaian dengan SDLC dan Literatur", 3)
    add_para(doc, "ISO/IEC/IEEE 12207 menempatkan lifecycle sebagai kerangka proses yang dapat disesuaikan dan dilaksanakan berulang atau paralel [15]. NagariSDLC sejalan dengan prinsip tersebut karena tidak memaksakan satu model delivery. Task serta perbaikan dapat iteratif, sementara keputusan formal tetap menggunakan gate. QA dan Siber berjalan paralel; SIT dan UAT berulang ketika revisi mayor terjadi. Struktur ini lebih tepat disebut workflow governance hibrida daripada waterfall murni.")
    add_para(doc, "Prinsip agile menekankan kolaborasi dan respons terhadap perubahan [25], sedangkan penelitian Dybå dan Dingsøyr mengingatkan bahwa manfaat metode perlu dipahami sesuai konteks [26]. Implementasi NagariSDLC menerima perubahan melalui Change Request, task revisi, return round, dan supersession. Pada saat yang sama, sistem menahan approval ketika artefak berubah. Dengan cara ini, adaptasi tidak menghapus kebutuhan akan bukti dan akuntabilitas.")
    add_para(doc, "Continuous software engineering berupaya mengurangi keterputusan antaraktivitas [27]. NagariSDLC belum mengotomatisasi pipeline deployment, tetapi telah mengurangi keterputusan informasi: task menjadi basis pengujian, hasil pengujian menjadi prasyarat rilis, dan rilis memiliki keputusan eksplisit. Integrasi CI/CD di masa depan akan bernilai apabila pipeline menyampaikan evidence kembali ke project record, bukan berjalan sebagai kanal yang terpisah.")

    add_heading(doc, "4.2.3 Analisis Workflow dan State Machine", 3)
    add_para(doc, "Kompleksitas workflow berasal dari kombinasi fase linear, revisi, hold, penolakan, dua jalur paralel, dan status legacy. Model satu kolom status proyek tidak cukup untuk menyatakan kemajuan QA dan Siber secara bersamaan. Solusinya adalah mempertahankan status proyek sebagai fase global serta qa_status dan cyber_status sebagai state lokal. PENDING_GOLIVE hanya terbuka ketika kedua state lokal PASSED.")
    add_para(doc, "Pemisahan state global dan lokal mengurangi ledakan kombinasi status. Jika setiap kombinasi QA–Siber dibuat menjadi status proyek, jumlah status akan bertambah dan maknanya sulit dipahami. Trade-off-nya adalah setiap endpoint harus menilai lebih dari satu kolom. TestingTrackService dan validator workflow berfungsi sebagai tempat konsolidasi aturan tersebut.")
    add_para(doc, "READY_FOR_UAT dan UAT_PASSED dipertahankan untuk kompatibilitas histori, bukan sebagai bukti bahwa alur aktif masih memiliki UAT final. Keputusan ini penting karena menghapus nilai lama dapat merusak audit, sedangkan mengaktifkannya kembali akan menimbulkan alur yang tidak sesuai. Dokumentasi membedakan status legacy, opsional, aktif, dan terminal agar pembaca tidak menyimpulkan alur hanya dari nama enum.")

    add_heading(doc, "4.2.4 Analisis Kontrol Akses dan Separation of Duties", 3)
    add_para(doc, "Model RBAC mempermudah pengelolaan permission melalui role [21]. NagariSDLC menambahkan project scope karena role tidak menyatakan hubungan seseorang dengan proyek tertentu. Seorang developer hanya boleh melihat atau bertindak pada proyek yang ditugaskan; tester juga dibatasi oleh disposisi jalur. Kombinasi role dan relation-based scope lebih sempit daripada akses role global.")
    add_para(doc, "Separation of duties terlihat pada pemisahan pelaksana, reviewer, dan pengambil keputusan. QA Tester membuat laporan, QA Lead melakukan sign-off; Pentester menguji, Cyber Lead melakukan sign-off; PM mengajukan rilis, Head of IT memutuskan quality gate. Persetujuan UAT dibagi ke sisi peminta dan IT dengan keputusan individual. Pemisahan ini menurunkan risiko satu orang membuat, memeriksa, dan menyetujui hasilnya sendiri.")
    add_para(doc, "Keterbatasannya adalah daftar role otorisasi masih tetap di kode. Administrasi dapat membuat role baru, tetapi role tersebut tidak otomatis memperoleh semantics fase. Kondisi ini aman dari sudut default-deny, tetapi kurang fleksibel untuk organisasi yang sering mengubah struktur. Pengembangan permission engine dinamis harus dilakukan hati-hati karena kesalahan konfigurasi dapat memiliki dampak lebih besar daripada daftar eksplisit.")

    add_heading(doc, "4.2.5 Analisis SIT, UAT, QA, dan Keamanan Siber", 3)
    add_para(doc, "ISO/IEC/IEEE 29119 menekankan proses pengujian yang dapat ditata dari perencanaan sampai penyelesaian [17]. NagariSDLC menerjemahkannya ke artefak praktis: skenario, pelaksana, hasil, bukti, komentar, approval, laporan, dan status. Draft dipisahkan dari final submission agar pengguna dapat menyimpan pekerjaan tanpa memicu transisi. Pembekuan bukti setelah lulus mencegah hasil yang telah disetujui berubah tanpa putaran baru.")
    add_para(doc, "Perbedaan revisi minor dan mayor pada UAT memiliki implikasi proses. Revisi minor mempertahankan konteks UAT dan membuka task perbaikan; approval ditahan sampai task resolved. Revisi mayor menganggap perubahan cukup besar untuk membatalkan relevansi hasil sebelumnya, sehingga SIT diulang penuh dan UAT dimulai dari tahap pertama. Keputusan ini lebih mahal, tetapi menjaga validitas bukti integrasi dan penerimaan.")
    add_para(doc, "SSDF, WSTG, dan SP 800-115 menekankan pengujian keamanan yang terencana, analisis temuan, mitigasi, serta integrasi keamanan dalam lifecycle [7], [18], [19]. Jalur Siber memenuhi kerangka proses dasar tersebut melalui disposisi, tested scenarios, laporan, sign-off, return round, dan task perbaikan. Namun kontrol ini belum setara dengan program AppSec lengkap tanpa threat modeling, baseline ASVS, dependency scanning, SAST/DAST, secret scanning, dan vulnerability response operasional.")

    add_heading(doc, "4.2.6 Analisis Keamanan Aplikasi dan Audit Trail", 3)
    add_para(doc, "Penggunaan cookie HttpOnly mengurangi paparan token terhadap JavaScript, tetapi tidak menghilangkan risiko XSS maupun CSRF [23]. Karena itu kontrol perlu dipandang sebagai defense in depth: validasi input, output encoding, kebijakan origin, X-Requested-With, SameSite yang tepat, TLS, pembatasan sesi, serta otorisasi server tetap dibutuhkan. Bearer compatibility memperluas skenario penggunaan sekaligus memperluas cakupan pengujian.")
    add_para(doc, "Audit trail memiliki dua dimensi: kelengkapan informasi dan perlindungan histori. NagariSDLC mencatat aktor, waktu, status, keputusan, laporan, serta bukti. Perlindungan dilakukan melalui soft delete, foreign key RESTRICT, status superseded/revoked, dan larangan penghapusan evidence yang masih dirujuk. Keputusan approved pada putaran lama tetap tersimpan sehingga kronologi tidak direvisi secara retrospektif.")
    add_para(doc, "Masih diperlukan kebijakan retensi formal: berapa lama dokumen disimpan, siapa yang boleh memulihkan soft delete, kapan data boleh dimusnahkan, dan bagaimana ekspor audit dilakukan. Tanpa kebijakan tersebut, perlindungan teknis dapat menyebabkan akumulasi data tanpa aturan. Kebijakan perlu mempertimbangkan kebutuhan audit internal, klasifikasi informasi, dan keputusan resmi instansi.")

    add_heading(doc, "4.2.7 Analisis Data, Transaksi, dan Kompatibilitas", 3)
    add_para(doc, "Model data 17 tabel domain dan delapan tabel framework menggambarkan pemisahan antara data bisnis dan fasilitas infrastruktur aplikasi. Project menjadi pusat relasi. Approval UAT dinormalisasi karena membutuhkan query per orang, snapshot putaran, token, dan status yang konsisten. Data wizard SIT/UAT tetap pada JSON karena strukturnya berkembang dan kompatibilitas baris lama masih dibutuhkan.")
    add_para(doc, "Keputusan mempertahankan JSON mempunyai trade-off. Penambahan field lebih mudah, tetapi constraint dan query lintas key lebih lemah. Oleh karena itu key yang memengaruhi workflow tidak boleh dibaca bebas dari berbagai tempat. Contohnya, penanda UAT restart hanya dibaca melalui Project::isUatRestartPending(), sedangkan task approval dinormalisasi agar key task_ selalu konsisten.")
    add_para(doc, "Transaksi database menjaga atomicity pada operasi multi-tabel. Tanpa transaksi, proyek dapat berpindah status sementara histori atau notifikasi gagal tersimpan. Constraint foreign key melengkapi validasi aplikasi. Kombinasi service transaction dan referential integrity memberi dua lapis perlindungan, walaupun pengujian concurrency dan idempotency tetap diperlukan untuk skenario request ganda.")

    add_heading(doc, "4.2.8 Analisis Antarmuka dan Pengalaman Pengguna", 3)
    add_para(doc, "Antarmuka workflow harus membantu pengguna menjawab tiga hal: kondisi saat ini, alasan kondisi tersebut, dan aksi berikutnya. Dashboard, detail proyek, Kanban, wizard, matriks approval, quality gate, notifikasi, dan activity log membagi informasi sesuai konteks. Penggunaan Bahasa Indonesia mengurangi beban terminologi bagi pengguna internal, sedangkan istilah teknis dipertahankan ketika menjadi nama domain resmi seperti SIT dan UAT.")
    add_para(doc, "Wizard SIT/UAT mengurangi kepadatan satu formulir besar, tetapi berpotensi menyembunyikan ketergantungan antarbagian. Karena itu ringkasan, status penyimpanan draft, indikator read-only, dan blocker perlu terlihat. Pesan error server harus diterjemahkan menjadi arahan yang dapat dilakukan, bukan hanya kode status. can_resubmit dan resubmit_blocker dari server adalah contoh pengiriman alasan bisnis ke UI.")
    add_para(doc, "Evaluasi usability formal belum dilakukan. Berdasarkan kerangka Nielsen [34], pekerjaan lanjutan perlu menilai learnability, efficiency, error recovery, consistency, dan satisfaction melalui task-based usability test. Accessibility juga perlu diperiksa untuk navigasi keyboard, label form, kontras, focus state, tabel kompleks, dan pesan validasi. Laporan hanya menyatakan dukungan desain yang tersedia, bukan hasil pengujian pengguna.")

    add_heading(doc, "4.2.9 Kelebihan Sistem", 3)
    add_list(doc, [
        "Alur proyek terpusat dengan 27 status dan otorisasi yang terdokumentasi.",
        "SIT/UAT bertahap dengan bukti dan approval individual.",
        "QA dan keamanan siber paralel tanpa status gabungan yang ambigu.",
        "Perlindungan audit trail melalui soft delete, supersession, dan larangan hard delete bukti.",
        "Kontrak API konsisten dan satu lapisan layanan pada front end.",
        "Kompatibilitas data lama dipertahankan tanpa menghidupkan kembali alur yang dipensiunkan.",
    ])
    add_para(doc, "Kelebihan utama bukan jumlah modul, melainkan hubungan antarmodul. Status terkait dengan task; task terkait dengan SIT/UAT; hasil UAT terkait dengan approver; kegagalan QA/Siber terkait dengan return round dan task perbaikan; rilis terkait dengan dua jalur lulus serta keputusan Head of IT. Hubungan tersebut membentuk traceability chain yang lebih bernilai daripada kumpulan fitur yang berdiri sendiri.")

    add_heading(doc, "4.2.10 Kekurangan dan Batasan", 3)
    add_list(doc, [
        "Target database, realtime, queue, object storage, backup, monitoring, dan SLA produksi belum diputuskan.",
        "Reverb belum aktif pada template environment yang memakai broadcast log.",
        "Role baru yang dibuat melalui administrasi belum berfungsi tanpa perubahan daftar otorisasi di kode.",
        "menu_access hanya menyaring tampilan dan tidak menutup route secara mandiri.",
        "Chat masih memakai polling dan belum menggunakan WebSocket.",
        "Beberapa data SIT/UAT tetap berbentuk JSON karena kebutuhan kompatibilitas dan evolusi fitur.",
    ])
    add_para(doc, "Batasan tersebut menunjukkan perbedaan antara software yang telah memiliki fungsi inti dan layanan yang siap dioperasikan pada skala organisasi. Production readiness membutuhkan konfigurasi, kapasitas, keamanan jaringan, backup, observability, runbook, support model, dan acceptance operasional. Karena keputusan tersebut belum tersedia, laporan tidak mengklaim kesiapan produksi penuh walaupun status workflow memiliki LIVE_PRODUCTION.")

    add_heading(doc, "4.2.11 Kendala dan Solusi", 3)
    challenge_rows = [
        ["Percabangan status kompleks", "Memusatkan transisi dan prasyarat di ProjectWorkflowService."],
        ["QA/Siber paralel", "Menyimpan qa_status dan cyber_status secara independen."],
        ["Revisi UAT mayor", "Mengarsipkan siklus, SIT ulang penuh, dan UAT ulang dari Tahap 1."],
        ["Approval lintas akun", "Menyediakan internal account dan external link dengan verifikasi nomor HP."],
        ["Jejak audit berisiko hilang", "Soft delete, RESTRICT, superseded/revoked, dan larangan hapus bukti."],
        ["Data legacy", "Membaca key lama melalui helper normalisasi tanpa menulis kembali alur pensiun."],
        ["Pengajuan ulang QA/Siber", "Return round dan gerbang task perbaikan per jalur."],
    ]
    add_table(doc, ["Kendala", "Solusi"], challenge_rows, [3000, 4938], "Tabel 4.14 Kendala dan solusi", "tbl_4_14", bid); bid += 1
    add_para(doc, "Solusi pada tabel dipilih pada lapisan yang paling bertanggung jawab. Status paralel tidak diselesaikan dengan label UI, melainkan model data dan service; histori tidak dilindungi hanya dengan menyembunyikan tombol hapus, tetapi melalui aturan delete dan status supersession. Pendekatan ini mengurangi kemungkinan solusi terlewati ketika request datang dari klien berbeda.")

    add_heading(doc, "4.2.12 Potensi Pengembangan", 3)
    add_list(doc, [
        "Menetapkan arsitektur produksi, CI/CD, staging, backup, monitoring, logging, dan disaster recovery.",
        "Mengaktifkan Reverb dan queue worker setelah kebutuhan realtime serta kapasitas server diputuskan.",
        "Menambah pengujian end-to-end untuk alur kritis, terutama approval eksternal dan siklus revisi.",
        "Menentukan kebijakan retensi, terminal CANCELLED, pemulihan soft delete, dan pemusnahan data yang sah.",
        "Mengevaluasi bundle front end dan pemisahan chunk agar peringatan ukuran aset dapat dikurangi.",
        "Menyusun panduan operasional, matriks RACI resmi, dan SOP penggunaan untuk setiap peran.",
    ])
    roadmap_rows = [
        ["Prioritas 1", "Production readiness", "Environment, TLS, database, storage, backup, queue, monitoring, runbook", "Prasyarat penggunaan operasional"],
        ["Prioritas 2", "Security assurance", "Threat model, ASVS baseline, SAST/DAST, dependency/secret scan, pentest", "Risiko keamanan lebih terukur"],
        ["Prioritas 3", "Quality engineering", "E2E, performance, accessibility, usability, mutation/transition coverage", "Bukti kualitas lebih lengkap"],
        ["Prioritas 4", "Delivery integration", "CI/CD, deployment evidence, release automation, rollback drill", "Kontinuitas dari kode ke produksi"],
        ["Prioritas 5", "Governance analytics", "Lead time, defect escape, approval aging, bottleneck, audit export", "Perbaikan proses berbasis data"],
    ]
    add_table(doc, ["Urutan", "Program", "Ruang Lingkup", "Hasil Diharapkan"], roadmap_rows, [1100, 1800, 3300, 1738], "Tabel 4.15 Roadmap pengembangan", "tbl_4_15", bid); bid += 1
    add_para(doc, "Urutan roadmap menempatkan kesiapan dan keamanan sebelum fitur tambahan. Realtime atau dashboard analitik akan memberikan nilai terbatas apabila backup, monitoring, dan prosedur pemulihan belum ada. Setelah fondasi operasional tersedia, metrik governance dapat digunakan untuk mengidentifikasi bottleneck tanpa menjadikan jumlah aktivitas sebagai satu-satunya ukuran produktivitas.")

    add_heading(doc, "4.2.13 Kontribusi Mahasiswa", 3)
    contribution = add_para(doc, "Kontribusi mahasiswa mencakup pengembangan fullstack, analisis aturan bisnis, integrasi front end–back end, workflow, pengujian, dokumentasi, dan perbaikan berbasis temuan. Kontribusi teknis tidak hanya berupa penulisan halaman dan endpoint, tetapi juga menjaga konsistensi kontrak, memusatkan aturan domain, melindungi histori, serta menerjemahkan proses bisnis menjadi state dan prasyarat yang dapat diuji. Rincian harian wajib disesuaikan dengan log asli pada Lampiran E dan tidak boleh direkonstruksi tanpa bukti.")
    contribution.paragraph_format.keep_together = True

    add_chapter_heading(doc, "V", "PENUTUP")
    add_heading(doc, "5.1 Kesimpulan", 2)
    add_list(doc, [
        "NagariSDLC berhasil membentuk sistem tata kelola SDLC terpusat untuk mengelola proyek teknologi dari pengajuan sampai produksi.",
        "Mesin status, RBAC, cakupan proyek, validasi prasyarat, histori, dan notifikasi menegakkan alur serta pemisahan tanggung jawab.",
        "SIT dan UAT dikelola melalui wizard tiga tahap dengan bukti, revisi, dan approval individual yang mempertahankan audit trail.",
        "QA dan keamanan siber berjalan paralel dengan status independen, return round, task perbaikan, dan gerbang pengajuan ulang per jalur.",
        "Rilis hanya dapat diajukan setelah kedua jalur lulus dan diputuskan Head of IT pada quality gate; tidak ada UAT final setelah QA/Siber.",
        "Arsitektur Laravel–React–MySQL dan kontrak REST menyediakan fondasi pengembangan yang terstruktur, sementara keputusan infrastruktur produksi masih perlu dilengkapi.",
    ])
    add_heading(doc, "5.2 Saran", 2)
    add_list(doc, [
        "Lengkapi seluruh placeholder identitas, profil resmi, struktur organisasi, logo, foto, tanda tangan, dan bukti kegiatan sebelum pengesahan.",
        "Render kode Mermaid pada berkas pendamping, periksa keterbacaan, lalu masukkan setiap gambar pada placeholder dengan caption yang tetap.",
        "Lakukan pengujian ulang backend, ESLint, build, dan pengujian end-to-end setelah perubahan source code terakhir sebelum laporan dinyatakan final.",
        "Tetapkan konfigurasi dan SOP produksi untuk database, queue, storage, CORS, Reverb, backup, monitoring, logging, serta retensi data.",
        "Pertahankan seluruh transisi proyek melalui ProjectWorkflowService dan seluruh penulisan return round melalui ProjectReturnRoundService.",
        "Hindari hard delete pada approval, histori, laporan pengujian, dan bukti dokumen kecuali telah ada keputusan retensi resmi.",
    ])

    add_chapter_heading(doc, "", "DAFTAR PUSTAKA")
    # Ubah heading khusus agar tidak tampil sebagai BAB kosong.
    p = doc.paragraphs[-1]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(p, "DAFTAR PUSTAKA", size=14, bold=True)
    references = [
        "[1] I. Sommerville, Software Engineering, 10th ed. Pearson, 2016.",
        "[2] R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner's Approach, 9th ed. McGraw-Hill, 2020.",
        "[3] R. T. Fielding, Architectural Styles and the Design of Network-based Software Architectures. University of California, Irvine, 2000.",
        "[4] Laravel, Laravel 13.x Documentation. https://laravel.com/docs/13.x [Diakses: 28 Agustus 2026].",
        "[5] Meta Open Source, React Documentation. https://react.dev [Diakses: 28 Agustus 2026].",
        "[6] Vite, Vite Documentation. https://vite.dev [Diakses: 28 Agustus 2026].",
        "[7] OWASP Foundation, Web Security Testing Guide, Version 4.2. https://owasp.org/www-project-web-security-testing-guide/v42/ [Diakses: 28 Agustus 2026].",
        "[8] Tailwind Labs, Tailwind CSS Documentation. https://tailwindcss.com/docs [Diakses: 28 Agustus 2026].",
        "[9] ISO/IEC/IEEE 12207:2017, Systems and software engineering - Software life cycle processes. ISO, 2017.",
        "[10] ISO/IEC/IEEE 29119-1:2022, Software and systems engineering - Software testing - Part 1: General concepts. ISO, 2022. https://www.iso.org/standard/81291.html [Diakses: 28 Agustus 2026].",
        "[11] Laravel, Laravel Sanctum Documentation. https://laravel.com/docs/13.x/sanctum [Diakses: 28 Agustus 2026].",
        "[12] Oracle, MySQL Reference Manual. https://dev.mysql.com/doc/ [Diakses: 28 Agustus 2026].",
        "[13] Mermaid, Mermaid Documentation. https://mermaid.js.org [Diakses: 28 Agustus 2026].",
        "[14] Tim Proyek NagariSDLC, Dokumentasi internal: AI_HANDOFF, PROJECT_SUMMARY, WORKFLOW, ARCHITECTURE, DATA_MODEL, API_REFERENCE, dan PRD, snapshot 26 Agustus 2026.",
        "[15] ISO/IEC/IEEE 12207:2026, Systems and software engineering - Software life cycle processes. ISO, 2026. https://www.iso.org/standard/90219.html [Diakses: 28 Agustus 2026].",
        "[16] ISO/IEC 25010:2023, Systems and software engineering - Systems and software Quality Requirements and Evaluation (SQuaRE) - Product quality model. ISO, 2023. https://www.iso.org/standard/78176.html [Diakses: 28 Agustus 2026].",
        "[17] ISO/IEC/IEEE 29119-2:2021, Software and systems engineering - Software testing - Part 2: Test processes. ISO, 2021. https://www.iso.org/standard/79428.html [Diakses: 28 Agustus 2026].",
        "[18] M. Souppaya, K. Scarfone, and D. Dodson, Secure Software Development Framework (SSDF) Version 1.1: Recommendations for Mitigating the Risk of Software Vulnerabilities, NIST SP 800-218. National Institute of Standards and Technology, 2022. doi: 10.6028/NIST.SP.800-218.",
        "[19] K. Scarfone, M. Souppaya, A. Cody, and A. Orebaugh, Technical Guide to Information Security Testing and Assessment, NIST SP 800-115. National Institute of Standards and Technology, 2008. doi: 10.6028/NIST.SP.800-115.",
        "[20] OWASP Foundation, Application Security Verification Standard, Version 5.0.0. https://owasp.org/www-project-application-security-verification-standard/ [Diakses: 28 Agustus 2026].",
        "[21] R. S. Sandhu, E. J. Coyne, H. L. Feinstein, and C. E. Youman, 'Role-Based Access Control Models,' Computer, vol. 29, no. 2, pp. 38-47, 1996, doi: 10.1109/2.485845.",
        "[22] R. Fielding, M. Nottingham, and J. Reschke, HTTP Semantics, RFC 9110. IETF, 2022. doi: 10.17487/RFC9110.",
        "[23] A. Barth, HTTP State Management Mechanism, RFC 6265. IETF, 2011. https://www.rfc-editor.org/info/rfc6265/ [Diakses: 28 Agustus 2026].",
        "[24] Object Management Group, Unified Modeling Language, Version 2.5.1, 2017. https://www.omg.org/spec/UML/2.5.1 [Diakses: 28 Agustus 2026].",
        "[25] K. Beck et al., Manifesto for Agile Software Development and Principles behind the Agile Manifesto, 2001. https://agilemanifesto.org/ [Diakses: 28 Agustus 2026].",
        "[26] T. Dybå and T. Dingsøyr, 'Empirical studies of agile software development: A systematic review,' Information and Software Technology, vol. 50, no. 9-10, pp. 833-859, 2008, doi: 10.1016/j.infsof.2008.01.006.",
        "[27] B. Fitzgerald and K.-J. Stol, 'Continuous software engineering: A roadmap and agenda,' Journal of Systems and Software, vol. 123, pp. 176-189, 2017, doi: 10.1016/j.jss.2015.06.063.",
        "[28] L. E. Lwakatare et al., 'DevOps in practice: A multiple case study of five companies,' Information and Software Technology, vol. 114, pp. 217-230, 2019, doi: 10.1016/j.infsof.2019.06.010.",
        "[29] B. W. Boehm, 'A Spiral Model of Software Development and Enhancement,' Computer, vol. 21, no. 5, pp. 61-72, 1988, doi: 10.1109/2.59.",
        "[30] L. Bass, P. Clements, and R. Kazman, Software Architecture in Practice, 4th ed. Addison-Wesley Professional, 2021.",
        "[31] M. Fowler, UML Distilled: A Brief Guide to the Standard Object Modeling Language, 3rd ed. Addison-Wesley Professional, 2004.",
        "[32] J. Humble and D. Farley, Continuous Delivery: Reliable Software Releases through Build, Test, and Deployment Automation. Addison-Wesley Professional, 2010.",
        "[33] N. Forsgren, J. Humble, and G. Kim, Accelerate: The Science of Lean Software and DevOps. IT Revolution Press, 2018.",
        "[34] J. Nielsen, Usability Engineering. Academic Press, 1993.",
        "[35] Bank Nagari, 'Visi dan Misi' dan profil kantor pusat. https://www.banknagari.co.id/profile?page=hkcukYGoSeEHiTNFnflSWg%3D%3D [Diakses: 28 Agustus 2026].",
        "[36] Bank Nagari, 'Sejarah.' https://banknagari.co.id/profile?page=G1lnugtlDJSwW%2FaHA5UGAQ%3D%3D [Diakses: 28 Agustus 2026].",
    ]
    for ref in references:
        p = add_para(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False, after=4, line=1.15)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)


def build_appendices(doc: Document):
    bid = 3000
    p = add_chapter_heading(doc, "", "LAMPIRAN")
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_run(p, "LAMPIRAN", size=14, bold=True)

    app = add_heading(doc, "Lampiran A - Daftar Placeholder Tangkapan Layar", 2)
    app.paragraph_format.page_break_before = True
    add_bookmark(app, "app_a", bid); bid += 1
    add_para(doc, "Gunakan daftar berikut sebagai checklist penggantian placeholder. Setiap gambar harus berasal dari environment yang diizinkan, menampilkan data uji, dipotong secara proporsional, dan tidak memperlihatkan token, kredensial, nomor telepon, surel internal, dokumen rahasia, atau informasi pribadi.")
    screenshot_rows = [[str(i), title, f"Gambar 4.{15+i}"] for i, title in enumerate([
        "Login", "Dashboard", "Daftar proyek", "Detail proyek", "Papan Kanban", "Detail task", "Wizard SIT", "Wizard UAT",
        "Tugas QA", "Tugas keamanan siber", "Matriks persetujuan UAT", "Quality Gate", "Manajemen pengguna",
        "Manajemen grup dan role", "Document Vault", "Activity Log",
    ], 1)]
    add_table(doc, ["No.", "Tangkapan Layar", "Caption"], screenshot_rows, [650, 4300, 2988])

    app = add_heading(doc, "Lampiran B - Indeks Kode Mermaid", 2)
    app.paragraph_format.page_break_before = True
    add_bookmark(app, "app_b", bid); bid += 1
    add_para(doc, "Kode lengkap berada pada berkas pendamping Kode-Mermaid-Laporan-KP-NagariSDLC.md. Salin satu blok mermaid ke https://mermaid.live, atur tema netral, lalu ekspor sebagai SVG atau PNG beresolusi tinggi. Gunakan SVG bila Word yang digunakan dapat mempertahankan kualitas vektor.")
    mermaid_rows = [
        ["D-01", "Alur end-to-end NagariSDLC", "Gambar 4.2"],
        ["D-01A", "Proses sebelum sistem terpusat", "Gambar 4.1"],
        ["D-02", "Struktur unit representatif", "Gambar 2.4"],
        ["D-03", "Arsitektur sistem", "Gambar 4.3"],
        ["D-04", "Use case aktor utama", "Gambar 4.4"],
        ["D-05", "State machine proyek", "Gambar 4.5"],
        ["D-06", "Alur SIT", "Gambar 4.6"],
        ["D-07", "Alur UAT", "Gambar 4.7"],
        ["D-08", "QA dan keamanan siber", "Gambar 4.8"],
        ["D-09", "Relasi approver UAT", "Gambar 4.9"],
        ["D-10", "Putaran pengembalian", "Gambar 4.10"],
        ["D-11", "ERD NagariSDLC", "Gambar 4.11"],
        ["D-12", "Class relationship inti", "Gambar 4.12"],
        ["D-13", "Sequence login", "Gambar 4.13"],
        ["D-14", "Sequence approval UAT eksternal", "Gambar 4.14"],
        ["D-15", "Sitemap halaman", "Gambar 4.15"],
    ]
    add_table(doc, ["Kode", "Diagram", "Penempatan"], mermaid_rows, [1200, 4400, 2338])

    app = add_heading(doc, "Lampiran C - Endpoint API Representatif", 2)
    app.paragraph_format.page_break_before = True
    add_bookmark(app, "app_c", bid); bid += 1
    endpoints = [
        ["POST", "/auth/login", "Login dan pembuatan sesi"],
        ["POST", "/auth/logout", "Mencabut sesi aktif"],
        ["POST", "/auth/refresh", "Rotasi token sesi"],
        ["POST", "/auth/forgot-password", "Permintaan reset tanpa enumerasi akun"],
        ["POST", "/auth/reset-password", "Reset kata sandi dan cabut seluruh token"],
        ["GET/POST", "/projects", "Daftar terscope / pengajuan proyek"],
        ["GET/PATCH", "/projects/{id}", "Detail / perubahan proyek"],
        ["PATCH", "/projects/{id}/status", "Transisi status resmi"],
        ["POST", "/projects/{id}/team", "Alokasi tim"],
        ["GET", "/projects/{id}/timeline", "Riwayat status"],
        ["GET", "/projects/{id}/sit-gate", "Prasyarat masuk SIT"],
        ["POST", "/projects/{id}/sit-approval", "Approval SIT"],
        ["PUT", "/projects/{id}/uat-execution/draft", "Simpan draft eksekusi UAT"],
        ["POST", "/projects/{id}/uat-execution", "Finalisasi eksekusi UAT"],
        ["GET", "/projects/{id}/uat-approval-matrix", "Matrix approver terbaru"],
        ["POST", "/projects/{id}/uat-approval-rounds", "Membuka putaran approval"],
        ["POST", "/projects/{id}/uat-approval-rounds/sync", "Sinkronisasi peserta"],
        ["POST", "/projects/{id}/uat-approvers/{approver}/decision", "Keputusan approver internal"],
        ["GET/POST", "/projects/{projectId}/tasks", "Daftar / buat task"],
        ["PATCH/DELETE", "/tasks/{taskId}", "Ubah / hapus task"],
        ["POST", "/tasks/{taskId}/request-revision", "Permintaan revisi task"],
        ["GET/POST", "/projects/{projectId}/chat", "Diskusi proyek"],
        ["GET/POST", "/documents", "Daftar / unggah dokumen"],
        ["GET", "/documents/{id}/download", "Unduh dokumen terotorisasi"],
        ["DELETE", "/documents/{id}", "Hapus dengan proteksi audit trail"],
        ["POST", "/qa-requests/submit|assign|report|sign-off", "Empat langkah QA"],
        ["POST", "/cyber-requests/submit|assign|report|sign-off", "Empat langkah keamanan siber"],
        ["GET/POST", "/release-requests", "Daftar / pengajuan rilis"],
        ["GET", "/quality-gate/queue", "Antrean quality gate"],
        ["POST", "/quality-gate/approve|reject", "Keputusan rilis"],
        ["GET", "/dashboard/summary", "Ringkasan sesuai scope"],
        ["GET", "/dashboard/analytics", "Analitik global Super Admin"],
        ["GET/PATCH", "/notifications", "Notifikasi dan status baca"],
        ["GET", "/activity-logs", "Jejak aktivitas"],
        ["GET/POST/PATCH/DELETE", "/users|roles|groups|divisions", "Master dan administrasi"],
        ["GET", "/health", "Pemeriksaan koneksi database"],
    ]
    add_table(doc, ["Method", "Endpoint", "Fungsi"], endpoints, [1200, 3650, 3088], "Tabel C.1 Daftar endpoint API", "tbl_c_1", bid); bid += 1

    app = add_heading(doc, "Lampiran D - Kamus Data Ringkas", 2)
    app.paragraph_format.page_break_before = True
    add_bookmark(app, "app_d", bid); bid += 1
    data_dict = [
        ["projects.req_id", "string", "Kode REQ-YYYY-NNN"],
        ["projects.status", "enum/string", "27 status ProjectStatus"],
        ["projects.qa_status", "enum/string", "Status jalur QA"],
        ["projects.cyber_status", "enum/string", "Status jalur keamanan siber"],
        ["projects.sit_uat_data", "json", "Data aktif dan histori wizard SIT/UAT"],
        ["projects.priority", "string", "High, Medium, atau Low"],
        ["project_tasks.status", "enum/string", "todo, in_progress, hold, done, take_down"],
        ["project_tasks.return_round_id", "FK nullable", "Task perbaikan putaran QA/Siber"],
        ["roles.menu_access", "json nullable", "Pembatasan menu yang bersifat mengurangi"],
        ["test_reports.tested_scenarios", "text nullable", "Ruang lingkup pengujian bebas"],
        ["test_reports.evidence_document_ids", "json", "ID Document Vault untuk bukti"],
        ["uat_approval_rounds.status", "string", "active, completed, superseded"],
        ["uat_approvers.decision_status", "string", "pending, approved, rejected, revoked"],
        ["project_return_rounds.status", "string", "OPEN atau RESUBMITTED"],
        ["release_requests.rollback_plan", "text", "Rencana pemulihan rilis"],
    ]
    add_table(doc, ["Field", "Tipe", "Keterangan"], data_dict, [2800, 1500, 3638], "Tabel D.1 Kamus data ringkas", "tbl_d_1", bid); bid += 1

    app = add_heading(doc, "Lampiran E - Log Kegiatan Kerja Praktek", 2)
    app.paragraph_format.page_break_before = True
    add_bookmark(app, "app_e", bid); bid += 1
    add_para(doc, "Isi tabel berikut menggunakan catatan harian asli. Jangan mengubah jadwal mingguan pada Bab II menjadi seolah-olah kegiatan harian apabila tanggal dan bukti tidak tersedia.")
    log_rows = [[str(i), "[ HARI/TANGGAL ]", "[ URAIAN KEGIATAN FAKTUAL ]", "[ NAMA/FILE DOKUMENTASI ]"] for i in range(1, 25)]
    add_table(doc, ["No.", "Hari/Tanggal", "Kegiatan", "Dokumentasi"], log_rows, [600, 1500, 3838, 2000], "Tabel E.1 Log kegiatan", "tbl_e_1", bid); bid += 1

    app = add_heading(doc, "Lampiran F - Formulir Bimbingan dan Penilaian", 2)
    app.paragraph_format.page_break_before = True
    add_bookmark(app, "app_f", bid); bid += 1
    add_figure_placeholder(doc, "Lampiran F.1 Formulir bimbingan Kerja Praktek", "app_f_fig_1", bid, "Masukkan hasil pemindaian formulir bimbingan yang lengkap dan ditandatangani.", lines=14); bid += 1
    add_figure_placeholder(doc, "Lampiran F.2 Lembar penilaian pembimbing lapangan", "app_f_fig_2", bid, "Masukkan hasil pemindaian lembar penilaian pembimbing lapangan.", lines=14)


if __name__ == "__main__":
    build()
